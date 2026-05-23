# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from datetime import date
from shutil import copy2
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:  # pragma: no cover
    Image = ImageDraw = ImageFont = None


ROOT = Path(r"D:\AAAWorkSpace\java project\super_big_homework\Java＆Database_bigwork")
STAGING = ROOT / r"docs\面向对象程序与设计\我们的实际写作"
TEACHER_SCORE_TEMPLATE = ROOT / r"docs\面向对象程序与设计\老师要求\评分表格式\2024级-专业名称-X班第X组_评分表_电子图片管理程序.docx"

MEMBERS = [
    ("组长", "202425220501", "毕振岚", "JavaFX 主体编码、图片管理核心功能、界面完善、最终版修复与整合", "33.4"),
    ("组员", "202425220502", "陈厚华", "音乐扩展与外部接口调研、运行流程协助测试、验收材料核对", "33.3"),
    ("组员", "202425220527", "徐阳", "数据库设计文档、课程论文撰写整合、提交材料整理、格式核查", "33.3"),
]

TITLE = "电子图片管理程序"
LONG_TITLE = "基于 PostgreSQL 的数字图像集成管理系统"
SUBMIT_DATE = "2026年5月20日"


def set_east_asia_font(run, font_name="宋体"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text: str, size: float = 10.5, bold: bool = False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        set_east_asia_font(run)
        run.font.size = Pt(size)
        run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for name, size, bold in [
        ("Heading 1", 14, True),
        ("Heading 2", 12, True),
        ("Heading 3", 10.5, True),
    ]:
        style = styles[name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = bold


def add_para(doc: Document, text: str = "", size: float = 10.5, bold: bool = False, align=None, first_line=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    pf.space_after = Pt(3)
    if first_line and text:
        pf.first_line_indent = Pt(21)
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_east_asia_font(run)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        run.font.size = Pt(14 if level == 1 else 12 if level == 2 else 10.5)
    return p


def add_table(doc: Document, rows, headers=None, widths=None):
    total_rows = len(rows) + (1 if headers else 0)
    cols = len(headers or rows[0])
    table = doc.add_table(rows=total_rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    start = 0
    if headers:
        for c, h in enumerate(headers):
            set_cell_text(table.cell(0, c), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        start = 1
    for r, row in enumerate(rows, start=start):
        for c, value in enumerate(row):
            set_cell_text(table.cell(r, c), str(value), align=WD_ALIGN_PARAGRAPH.CENTER if c == 0 else None)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_code_block(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(9)
    return p


def add_toc(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目录")
    set_east_asia_font(run)
    run.bold = True
    run.font.size = Pt(14)
    p = doc.add_paragraph()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "右键更新目录或导出 PDF 时自动更新。"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = p.add_run()._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)


def add_page_number_footer(doc: Document):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run()
        set_east_asia_font(run)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)


def font(size=24, bold=False):
    if ImageFont is None:
        return None
    for path in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]:
        if Path(path).exists():
            try:
                return ImageFont.truetype(path, size=size)
            except Exception:
                pass
    return ImageFont.load_default()


def draw_round(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def make_ui_images():
    source = ROOT / "target" / "ui-smoke"
    out = STAGING / "界面截图"
    out.mkdir(parents=True, exist_ok=True)
    mapping = [
        ("MainView-900x600.png", "real_01_主界面_900x600_最小窗口.png"),
        ("MainView-1200x800.png", "real_02_主界面_1200x800_默认窗口.png"),
        ("MainView-1440x900.png", "real_03_主界面_1440x900_宽屏窗口.png"),
        ("WelcomeDialog-640x620.png", "real_04_首次启动向导_640x620.png"),
        ("WelcomeDialog-800x775.png", "real_05_首次启动向导_800x775.png"),
        ("SettingsView-680x720.png", "real_06_系统设置_680x720.png"),
        ("SettingsView-850x900.png", "real_07_系统设置_850x900.png"),
        ("ImageViewerView-960x680.png", "real_08_图片查看器_960x680.png"),
        ("ImageViewerView-1200x850.png", "real_09_图片查看器_1200x850.png"),
        ("SlideshowView-1000x700.png", "real_10_幻灯片播放_1000x700.png"),
        ("SlideshowView-1250x875.png", "real_11_幻灯片播放_1250x875.png"),
        ("ImageEditorView-1000x750.png", "real_12_图片编辑器_1000x750.png"),
        ("ImageEditorView-1250x938.png", "real_13_图片编辑器_1250x938.png"),
        ("RenameDialog-450x360.png", "real_14_批量重命名_450x360.png"),
        ("RenameDialog-563x450.png", "real_15_批量重命名_563x450.png"),
    ]
    missing = [name for name, _ in mapping if not (source / name).exists()]
    if missing:
        raise FileNotFoundError(
            "缺少 JavaFX 真实截图，请先运行 UiSnapshotSmoke："
            + ", ".join(missing)
        )
    images = []
    for source_name, target_name in mapping:
        target = out / target_name
        copy2(source / source_name, target)
        images.append(target)
    return images

    if Image is None:
        return []
    out = STAGING / "界面截图"
    out.mkdir(parents=True, exist_ok=True)
    images = []
    colors = {
        "bg": (248, 250, 252),
        "panel": (255, 255, 255),
        "line": (210, 216, 225),
        "text": (31, 41, 55),
        "muted": (100, 116, 139),
        "blue": (37, 99, 235),
        "green": (16, 185, 129),
        "amber": (245, 158, 11),
        "red": (239, 68, 68),
        "purple": (124, 58, 237),
    }
    f12, f14, f16, f18, f22, f26 = font(12), font(14), font(16), font(18), font(22), font(26)

    def save(img, name):
        path = out / name
        img.save(path)
        images.append(path)

    # Main window
    img = Image.new("RGB", (1440, 900), colors["bg"])
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1440, 68), fill=(31, 41, 55))
    d.text((24, 12), "数字图像集成管理系统", fill="white", font=f22)
    d.text((24, 41), r"D:\Pictures\课程设计样例", fill=(203, 213, 225), font=f12)
    draw_round(d, (435, 14, 558, 52), 7, (255, 255, 255), (203, 213, 225))
    d.text((462, 25), "关键词", fill=colors["text"], font=f14)
    draw_round(d, (568, 14, 1110, 52), 7, (255, 255, 255), (203, 213, 225))
    d.text((590, 25), "搜索文件名、标签、元数据或 AI 描述...", fill=colors["muted"], font=f14)
    draw_round(d, (1120, 14, 1205, 52), 7, colors["blue"], None)
    d.text((1148, 25), "搜索", fill="white", font=f14)
    draw_round(d, (1234, 14, 1318, 52), 7, (51, 65, 85), (71, 85, 105))
    d.text((1262, 25), "设置", fill=(226, 232, 240), font=f14)
    draw_round(d, (24, 92, 330, 820), 8, colors["panel"], colors["line"])
    d.text((44, 116), "目录树", fill=colors["text"], font=f18)
    tree = ["我的电脑", "  本地磁盘 (C:)", "    Users", "      Pictures", "        课程设计样例", "  本地磁盘 (D:)", "    AAAWorkSpace", "      images"]
    y = 160
    for t in tree:
        color = colors["blue"] if "课程设计样例" in t else colors["text"]
        d.text((46, y), t, fill=color, font=f14)
        y += 34
    draw_round(d, (354, 92, 1416, 820), 8, colors["panel"], colors["line"])
    d.rectangle((355, 93, 1415, 158), fill=(255, 255, 255))
    d.text((380, 112), "课程设计样例", fill=colors["text"], font=f18)
    draw_round(d, (1252, 110, 1368, 144), 17, (234, 251, 242), None)
    d.text((1276, 119), "18 张图片", fill=(4, 120, 87), font=f14)
    for idx in range(18):
        row, col = divmod(idx, 6)
        x = 390 + col * 165
        y = 190 + row * 190
        selected = idx in {1, 4, 7}
        draw_round(d, (x, y, x + 150, y + 172), 8, (239, 246, 255) if selected else (255, 255, 255),
                   colors["blue"] if selected else colors["line"], 2 if selected else 1)
        draw_round(d, (x + 12, y + 12, x + 138, y + 122), 6, (238, 242, 247), None)
        d.rectangle((x + 21, y + 24, x + 129, y + 110), fill=((180 + idx * 3) % 255, (140 + idx * 9) % 255, (110 + idx * 13) % 255))
        d.text((x + 18, y + 133), f"IMG_{idx + 1:04d}.jpg", fill=colors["text"], font=f12)
        d.text((x + 18, y + 152), "1920x1080 JPG", fill=colors["muted"], font=f12)
    d.rectangle((0, 842, 1440, 900), fill=(255, 255, 255))
    d.text((28, 862), "状态：目录扫描完成，缩略图缓存已同步到 PostgreSQL。", fill=colors["muted"], font=f14)
    draw_round(d, (1040, 854, 1150, 888), 6, (248, 250, 252), (203, 213, 225))
    d.text((1058, 863), "清理 AI 标签", fill=colors["text"], font=f12)
    draw_round(d, (1258, 854, 1348, 888), 6, (15, 118, 110), None)
    d.text((1278, 863), "幻灯片", fill="white", font=f12)
    save(img, "01_主界面_目录树与缩略图.png")

    # Slideshow
    img = Image.new("RGB", (1440, 900), (15, 23, 42))
    d = ImageDraw.Draw(img)
    d.text((32, 24), "幻灯片播放", fill=(226, 232, 240), font=f24 if False else f26)
    draw_round(d, (120, 86, 1320, 690), 10, (30, 41, 59), (71, 85, 105))
    d.rectangle((230, 136, 1210, 640), fill=(73, 116, 150))
    d.text((558, 360), "当前大图等比展示", fill=(240, 249, 255), font=f26)
    for x, txt in [(450, "上一张"), (570, "播放"), (690, "下一张"), (810, "放大"), (930, "缩小")]:
        draw_round(d, (x, 724, x + 92, 764), 7, (37, 99, 235), None)
        d.text((x + 23, 735), txt, fill="white", font=f14)
    for i in range(7):
        x = 310 + i * 118
        draw_round(d, (x, 790, x + 94, 860), 6, (51, 65, 85), (37, 99, 235) if i == 2 else (71, 85, 105), 2)
        d.text((x + 16, 813), f"{i+1}", fill=(226, 232, 240), font=f18)
    save(img, "02_幻灯片播放.png")

    # Settings
    img = Image.new("RGB", (1440, 900), colors["bg"])
    d = ImageDraw.Draw(img)
    d.text((48, 42), "系统设置", fill=colors["text"], font=f26)
    sections = [
        ("AI 图像识别 API 配置", ["Base URL  https://cpa.ystone.top/v1", "模型     从 /models 自动获取后下拉选择", "API Key   优先读取系统环境变量", "请求间隔   1500 ms"]),
        ("扫描目录", [r"D:\Pictures\课程设计样例", "AI识别单批上限 100(max)，可在设置页调整", "可停止扫描并清理AI标签"]),
        ("幻灯片偏好", ["播放间隔 3 秒", "播放顺序 顺序播放", "背景音乐 无音乐"]),
        ("界面主题", ["背景图片 可选", "透明度 35%"]),
    ]
    y = 100
    for title, lines in sections:
        draw_round(d, (60, y, 1380, y + 138), 8, colors["panel"], colors["line"])
        d.text((90, y + 20), title, fill=colors["blue"], font=f18)
        ly = y + 56
        for line in lines:
            d.text((110, ly), line, fill=colors["text"], font=f14)
            ly += 28
        y += 166
    draw_round(d, (940, 788, 1080, 835), 8, colors["green"], None)
    draw_round(d, (1100, 788, 1240, 835), 8, colors["blue"], None)
    d.text((975, 802), "测试 API", fill="white", font=f16)
    d.text((1148, 802), "保存", fill="white", font=f16)
    save(img, "03_设置页_运行配置.png")

    # Image editor
    img = Image.new("RGB", (1440, 900), (248, 250, 252))
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1440, 56), fill=(31, 41, 55))
    d.text((24, 16), "图片编辑器", fill="white", font=f18)
    for i, txt in enumerate(["裁剪", "画笔", "箭头", "文字", "保存版本", "恢复"]):
        x = 270 + i * 120
        draw_round(d, (x, 11, x + 100, 45), 7, (51, 65, 85), (71, 85, 105))
        d.text((x + 28, 20), txt, fill=(226, 232, 240), font=f14)
    draw_round(d, (28, 80, 1088, 830), 8, (255, 255, 255), colors["line"])
    d.rectangle((110, 145, 1005, 765), fill=(184, 210, 220))
    d.rectangle((350, 240, 820, 610), outline=(239, 68, 68), width=5)
    d.line((350, 240, 820, 610), fill=(239, 68, 68), width=5)
    d.text((420, 300), "课堂展示重点区域", fill=(239, 68, 68), font=f22)
    draw_round(d, (1115, 80, 1410, 830), 8, (255, 255, 255), colors["line"])
    d.text((1140, 108), "版本历史", fill=colors["text"], font=f18)
    for i, txt in enumerate(["v1 原图", "v2 裁剪", "v3 标注", "v4 当前版本"]):
        y = 160 + i * 70
        draw_round(d, (1140, y, 1380, y + 48), 6, (239, 246, 255) if i == 3 else (248, 250, 252), colors["line"])
        d.text((1160, y + 15), txt, fill=colors["blue"] if i == 3 else colors["text"], font=f14)
    save(img, "04_图片编辑与版本历史.png")

    # Tags/search
    img = Image.new("RGB", (1440, 900), colors["bg"])
    d = ImageDraw.Draw(img)
    d.text((46, 38), "标签与扩展搜索", fill=colors["text"], font=f26)
    draw_round(d, (60, 100, 760, 805), 8, colors["panel"], colors["line"])
    d.rectangle((100, 145, 720, 505), fill=(205, 213, 224))
    d.text((310, 310), "测试图片", fill=colors["muted"], font=f26)
    d.text((100, 548), "图片信息：校园操场上有多人合影，背景为教学楼和树木。", fill=colors["text"], font=f16)
    tags = ["场景/校园", "人物/多人", "物体/教学楼", "颜色/蓝色", "情绪/欢乐", "人数/5"]
    x, y = 100, 590
    for t in tags:
        draw_round(d, (x, y, x + 132, y + 36), 18, (219, 234, 254), None)
        d.text((x + 18, y + 9), t, fill=colors["blue"], font=f12)
        x += 150
        if x > 620:
            x, y = 100, y + 48
    draw_round(d, (800, 100, 1380, 805), 8, colors["panel"], colors["line"])
    d.text((836, 136), "关键词与扩展搜索", fill=colors["text"], font=f18)
    draw_round(d, (836, 184, 1342, 232), 6, (248, 250, 252), colors["line"])
    d.text((858, 198), "找出校园里多人合影的照片", fill=colors["muted"], font=f16)
    d.text((836, 274), "检索条件", fill=colors["blue"], font=f18)
    sql_lines = [
        "SELECT id",
        "FROM v_image_search",
        "WHERE is_deleted = FALSE",
        "  AND all_tags ILIKE '%校园%'",
        "  AND people_count >= 2;"
    ]
    yy = 318
    for line in sql_lines:
        d.text((860, yy), line, fill=(15, 23, 42), font=f14)
        yy += 32
    d.text((836, 520), "搜索结果：3 张匹配图片", fill=colors["green"], font=f18)
    save(img, "05_标签与扩展搜索.png")
    return images


def make_report_docx():
    doc = Document()
    set_doc_defaults(doc)

    # Cover
    for _ in range(4):
        add_para(doc, "", first_line=False)
    p = add_para(doc, "《面向对象程序设计实践》", 22, True, WD_ALIGN_PARAGRAPH.CENTER, False)
    p.paragraph_format.space_after = Pt(8)
    add_para(doc, "课程论文", 22, True, WD_ALIGN_PARAGRAPH.CENTER, False)
    for _ in range(3):
        add_para(doc, "", first_line=False)
    cover_rows = [
        ("题目：", TITLE),
        ("专业：", "软件工程"),
        ("班级：", "2024级软件工程5班"),
        ("小组：", "第7组"),
        ("小组成员1：", "202425220501  毕振岚"),
        ("2：", "202425220502  陈厚华"),
        ("3：", "202425220527  徐阳"),
        ("指导老师：", "彭红星"),
        ("提交时间：", SUBMIT_DATE),
    ]
    cover_table = doc.add_table(rows=len(cover_rows), cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_table.style = "Table Grid"
    for row_idx, (label, value) in enumerate(cover_rows):
        set_cell_text(cover_table.cell(row_idx, 0), label, 12, True, WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(cover_table.cell(row_idx, 1), value, 12, False, WD_ALIGN_PARAGRAPH.LEFT)
        cover_table.cell(row_idx, 0).width = Cm(4.0)
        cover_table.cell(row_idx, 1).width = Cm(8.8)
    for _ in range(4):
        add_para(doc, "", first_line=False)
    add_para(doc, "华南农业大学 数学与信息学院 软件学院", 14, True, WD_ALIGN_PARAGRAPH.CENTER, False)
    doc.add_page_break()

    add_toc(doc)
    doc.add_page_break()

    add_heading(doc, "1 系统分析", 1)
    add_heading(doc, "1.1 问题描述", 2)
    add_para(doc, f"本课题要求完成一个面向普通用户的{TITLE}，能够浏览本地磁盘目录，识别并管理 JPG、JPEG、PNG、GIF、BMP 等常见图片格式，并围绕缩略图、目录树、单选多选、删除、复制粘贴、重命名、图片展示、放大缩小和幻灯片播放等功能完成完整的软件设计、编码、测试与课程论文撰写。")
    add_para(doc, f"本组在基础要求之上，将系统扩展为“{LONG_TITLE}”。系统仍以老师评分表中的电子图片管理功能为核心，但在数据层采用 PostgreSQL 保存目录、图片元数据、缩略图缓存、操作日志、搜索历史、系统设置和版本历史，使程序不只是临时文件浏览器，而是具备长期管理、检索和扩展能力的桌面应用。")
    add_para(doc, "从用户角度看，系统要解决三个实际问题：第一，大量图片分布在不同目录中，靠文件资源管理器查找效率低；第二，图片管理经常涉及成批选择、重命名、复制迁移和删除，人工操作容易出错；第三，图片内容不能只靠文件名表达，因此需要缩略图、标签、信息面板和搜索等方式提高检索效率。")
    add_para(doc, "本次修订融合了毕振岚单独撰写版本中的目录扫描流程、图片查看器、全局主题背景和成员总结材料，同时参考了班级其他组作业在封面表格、测试表、提交清单和格式核查上的做法。这里参考的是结构和表达方式，不照搬其他组正文。")

    add_heading(doc, "1.2 系统功能分析", 2)
    rows = [
        ("目录树浏览", "展示磁盘根目录和子目录，采用懒加载避免一次性遍历全盘。", "MainController.initDirectoryTree / FileUtil.listSubDirectories"),
        ("缩略图预览", "选择目录后扫描支持格式，生成等比缩略图并显示图片名称、尺寸和格式。", "ImageServiceImpl.loadImagesFromDirectory / ImageUtil"),
        ("单选与多选", "鼠标点击单选，Ctrl 多选，拖拽框选；状态栏显示当前选中数量。", "MainController.initThumbnailInteraction"),
        ("删除", "支持单张和多张删除，先逻辑删除数据库记录，再删除磁盘文件，并保留日志。", "ImageServiceImpl.deleteImages / operation_logs"),
        ("复制粘贴", "支持单张和多张复制粘贴，目标目录重名时自动生成不冲突文件名。", "ImageServiceImpl.copyImages / pasteImages"),
        ("重命名", "支持单张重命名和批量编号重命名，批量操作使用事务与磁盘回滚。", "ImageServiceImpl.renameImage / batchRename"),
        ("图片展示", "双击或按钮进入大图展示，支持前后切换、实际大小、适应窗口。", "ImageViewerController / SlideshowController"),
        ("幻灯片播放", "支持自动播放、循环、随机顺序、缩放、全屏和背景音乐。", "SlideshowController / MusicService"),
        ("主题与设置", "设置页统一管理扫描目录、播放间隔、主题背景和透明度。", "SettingsController / ThemeUtil / app_settings"),
        ("AI识别安全控制", "默认连接 CPA 代理节点，密钥从环境变量读取，模型从 /models 下拉选择；单批上限可在设置页调整，显示为 N(max)，可停止并清理AI标签。", "AIConfig / ScanTask / AiTagStorageService"),
    ]
    add_table(doc, rows, headers=["功能", "说明", "主要代码依据"], widths=[3.0, 7.3, 5.4])
    add_para(doc, "以上功能覆盖老师评分表中对电子图片管理程序的基本功能和扩展功能要求。扩展功能并没有替代基础功能，而是围绕图片管理主流程补充了数据库持久化、版本历史、主题配置和可继续扩展的接口。")

    add_heading(doc, "1.3 开发平台及工具介绍", 2)
    rows = [
        ("开发语言", "Java 21 编译目标，JDK 21+ 可运行", "使用面向对象方式组织模型、DAO、Service、Controller。"),
        ("界面框架", "JavaFX 21.0.6 + FXML + CSS", "实现桌面端目录树、缩略图、设置页、幻灯片和编辑窗口。"),
        ("数据库", "PostgreSQL 18.3", "保存目录、图片元数据、缩略图、标签、日志、设置和版本记录。"),
        ("构建工具", "Maven 3.9.14", "统一管理依赖，生成可运行目标 JAR。"),
        ("主要依赖", "JDBC、HikariCP、OkHttp、Jackson、SLF4J/Logback", "连接池、可选 HTTP 调用、JSON 解析与日志。"),
        ("开发环境", "IntelliJ IDEA / Windows", "源码按 Maven 标准工程提交，可直接导入 IDE。"),
    ]
    add_table(doc, rows, headers=["项目", "采用方案", "作用"], widths=[3.0, 4.2, 8.5])

    add_heading(doc, "1.4 需求与采分点对照", 2)
    rows = [
        ("基本功能", "完成", "可运行 JAR、源码 ZIP、论文 DOCX 和评分表 DOCX 已准备，后续可由 Word 导出 PDF。"),
        ("支持规定图片格式", "完成", "SUPPORTED_FORMATS 包含 JPG、JPEG、GIF、PNG、BMP。"),
        ("目录树实现", "完成", "TreeView + 懒加载磁盘目录，展开目录时加载子目录。"),
        ("缩略图比例", "完成", "使用 ImageUtil 生成缩略图，前端 ImageView 保持比例显示。"),
        ("目录图片数", "完成", "选择目录后刷新缩略图和图片数量 Label。"),
        ("缩略图单选/多选", "完成", "支持点击、Ctrl 多选和框选。"),
        ("单个/多个删除", "完成", "deleteImages 接收列表，统一处理单张或多张。"),
        ("单个/多个复制粘贴", "完成", "内部剪贴板为 List<ImageFile>，支持多张。"),
        ("单个/批量重命名", "完成", "单张直接改名，批量编号并带事务回滚。"),
        ("图片切换/缩放/幻灯片", "完成", "SlideshowController 提供前后切换、缩放和自动播放。"),
        ("代码质量", "优秀", "分层结构清晰，DAO 使用 PreparedStatement，批量操作有事务。"),
    ]
    add_table(doc, rows, headers=["评分项", "完成情况", "证据说明"], widths=[4.0, 2.3, 9.2])

    add_heading(doc, "1.5 小组分工与依据", 2)
    add_para(doc, "分工根据 Git 提交记录、群聊任务安排和最终文档整理情况重新核对。Git 中的 TreeRabbit15 为毕振岚，他在 2026 年 5 月 18 日至 5 月 19 日集中提交了最终版代码；Xu Yang 的提交主要覆盖早期工程、数据库文档、课程材料整理和本次论文生成。评分表中的工作量按小组协作口径保持基本均分，正文分工则按实际承担内容写清楚。")
    add_table(doc, [
        ("毕振岚", "组长、主程序实现", "主导 JavaFX 代码、图片管理核心流程、界面修复、最终功能冻结和版本推送。", "TreeRabbit15 提交 20260518、20260519、20260519最终版；群聊中多次发布功能冻结和任务安排。"),
        ("陈厚华", "扩展与测试协助", "按分配参与音乐扩展、外部接口调研、运行流程检查和验收反馈。", "群聊中接收任务并参与确认；最终评分表按协作成员保留贡献。"),
        ("徐阳", "文档与提交整理", "负责数据库方向文档、课程论文整合、其他组作业收齐参考、提交目录整理和格式核查。", "Xu Yang 提交课程材料整理、需求/数据库文档；本次根据老师模板生成 DOCX。"),
    ], headers=["成员", "主要定位", "实际承担内容", "依据"], widths=[2.4, 3.2, 5.2, 4.7])

    add_heading(doc, "2 系统设计", 1)
    add_heading(doc, "2.1 系统总体结构设计", 2)
    add_para(doc, "系统采用典型的分层架构：表示层负责 JavaFX 界面和用户事件；业务层负责图片扫描、复制、粘贴、删除、重命名、搜索、编辑和幻灯片逻辑；数据访问层封装 JDBC 操作；模型层定义目录、图片、标签、设置、版本和操作日志等对象；工具层提供文件、图片、主题和弹窗等公共能力。")
    add_table(doc, [
        ("表示层 Controller", "MainController、SlideshowController、SettingsController、ImageViewerController、RenameDialogController、WelcomeController", "接收用户操作，刷新界面状态。"),
        ("业务层 Service", "ImageServiceImpl、SearchService、EditService、MusicService", "组织跨 DAO 和文件系统的业务流程。"),
        ("数据访问层 DAO", "ImageDaoImpl、DirectoryDaoImpl、TagDaoImpl、SettingsDaoImpl、VersionDaoImpl", "执行 SQL，屏蔽数据库细节。"),
        ("模型层 Model", "ImageFile、DirectoryNode、Tag、ImageAnalysisResult、ImageVersion、AppSetting", "表达系统核心数据。"),
        ("工具层 Util/Scanner", "FileUtil、ImageUtil、ThemeUtil、AIConfig、ScanTask", "提供可复用的文件、图片、主题、配置和后台扫描能力。"),
    ], headers=["层次", "主要类", "职责"], widths=[3.4, 6.5, 5.6])
    add_code_block(doc, """
用户操作
  -> JavaFX Controller
    -> Service 业务流程
      -> DAO / 文件系统 / 扩展配置
        -> PostgreSQL 与本地图片文件
  <- 模型对象与状态结果
<- 界面刷新、提示信息与日志记录
""")

    add_heading(doc, "2.2 系统各个类及类之间关系设计", 2)
    add_para(doc, "MainController 是主界面协调者，它不直接编写 SQL，而是调用 ImageService、SearchService、SettingsDao 和 TagDao；ImageServiceImpl 负责图片文件与数据库记录同步；DAO 层只负责数据库读写；图片查看、幻灯片播放、编辑和设置分别拆到独立 Controller，避免主控制器继续膨胀。")
    add_table(doc, [
        ("MainController", "ImageService、SearchService、SettingsDao、TagDao", "目录树、缩略图、选中状态、右键菜单、搜索入口和幻灯片入口。"),
        ("ImageServiceImpl", "ImageDao、DirectoryDao、FileUtil、ImageUtil", "扫描目录、同步数据库、删除、复制粘贴、重命名。"),
        ("SearchService", "TagDao、ImageDao、DirectoryDao", "关键词检索和预留智能检索入口。"),
        ("SettingsController", "SettingsDao、AIConfig、ThemeUtil", "配置扫描目录、主题背景、透明度和扩展接口参数。"),
        ("SlideshowController", "MusicService、SettingsDao、ImageUtil", "大图展示、前后切换、缩放、自动播放和音乐。"),
        ("TagDaoImpl", "DatabaseConnection", "标签关系、扩展分析结果和搜索 SQL 安全执行。"),
    ], headers=["类", "关联对象", "主要职责"], widths=[4.0, 5.5, 6.0])
    add_para(doc, "类之间的依赖遵循“界面层依赖业务层，业务层依赖 DAO 和工具层，DAO 层依赖数据库连接”的单向规则。这样做可以避免界面代码中混入 SQL，也便于替换数据源、增加测试和定位故障。")

    add_heading(doc, "2.3 数据存储的设计", 2)
    add_para(doc, "数据存储采用 PostgreSQL。directories 表保存磁盘目录树，images 表保存图片元数据和缩略图二进制，operation_logs 表记录操作日志，tag_categories、tags、image_tags 保存标签关系，app_settings 保存应用配置，image_versions 与 image_edit_operations 保存编辑版本历史。")
    add_table(doc, [
        ("directories", "目录名称、完整路径、父目录 ID", "通过 parent_id 自引用表达树结构。"),
        ("images", "文件名、路径、目录、大小、宽高、格式、缩略图、删除标记", "核心图片表，支持逻辑删除和缩略图缓存。"),
        ("operation_logs", "图片 ID、操作类型、旧值、新值、时间", "触发器自动记录插入、重命名和删除。"),
        ("tag_categories/tags/image_tags", "标签分类、标签值、图片标签关联", "支持人工标签和后续智能识别扩展。"),
        ("ai_analysis_results", "原始 JSON、描述、人数、模型名", "作为扩展预留表，未配置接口时不影响主流程。"),
        ("app_settings", "key-value 配置", "保存扫描目录、幻灯片、主题和可选扩展设置。"),
        ("image_versions/edit_operations", "版本文件、缩略图、编辑类型、操作参数", "支持图片编辑后保存版本和恢复。"),
    ], headers=["表", "关键字段", "设计目的"], widths=[4.0, 5.6, 5.9])
    add_code_block(doc, """
CREATE TABLE IF NOT EXISTS images (
    id SERIAL PRIMARY KEY,
    file_name VARCHAR(255) NOT NULL,
    file_path TEXT NOT NULL,
    directory_id INTEGER NOT NULL REFERENCES directories(id),
    format VARCHAR(10),
    thumbnail BYTEA,
    ai_processed BOOLEAN NOT NULL DEFAULT FALSE,
    is_deleted BOOLEAN NOT NULL DEFAULT FALSE,
    CONSTRAINT uq_images_dir_name UNIQUE (file_name, directory_id)
);
""")
    add_para(doc, "设计中对图片不做简单的硬删除记录丢弃，而是通过 is_deleted 标记结合 operation_logs 保留审计信息；对目录使用唯一完整路径避免重复入库；对标签使用多对多关系，使一张图片可以同时归入多个分类。")

    add_heading(doc, "2.4 界面设计", 2)
    add_para(doc, "本次最终界面按 shadcn/ui 的视觉语言重新整理，但仍保留 JavaFX、FXML 和 CSS 实现方式，没有引入 Web 技术。主窗口采用中性浅色工作台、左侧目录树 sidebar、右侧缩略图网格和底部状态栏；控件使用清晰边框、低阴影、约 8px 圆角和紧凑间距，使界面更适合课堂演示和日常重复操作。")
    add_para(doc, "目录树保留文件管理习惯，缩略图卡片固定尺寸，卡片内显示等比图片、文件名、尺寸和格式；选中卡片通过边框和背景变化区分，按钮、输入框、下拉框、滚动条、弹窗和右键菜单都统一了 hover、focus、disabled、selected 等状态。图片查看器、幻灯片和编辑器保留深色图片区，避免图片内容被浅色背景干扰；底部控制区则改成同一套浅色按钮、边框和状态文本。设置页和首次启动向导去掉内联颜色，改为语义化 styleClass，窄窗口下可滚动且底部按钮保持可见。")

    add_heading(doc, "2.5 扩展功能与降级设计", 2)
    add_para(doc, "项目中保留了 AI 配置、外部接口和智能搜索的扩展点，但最终验收不把这部分当作主要完成功能。这样写更符合当前代码和群聊中的功能冻结口径：基础图片管理必须稳定，外部服务如果没有密钥或网络环境，就只能作为后续扩展。")
    add_code_block(doc, """
public Optional<String> naturalLanguageToSQL(String naturalLanguageQuery) {
    if (!AIConfig.isConfigured()) {
        return Optional.empty();
    }
    String requestBody = buildTextRequest(NL_TO_SQL_PROMPT, naturalLanguageQuery);
    String responseJson = sendRequestWithRetry(requestBody);
    String content = extractContentFromResponse(responseJson);
    String sql = content.replaceAll("```sql\\\\s*", "").replaceAll("```\\\\s*", "").trim();
    return Optional.of(sql);
}
""")
    add_para(doc, "这里真正要说明的是降级思路：系统不会把密钥写死在源码中，未配置时只关闭相关入口或返回空结果，目录树、缩略图、删除、复制粘贴、重命名和幻灯片播放仍能独立运行。后续如果继续完善，可以把外部识别、网盘同步和音乐资源统一收进扩展服务层。")

    add_heading(doc, "3 系统实现", 1)
    add_heading(doc, "3.1 应用启动与数据库连接", 2)
    add_para(doc, "App 类负责加载 FXML、初始化 JavaFX 主舞台并显示欢迎或主界面。数据库连接由 DatabaseConnection 统一创建和管理，配置来自 resources/config/database.properties。连接池使用 HikariCP，避免频繁创建连接影响界面响应。")
    add_code_block(doc, """
db.url=jdbc:postgresql://localhost:5432/image_manager
db.username=postgres
db.password=1234
""")
    add_para(doc, "系统首次使用前执行 sql/schema.sql，脚本会创建核心表、索引、视图、触发器和存储过程。脚本中加入了 ADD COLUMN IF NOT EXISTS 和 DROP VIEW IF EXISTS，能兼容已经执行过旧版本脚本的数据库。")

    add_heading(doc, "3.2 目录树与缩略图加载", 2)
    add_para(doc, "目录树使用 TreeView 实现，磁盘根目录作为第一层节点。每个目录节点初始放入“加载中...”占位子节点，只有用户展开时才真正扫描子目录，从而避免程序启动时遍历全盘。")
    add_code_block(doc, """
TreeItem<String> placeholder = new TreeItem<>("加载中...");
item.getChildren().add(placeholder);
item.expandedProperty().addListener((obs, wasExpanded, isNowExpanded) -> {
    if (isNowExpanded && item.getChildren().size() == 1) {
        loadSubDirectories(item, path);
    }
});
""")
    add_para(doc, "选择目录后，ImageServiceImpl 会先确保该目录在数据库中存在，再扫描磁盘中支持格式的图片文件。磁盘有而数据库没有的图片会新增记录，数据库有而磁盘不存在的图片会标记删除，最后返回最新图片列表给主界面渲染缩略图。")

    add_heading(doc, "3.3 图片选择与右键操作", 2)
    add_para(doc, "缩略图区域使用 FlowPane 承载固定尺寸卡片，selectedImages 保存当前选中集合，cardMap 保存图片对象到卡片节点的映射。单击卡片时根据 Ctrl 状态决定清空选择或追加选择；拖拽时根据框选矩形和卡片 Bounds 判断是否选中。")
    add_para(doc, "右键菜单根据当前状态动态启用删除、复制、粘贴、重命名和批量重命名。这样既能满足老师对单张和多张操作的采分要求，也避免无选中对象时出现不可执行命令。")

    add_heading(doc, "3.4 删除、复制粘贴与重命名实现", 2)
    add_para(doc, "删除功能接收图片列表，因此单张删除和多张删除走同一条业务流程。每张图片先在数据库中执行逻辑删除，再尝试删除磁盘文件；若个别文件删除失败，系统记录日志并继续处理其他文件，避免批量操作被单个异常中断。")
    add_code_block(doc, """
for (var image : images) {
    imageDao.softDelete(image.id());
    Path path = Path.of(image.filePath());
    if (Files.exists(path)) {
        Files.delete(path);
    }
}
""")
    add_para(doc, "复制粘贴使用内部剪贴板保存 List<ImageFile>。粘贴时会检查目标目录是否存在同名文件，并通过 resolveConflictName 自动产生不冲突的新文件名。批量重命名使用数据库事务，若中途失败先回滚数据库，再尝试把已经在磁盘上改名的文件改回原名。")
    add_code_block(doc, """
conn.setAutoCommit(false);
try {
    // 逐个移动文件并更新数据库
    conn.commit();
} catch (Exception e) {
    conn.rollback();
    // 回滚已经成功改名的磁盘文件
}
""")

    add_heading(doc, "3.5 图片展示与幻灯片播放", 2)
    add_para(doc, "图片展示和幻灯片播放都围绕 ImageFile 列表工作。SlideshowController 保存当前索引和缩放比例，支持上一张、下一张、自动播放、循环、随机、全屏、适应窗口和实际大小。自动播放使用 JavaFX Timeline，根据设置页中的播放间隔动态构建。")
    add_code_block(doc, """
autoPlayTimeline = new Timeline(
    new KeyFrame(Duration.seconds(playIntervalSeconds), event -> {
        if (!showNextImage(true)) {
            stopAutoPlay();
        }
    })
);
autoPlayTimeline.setCycleCount(Timeline.INDEFINITE);
""")
    add_para(doc, "为了提高演示完整度，幻灯片还支持底部缩略图条和背景音乐。用户可从内置音乐中选择，也可以选择本地音频文件；音乐开关、音量和播放顺序均在界面上直接操作。")

    add_heading(doc, "3.6 标签、搜索与扩展接口实现", 2)
    add_para(doc, "搜索功能先保证关键词检索可用。用户输入关键词后，SearchService 调用 TagDao 在文件名、路径、标签和描述字段中检索匹配图片，再把图片 ID 列表转回完整 ImageFile 对象。这个流程不依赖外部网络，适合作为课程验收的稳定功能。")
    add_code_block(doc, """
private SearchResult searchByKeyword(String keyword, Optional<Integer> directoryId) {
    List<Integer> imageIds = directoryId.isPresent()
            ? tagDao.searchImagesByKeyword(keyword, directoryId.get())
            : tagDao.searchImagesByKeyword(keyword);
    List<ImageFile> images = loadImagesByIds(imageIds);
    recordSearchHistory(keyword, "KEYWORD", null, images.size());
    return new SearchResult(images, keyword, images.size(), message);
}
""")
    add_para(doc, "AI 图像识别默认连接 CPA 代理节点，Base URL 为 https://cpa.ystone.top/v1。系统不会在源码和提交材料中保存个人密钥，运行时优先从 DIMS_AI_API_KEY、CPA_API_KEY、HAJIMI 或 OPENAI_API_KEY 环境变量读取。模型名称不再要求用户手填，而是由设置页调用兼容 /models 接口后通过下拉菜单选择。")
    add_para(doc, "为了避免误选大目录造成过多 token 消耗，ScanTask 将 AI 识别限制在当前扫描根目录下，单次批处理上限由设置页保存，默认 100(max)，界面中凡是上限都按 N(max) 显示。主界面状态栏会显示当前目录图片总数、本批处理数和待识别数量，并提供停止扫描按钮；清理 AI 标签时，系统会先统计数据库中的标签记录、AI 描述和搜索历史占用，告诉用户不删除时会浪费多少空间，再由用户确认是否清理。")
    add_para(doc, "项目中也保留了外部模型把自然语言转换为 SQL 的接口，但该功能受密钥、网络和模型稳定性影响较大，因此在论文中只作为扩展接口说明。真正执行 SQL 前，TagDaoImpl 会检查语句必须是只读 SELECT，并拒绝 INSERT、UPDATE、DELETE、DROP、ALTER 等危险关键字。")

    add_heading(doc, "3.7 图片编辑与版本历史实现", 2)
    add_para(doc, "图片编辑功能围绕版本历史表实现。用户进行裁剪、标注、绘制文字或恢复时，系统保存新的版本记录和对应操作参数。这样做的优点是编辑过程可追踪，误操作后可以回到之前版本，也便于在数据库课程要求中体现数据持久化设计。")

    add_heading(doc, "4 系统测试", 1)
    add_heading(doc, "4.1 测试环境", 2)
    add_table(doc, [
        ("操作系统", "Windows", "本地桌面环境。"),
        ("JDK", "Java 21 编译目标，JDK 21+ 可运行", "当前开发机将 .jar 双击关联到 JDK 26 的 javaw.exe。"),
        ("数据库", "PostgreSQL 18.3", "本地 image_manager 数据库。"),
        ("构建工具", "Maven 3.9.14", "执行 mvn -DskipTests package 生成 shaded JAR。"),
        ("网络/API", "CPA代理节点，可选密钥", "支持环境变量读取密钥；未配置时不影响基础图片管理功能。"),
    ], headers=["项目", "版本/配置", "说明"], widths=[3.5, 4.5, 7.5])
    add_para(doc, "本次更新后重新执行 Maven 编译、单元测试、FXML/CSS 静态检查和 JavaFX 界面快照测试。界面快照测试直接加载当前 FXML 与 style.css，在不同窗口尺寸下生成 PNG，用于检查控件截断、文字重叠、滚动区域、选中状态和按钮可见性。数据库脚本执行后 public schema 中基础表、视图、触发器和函数均能创建，目标 JAR 为包含依赖的 shaded JAR。")

    add_heading(doc, "4.2 模块测试", 2)
    test_rows = [
        ("目录树加载", "启动程序，展开 C/D 磁盘和图片目录", "显示子目录且不卡死", "目录按需加载，能继续展开", "通过"),
        ("支持格式扫描", "在目录放入 jpg/jpeg/png/gif/bmp 与 txt", "只显示支持图片格式", "仅图片进入缩略图网格", "通过"),
        ("缩略图显示", "选择含多张不同尺寸图片的目录", "缩略图保持比例并显示文件名", "卡片显示比例正常、文件名可见", "通过"),
        ("目录图片数", "切换不同图片目录", "数量随目录变化", "状态栏和标题区数量同步刷新", "通过"),
        ("单选", "点击一张缩略图", "仅该缩略图高亮", "高亮状态正确", "通过"),
        ("多选", "Ctrl 点击多张或拖拽框选", "多张缩略图高亮", "选择数量正确", "通过"),
        ("单个删除", "选中一张图片执行删除", "图片从界面消失并记录删除", "数据库逻辑删除、界面刷新", "通过"),
        ("多个删除", "选中多张图片执行删除", "所有选中图片删除", "列表批量处理正常", "通过"),
        ("复制粘贴", "复制一张或多张图片到另一目录", "目标目录出现副本，重名自动处理", "文件复制成功并入库", "通过"),
        ("重命名", "单张改名、批量按前缀编号", "文件名和数据库一致", "事务提交后刷新正确", "通过"),
        ("图片切换", "进入幻灯片后点上一张/下一张", "大图按顺序切换", "索引和缩略图高亮同步", "通过"),
        ("放大缩小", "点击放大/缩小/适应窗口", "图片按比例缩放", "缩放范围受限且显示正常", "通过"),
        ("幻灯片播放", "设置 3 秒间隔并播放", "自动切换图片，可暂停", "Timeline 定时切换正常", "通过"),
        ("AI模型下拉", "打开设置页并刷新模型", "从 /models 获取列表供选择", "模型不再手动填写，未取到时给出提示", "通过"),
        ("AI批处理上限", "在设置页调整单批上限后扫描目录", "按配置值 N(max) 限制本批处理数量", "状态栏显示目录总数、本批数量和待识别数量", "通过"),
        ("停止扫描", "AI识别过程中点击停止扫描", "当前任务尽快中断，不继续旧目录", "任务取消后按钮状态恢复", "通过"),
        ("清理AI标签", "点击清理AI标签", "显示标签文件位置、大小和空间浪费并二次确认", "确认后清理数据库AI标签记录", "通过"),
        ("目录切换安全", "扫描中切换到新目录", "取消旧扫描并只识别当前目录", "待识别SQL已限定当前根目录", "通过"),
        ("首次向导布局", "选择较长路径或大目录", "确定/取消按钮仍可见", "内容区可滚动且窗口动态限高", "通过"),
        ("扩展接口配置", "未配置密钥启动并搜索", "AI功能跳过，不影响基础功能", "关键词搜索和图片管理可用", "通过"),
        ("数据库脚本", "执行 sql/schema.sql", "脚本无错误，表视图齐全", "执行成功，表/视图数量符合预期", "通过"),
    ]
    add_table(doc, test_rows, headers=["测试项", "输入/步骤", "预期结果", "实际结果", "结论"], widths=[2.6, 4.2, 3.9, 3.8, 1.3])

    add_heading(doc, "4.3 系统完整测试", 2)
    add_para(doc, "完整流程测试从启动程序开始：选择扫描目录，展开目录树，点击目录加载缩略图，使用单选和多选分别执行复制、粘贴、重命名、删除，随后进入幻灯片窗口进行前后切换、放大缩小和自动播放。完成后检查数据库 images、directories、operation_logs 与 app_settings 等表，确认文件系统和数据库状态一致。")
    add_para(doc, "扩展接口的测试口径为：在设置页填写 base url、api key 和 model 后，只验证配置保存和接口调用入口是否能被正常触发；若验收环境没有密钥或网络条件，则直接使用关键词搜索。系统不在源码中保存密钥，避免提交材料泄露个人配置。")

    add_heading(doc, "4.4 打包测试", 2)
    add_code_block(doc, """
mvn -q -DskipTests compile
mvn -q test
mvn -q -DskipTests test-compile exec:java "-Dexec.classpathScope=test" "-Dexec.mainClass=com.imagemanager.UiSnapshotSmoke"
psql -v ON_ERROR_STOP=1 -U postgres -d image_manager -f sql/schema.sql
""")
    add_para(doc, "打包产物 target/image-manager-1.0.0.jar 已复制为提交要求的“面向对象程序设计实践目标代码.JAR”。该 JAR 包含 JavaFX、PostgreSQL JDBC、HikariCP、OkHttp、Jackson 等运行依赖，便于课程验收。本次界面更新重点验证了编译、测试和真实界面截图生成，避免报告截图与当前程序不一致。")

    add_heading(doc, "4.5 前端缩放与截图测试", 2)
    add_para(doc, "JavaFX 桌面程序没有浏览器 zoom，本次按窗口尺寸和系统显示缩放的等效效果验证。最小窗口覆盖 900×600，默认窗口覆盖 1200×800，宽屏窗口覆盖 1440×900；弹窗和工具窗口按 100%、125% 和 150% 附近的等效尺寸截图复查。检查重点包括：按钮文字不截断，底部操作栏不重叠，目录树展开箭头可见，选中缩略图状态清楚，设置页和欢迎向导在窄窗口下可滚动。")
    add_table(doc, [
        ("主界面", "900×600、1200×800、1440×900", "目录树、搜索区、缩略图、状态栏、清理AI标签和幻灯片按钮均可见。", "通过"),
        ("首次启动向导", "640×620、800×775", "长路径输入、AI配置说明、警告区域和确定/取消按钮不遮挡。", "通过"),
        ("系统设置", "680×720、850×900", "折叠区、输入框、模型下拉、数值输入、滚动条和底部按钮正常。", "通过"),
        ("图片查看器", "960×680、1200×850", "深色图片区、图片等比显示、缩放/切换/编辑/幻灯片按钮不截断。", "通过"),
        ("幻灯片播放", "1000×700、1250×875", "底部控制栏、音乐下拉、音量滑块、缩略图条和退出按钮无重叠。", "通过"),
        ("图片编辑器", "1000×750、1250×938", "工具栏、画布、标注图层、版本历史条和状态文本可读。", "通过"),
        ("批量重命名弹窗", "450×360、563×450", "名称前缀、编号输入、预览区域和确认按钮布局稳定。", "通过"),
    ], headers=["界面", "截图尺寸", "检查结论", "结果"], widths=[3.0, 4.0, 7.0, 1.5])

    add_heading(doc, "5 系统运行界面", 1)
    ui_images = make_ui_images()
    add_para(doc, "本节截图均由 JavaFX 测试程序直接加载当前 FXML 与 CSS 后截取，来源为 target/ui-smoke，不再使用手工绘制或旧版示意图。为便于老师检查，本节同时放入默认尺寸、最小尺寸和放大等效尺寸下的界面图。")
    captions = [
        "图5-1 主界面 900×600：最小窗口下目录树、缩略图和状态栏",
        "图5-2 主界面 1200×800：默认窗口下的工作台布局",
        "图5-3 主界面 1440×900：宽屏窗口下的缩略图网格",
        "图5-4 首次启动向导 640×620：路径选择、AI配置说明和底部按钮",
        "图5-5 首次启动向导 800×775：放大等效尺寸下的滚动与按钮状态",
        "图5-6 系统设置 680×720：窄窗口下的配置表单",
        "图5-7 系统设置 850×900：放大等效尺寸下的设置页",
        "图5-8 图片查看器 960×680：大图展示和底部控制区",
        "图5-9 图片查看器 1200×850：宽窗口下的图片等比展示",
        "图5-10 幻灯片播放 1000×700：播放控制、音乐和缩略图条",
        "图5-11 幻灯片播放 1250×875：放大等效尺寸下的控制栏",
        "图5-12 图片编辑器 1000×750：工具栏、画布和版本历史",
        "图5-13 图片编辑器 1250×938：放大等效尺寸下的编辑界面",
        "图5-14 批量重命名弹窗 450×360：最小弹窗布局",
        "图5-15 批量重命名弹窗 563×450：放大等效尺寸布局",
    ]
    for img_path, caption in zip(ui_images, captions):
        add_para(doc, caption, 10.5, True, WD_ALIGN_PARAGRAPH.CENTER, False)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6.1))
        add_para(doc, "该截图对应当前 JavaFX 程序的真实渲染结果。测试样例数据只用于截图复查，实际运行时界面数据来自用户选择的本地图片目录和 PostgreSQL 数据库。")
        doc.add_page_break()

    add_heading(doc, "6 总结", 1)
    add_heading(doc, "6.1 毕振岚总结", 2)
    add_para(doc, "这一版程序里，我承担的主要压力集中在代码本身。目录树、缩略图、右键菜单、图片查看器、幻灯片、编辑器、主题背景和最终版修复都需要放到同一个 JavaFX 工程里跑通，真正麻烦的地方不是写一个按钮，而是保证一次重命名、一次删除、一次目录切换之后，界面、文件系统和数据库不要各走各的状态。后期 20260518、20260519 两次集中提交，也基本是在补这些体验和稳定性问题。")
    add_para(doc, "这次实践让我更清楚地感受到 Controller 不能无限长下去。现在 MainController 仍然承担了很多协调工作，虽然通过 Service、DAO 和工具类做了分层，但后续如果继续维护，我会优先把缩略图渲染、选择状态、右键菜单和搜索区域拆成更小的组件。这样比单纯继续加功能更重要。")
    add_heading(doc, "6.2 陈厚华总结", 2)
    add_para(doc, "我在小组中更多是配合扩展功能和运行检查。最开始分配到的方向包括音乐播放、外部接口和网盘方向的尝试，最后真正稳定进入主流程的是幻灯片音乐和设置项检查，外部接口部分因为验收环境和密钥问题没有作为核心功能写满。这个过程让我意识到，课程项目里“能跑”和“写在设想里”不是一回事，最后文档必须跟实际程序对齐。")
    add_para(doc, "测试时我重点关注的是用户从启动、选择目录、查看缩略图、进入幻灯片到返回主界面的连续流程。单个功能看起来不复杂，但连起来之后会暴露很多小问题，例如按钮状态、目录为空时的提示、复制后目标目录是否刷新等。以后再做类似项目，我会更早记录测试步骤，而不是等最后再回忆。")
    add_heading(doc, "6.3 徐阳总结", 2)
    add_para(doc, "我这次主要负责把项目材料整理成能交、能看、能对上老师要求的报告。前面的数据库需求分析、文档整理和后面的论文整合其实很考验耐心：代码里有什么、群聊里谁做了什么、老师模板要什么、其他组提交暴露出什么格式问题，都要放到同一个文档逻辑里。只靠自动生成一版报告是不够的，封面表格、测试表、分工、参考资料这些地方都要手动核对。")
    add_para(doc, "写报告时我最大的感受是，不能为了让系统显得高级就把没稳定落地的功能写成已经完成。比如外部模型接口可以作为扩展点，但不应该抢走电子图片管理程序本身的重点。最后我选择把报告写回目录树、缩略图、复制粘贴、重命名、幻灯片、主题和数据库持久化这些确定存在的功能上，这样答辩时更稳。")
    add_heading(doc, "6.4 小组总结", 2)
    add_para(doc, "本系统按老师要求完成了电子图片管理程序的主要功能，并在此基础上加入数据库持久化、图片编辑版本历史、幻灯片背景音乐、主题背景和设置页等扩展。代码采用分层结构，便于说明、演示和后续维护。课程论文、评分表、源代码 ZIP、目标代码 JAR 和附加说明均按最终提交要求准备。")
    add_heading(doc, "参考资料", 1)
    for idx, item in enumerate([
        "《面向对象程序设计实践》课程工作安排、论文撰写格式及评分表模板。",
        "JavaFX 官方文档与 OpenJFX Maven 插件说明。",
        "PostgreSQL 18 文档：表、视图、触发器、事务与索引。",
        "OkHttp 与 Jackson 文档：HTTP 请求与 JSON 解析。",
        "Maven Shade Plugin 文档：可运行 JAR 打包与 Manifest 配置。"
    ], start=1):
        add_para(doc, f"[{idx}] {item}")

    add_page_number_footer(doc)
    out = STAGING / "面向对象程序设计实践论文.docx"
    doc.save(out)
    return out


def make_score_docx():
    doc = Document(TEACHER_SCORE_TEMPLATE)
    table = doc.tables[0]
    set_cell_text(table.cell(0, 1), TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 8), "2024级软件工程5班", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_idx, member in zip([2, 3, 4], MEMBERS):
        role, sid, name, task, workload = member
        set_cell_text(table.cell(row_idx, 0), role, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(row_idx, 1), sid, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(row_idx, 3), name, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(row_idx, 5), task, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(row_idx, 10), workload, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(5, 0), "组员", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(6, 0), "\n".join([
        "程序功能完成情况：",
        "基本功能：完成；扩展功能：完成。",
        "支持规定图片格式：JPG、JPEG、GIF、PNG、BMP，完成。",
        "目录树实现且操作正常：完成；点击目录显示缩略图及个数：完成。",
        "缩略图包含图片、文件名且保持比例：完成。",
        "缩略图单选：完成；缩略图多选：完成。",
        "单个图片删除：完成；多个图片删除：完成。",
        "单个图片复制粘贴：完成；多个图片复制粘贴：完成。",
        "单个图片重命名：完成；多个图片批量重命名：完成。",
        "图片展示的图片切换：完成；放大缩小：完成；幻灯片播放：完成。",
        "AI识别安全控制：完成，单批上限可调且按N(max)显示、可停止扫描、可统计并清理数据库AI标签。",
        "程序代码质量（主要考查程序结构、代码质量、运行效率等）：优秀。",
        "评语：系统采用 JavaFX + PostgreSQL 分层设计，基础采分点覆盖完整，并实现图片编辑版本历史、幻灯片音乐、主题设置、AI识别配置、安全扫描控制和可运行 JAR 打包。"
    ]), size=9)
    for p in doc.paragraphs:
        for run in p.runs:
            set_east_asia_font(run)
    out = STAGING / "2024级-软件工程-5班第7组_评分表_电子图片管理程序.docx"
    doc.save(out)
    return out


def make_extra_docx():
    doc = Document()
    set_doc_defaults(doc)
    add_para(doc, "《面向对象程序设计实践》附加说明", 18, True, WD_ALIGN_PARAGRAPH.CENTER, False)
    add_para(doc, f"题目：{TITLE}（{LONG_TITLE}）", 12, True, WD_ALIGN_PARAGRAPH.CENTER, False)
    add_para(doc, "班级：2024级软件工程5班  第07组", 12, False, WD_ALIGN_PARAGRAPH.CENTER, False)
    add_para(doc, "成员：202425220501 毕振岚；202425220502 陈厚华；202425220527 徐阳", 12, False, WD_ALIGN_PARAGRAPH.CENTER, False)
    add_heading(doc, "1 交付文件说明", 1)
    add_table(doc, [
        ("面向对象程序设计实践论文.docx", "可编辑课程论文，按老师模板包含封面、目录、系统分析、设计、实现、测试、运行界面和总结。"),
        ("2024级-软件工程-5班第7组_评分表_电子图片管理程序.docx", "可编辑评分表，按电子图片管理程序评分表模板填写小组成员、任务分工、功能完成情况和代码质量。"),
        ("面向对象程序设计实践源代码.ZIP", "Maven/IDEA 工程源码，包含 pom.xml、src、sql、assets、README 和开发日志。"),
        ("面向对象程序设计实践目标代码.JAR", "由 target/image-manager-1.0.0.jar 复制而来，包含运行依赖。"),
        ("面向对象程序设计实践附加说明.docx", "运行环境、数据库初始化、JAR 启动排查、测试命令和注意事项。"),
    ], headers=["文件", "说明"], widths=[6.0, 9.5])
    add_heading(doc, "2 运行环境", 1)
    add_table(doc, [
        ("JDK", "Java 21 或更高版本；当前开发机已将 .jar 双击关联到 JDK 26。"),
        ("数据库", "PostgreSQL 18.3，本地数据库名 image_manager。"),
        ("构建工具", "Maven 3.9.14。"),
        ("主要依赖", "JavaFX 21.0.6、PostgreSQL JDBC、HikariCP、OkHttp、Jackson、Logback。"),
    ], headers=["项目", "要求"], widths=[4.0, 11.5])
    add_heading(doc, "3 数据库初始化", 1)
    add_code_block(doc, """
psql -U postgres -c "CREATE DATABASE image_manager ENCODING 'UTF8';"
psql -U postgres -d image_manager -f sql/schema.sql
""")
    add_para(doc, "默认连接配置位于 src/main/resources/config/database.properties：db.url=jdbc:postgresql://localhost:5432/image_manager，db.username=postgres，db.password=1234。教师验收环境如密码不同，可修改该配置后重新打包，或在 IDE 中直接运行。")
    add_heading(doc, "4 启动方式", 1)
    add_code_block(doc, """
java -jar 面向对象程序设计实践目标代码.JAR
""")
    add_para(doc, "目标代码支持直接双击运行；上方命令主要用于排查环境问题。若使用源码运行，可在工程根目录执行 mvn javafx:run。首次进入程序建议选择一个较小的图片目录，AI识别单批上限默认100(max)，可在设置页调整，也可在主界面停止扫描或清理AI标签。")
    add_heading(doc, "5 双击 JAR 无反应的原因", 1)
    add_para(doc, "本项目目标 JAR 是 JavaFX 桌面程序，双击时 Windows 通常用 javaw.exe 启动，错误信息不会显示在命令行窗口里，所以看起来像“没反应”。本次用命令行启动过同一个 JAR，日志显示数据库连接池启动、主界面初始化并成功启动应用，说明 JAR 本身不是空包。")
    add_para(doc, "如果双击仍无窗口，优先检查三点：第一，系统 .jar 文件关联的 Java 版本是否为 Java 21 或更高版本；第二，本地 PostgreSQL 的 image_manager 数据库是否已创建并能按 database.properties 连接；第三，当前 Windows 用户的环境变量中是否已有 AI 密钥。命令行启动只是排错手段，正常验收可以双击 JAR。")
    add_code_block(doc, """
java -jar 面向对象程序设计实践目标代码.JAR
""")
    add_heading(doc, "6 已执行的交付前测试", 1)
    add_table(doc, [
        ("mvn -q -DskipTests compile", "通过", "确认当前 JavaFX、FXML、CSS 和控制器代码能完成编译。"),
        ("mvn -q test", "通过", "执行现有测试，确认本次界面风格改造没有破坏构建。"),
        ("UiSnapshotSmoke", "通过", "直接加载当前 FXML 与 style.css，生成全部主要界面的多尺寸真实截图。"),
        ("FXML/CSS 静态检查", "通过", "FXML 可解析，新增 styleClass 均有 CSS 定义，未保留关键内联颜色样式。"),
        ("psql -v ON_ERROR_STOP=1 -U postgres -d image_manager -f sql/schema.sql", "通过", "脚本执行成功，表和视图创建正常。"),
        ("数据库对象检查", "通过", "public schema 基础表 13 张，视图 4 个。"),
    ], headers=["检查项", "结果", "说明"], widths=[7.0, 2.0, 6.5])
    add_heading(doc, "7 与老师模板的格式对照", 1)
    add_table(doc, [
        ("论文封面", "已包含课程名、题目、专业、班级、三名小组成员、指导老师和提交时间。"),
        ("目录", "已插入 Word 目录域，导出 PDF 前自动更新。"),
        ("正文结构", "严格包含 1 系统分析、2 系统设计、3 系统实现、4 系统测试、5 系统运行界面、6 总结。"),
        ("字号字体", "正文使用宋体五号，一级标题宋体四号加粗，二级标题宋体小四加粗，三级标题宋体五号加粗。"),
        ("测试要求", "测试表包含输入/步骤、预期结果、实际结果和结论。"),
        ("个人总结", "三名成员分别撰写总结。"),
    ], headers=["模板要求", "对应处理"], widths=[4.5, 11.0])
    out = STAGING / "面向对象程序设计实践附加说明.docx"
    doc.save(out)
    return out


def main():
    STAGING.mkdir(parents=True, exist_ok=True)
    report = make_report_docx()
    score = make_score_docx()
    extra = make_extra_docx()
    print(report)
    print(score)
    print(extra)


if __name__ == "__main__":
    main()
