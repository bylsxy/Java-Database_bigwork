from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[4]
FINAL_DIR = ROOT / "docs" / "数据库系统基础" / "最终交付"
SOURCE = FINAL_DIR / "08_Gemini网页限额交付包" / "01_最终报告_Gemini限额版" / "Gemini返回回答.md"
TARGET = FINAL_DIR / "02_课程报告" / "第07组毕振岚-数据库课程设计报告.docx"
FIGURE_DIRS = [
    FINAL_DIR / "06_证据与清单" / "figures",
    FINAL_DIR / "07_Gemini交付包" / "01_Gemini_最终报告包" / "上传资料" / "05_界面截图与图表",
    FINAL_DIR / "09_豆包PPT最终交付包" / "可选素材_不用默认上传",
]


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("**", "")
    text = re.sub(r"(?<!\*)\*(?!\*)", "", text)
    text = text.replace("提示词上下文", "结构化上下文")
    text = text.replace("请求大模型生成", "请求智能服务生成")
    return text.strip()


def find_figure(name: str) -> Path | None:
    for directory in FIGURE_DIRS:
        candidate = directory / name
        if candidate.exists():
            return candidate
    return None


def set_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: float = 11, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_font(run, size=size, bold=bold)


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12.5)]:
        style = doc.styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cleaned = []
    for row in rows:
        cells = [clean_inline(cell) for cell in row]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        cleaned.append(cells)
    if not cleaned:
        return
    col_count = max(len(row) for row in cleaned)
    table = doc.add_table(rows=len(cleaned), cols=col_count)
    table.style = "Table Grid"
    for i, row in enumerate(cleaned):
        for j in range(col_count):
            cell = table.cell(i, j)
            cell.text = row[j] if j < len(row) else ""
            for paragraph in cell.paragraphs:
                set_paragraph_font(paragraph, size=10.5, bold=(i == 0))
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        rows.append([part.strip() for part in raw.split("|")])
        i += 1
    return rows, i


def add_markdown_doc() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document(str(TARGET)) if TARGET.exists() else Document()
    clear_document_body(doc)
    setup_styles(doc)

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line or line == "---":
            i += 1
            continue

        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        image_match = re.fullmatch(r"\[(.+\.(?:png|jpg|jpeg))\]", line, flags=re.I)
        if image_match:
            figure_name = image_match.group(1)
            figure = find_figure(figure_name)
            if figure:
                doc.add_picture(str(figure), width=Inches(5.9))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption = doc.add_paragraph(f"图：{Path(figure_name).stem}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_font(caption, size=10)
            else:
                paragraph = doc.add_paragraph(f"图：{Path(figure_name).stem}（见随附素材）")
                set_paragraph_font(paragraph, size=10)
            i += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            paragraph = doc.add_heading(clean_inline(heading.group(2)), level=level)
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        bullet = re.match(r"^[*+-]\s+(.+)$", line)
        if numbered:
            paragraph = doc.add_paragraph(clean_inline(numbered.group(1)), style="List Number")
        elif bullet:
            paragraph = doc.add_paragraph(clean_inline(bullet.group(1)), style="List Bullet")
        else:
            paragraph = doc.add_paragraph(clean_inline(line))
        set_paragraph_font(paragraph)
        i += 1

    doc.save(str(TARGET))


if __name__ == "__main__":
    add_markdown_doc()
    doc = Document(str(TARGET))
    chars = sum(len(paragraph.text.strip()) for paragraph in doc.paragraphs)
    print({"target": str(TARGET), "paragraphs": len(doc.paragraphs), "tables": len(doc.tables), "chars": chars})
