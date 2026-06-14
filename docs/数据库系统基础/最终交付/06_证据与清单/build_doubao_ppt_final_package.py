from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path.cwd()
FINAL_DIR = ROOT / "docs" / "数据库系统基础" / "最终交付"
GEMINI_PPT_DIR = FINAL_DIR / "08_Gemini网页限额交付包" / "02_豆包PPT策划_Gemini限额版"
GEMINI_REPORT_DIR = FINAL_DIR / "08_Gemini网页限额交付包" / "01_最终报告_Gemini限额版"
OLD_PPT_UPLOAD_DIR = FINAL_DIR / "07_Gemini交付包" / "02_Gemini_生成豆包PPT包" / "上传资料"
OUT_DIR = FINAL_DIR / "09_豆包PPT最终交付包"


IMAGE_SRC = OLD_PPT_UPLOAD_DIR / "04_界面截图与图表素材"
OPTIONAL_PPT = OLD_PPT_UPLOAD_DIR / "03_当前PPT草稿" / "当前PPT草稿_仅供覆盖范围参考.pptx"
GEMINI_OUTPUT = GEMINI_PPT_DIR / "Gemini返回的待交付豆包的输出.md"
GEMINI_REPORT = GEMINI_REPORT_DIR / "Gemini返回回答.md"


MAIN_IMAGES = [
    "图2_ER关系模式图.png",
    "图4_SQL对象与性能设计.png",
    "图5_NL2SQL安全链路.png",
    "图6_后台架构图.png",
    "real_02_主界面_1200x800_默认窗口.png",
    "real_08_数据库连接与初始化向导_760x680.png",
    "real_14_图片编辑器_1250x938.png",
]

ATLAS_IMAGES = [
    "图1_系统功能结构图.png",
    "图2_ER关系模式图.png",
    "图3_系统总体流程图.png",
    "图4_SQL对象与性能设计.png",
    "图5_NL2SQL安全链路.png",
    "图6_后台架构图.png",
    "图7_演示路线.png",
    "01_主界面_目录树与缩略图.png",
    "02_幻灯片播放.png",
    "03_设置页_运行配置.png",
    "04_图片编辑与版本历史.png",
    "05_标签与扩展搜索.png",
    "real_05_首次启动向导_800x775.png",
    "real_08_数据库连接与初始化向导_760x680.png",
    "real_14_图片编辑器_1250x938.png",
]


DOUBAO_PROMPT = """请严格参考老师要求和我同步上传的资料，为《基于 PostgreSQL 的数字图像集成管理系统》生成数据库系统课程设计答辩 PPT。

我已上传 3 个文件：
1. `01_豆包PPT生成主资料.docx`：PPT 的核心生成说明、评分导向、真实性边界、页面方案和演示路线。
2. `02_课程报告正文参考.docx`：课程报告正文，用于补全数据库设计、系统实现和测试说明，不要照搬成长段落。
3. `03_视觉素材图集.docx`：真实界面截图和数据库图表，请优先把这些图用于页面视觉。

请生成 16:9 横版 PPT，适合 15 分钟数据库系统课程设计答辩，建议 12 到 14 页。PPT 必须图像优先、少字、专业，不要做成普通软件介绍，也不要做廉价科技风、黑客风、花哨渐变或满屏卡片模板。视觉基调应明亮、克制、学术，建议用深墨色标题、青绿色数据库强调色和少量橙色交互强调色。

评分主线必须非常明确：数据库设计 30 分、功能设计 20 分、后台程序设计 15 分、界面设计 15 分、报告表述水平 10 分、PPT 表达能力 10 分。请把“数据库含金量”作为主线，突出 PostgreSQL 表结构、ER 关系、索引、视图、触发器、存储过程、递归 CTE、事务、连接池和安全查询，不要只讲图片查看功能。

必须体现的真实事实：
系统是 JavaFX + PostgreSQL 桌面端数字图像管理系统。数据库包含 13 张表、19 个索引、4 个视图、4 个函数、5 个触发器、5 个存储过程。核心关系包括 directories 自引用目录树、images 主表、tag_categories/tags/image_tags 多对多标签体系、ai_analysis_results 一对一 AI 结果、image_versions 一对多版本历史、operation_logs 审计日志。真实功能包括完整磁盘目录树懒加载、缩略图和元数据入库、bytea 缩略图、AI 标签、关键词搜索、NL2SQL 搜索、版本历史、幻灯片音乐、数据库初始化向导、离线降级、HikariCP、PreparedStatement 和事务回滚。

真实性边界：
WebDAV/云端同步、全网全盘搜索、语音搜索、SQLite 多数据库底层切换只能作为后续扩展，不要写成已完成主流程。AI 功能需要用户自行配置 OpenAI-compatible endpoint 和密钥。NL2SQL 只允许走只读视图 `v_image_search`，拦截非 SELECT 语句，并带超时和行数限制。

页面建议：
1. 封面：项目题目、第 07 组、成员与课程答辩。
2. 系统定位：这不是普通图片查看器，而是以 PostgreSQL 为核心的数据管理应用。
3. 数据库 ER 模型：放大展示 ER 图和核心关系。
4. 数据库对象全景：13 表、19 索引、4 视图、4 函数、5 触发器、5 存储过程、递归 CTE。
5. 后台基座：HikariCP、PreparedStatement、事务、初始化向导、离线降级。
6. AI 标签与 NL2SQL 安全链路：结果落库、只读视图、安全拦截。
7. 核心界面：目录树懒加载、缩略图、元数据。
8. 版本历史与操作审计：image_versions、operation_logs、恢复版本。
9. 幻灯片与图片编辑：体现 UI 和功能完整性。
10. 工程交付：JAR、Windows portable exe、源码包、验证日志。
11. 现场演示路线：先启动/初始化，再扫描目录，再搜索，再版本恢复，最后幻灯片。
12. 总结与扩展：已完成能力、数据库课程价值、后续扩展边界。

请直接生成 PPT，可包含简洁演讲者备注。不要把报告长段落照搬到页面上，每页只保留最关键的结论和图。"""


def ensure_inputs() -> None:
    required = [GEMINI_OUTPUT, GEMINI_REPORT, OPTIONAL_PPT]
    required.extend(IMAGE_SRC / name for name in ATLAS_IMAGES)
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("缺少输入文件：\n" + "\n".join(missing))


def reset_output_dir() -> None:
    out = OUT_DIR.resolve()
    final = FINAL_DIR.resolve()
    if OUT_DIR.exists():
        if final not in out.parents:
            raise RuntimeError(f"拒绝删除非最终交付目录下的路径：{OUT_DIR}")
        shutil.rmtree(OUT_DIR)
    (OUT_DIR / "上传文件_推荐").mkdir(parents=True, exist_ok=True)
    (OUT_DIR / "可选素材_不用默认上传").mkdir(parents=True, exist_ok=True)


def set_default_font(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Microsoft YaHei"


def add_markdown_text(doc: Document, text: str, max_lines: int | None = None) -> None:
    lines = text.splitlines()
    if max_lines is not None:
        lines = lines[:max_lines]
    list_buffer: list[str] = []

    def flush_list() -> None:
        nonlocal list_buffer
        for item in list_buffer:
            doc.add_paragraph(item, style="List Bullet")
        list_buffer = []

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush_list()
            continue
        if line.strip() == "---":
            flush_list()
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            flush_list()
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            flush_list()
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            flush_list()
            doc.add_heading(line[4:].strip(), level=3)
        elif re.match(r"^\s*[\*\-]\s+", line):
            list_buffer.append(re.sub(r"^\s*[\*\-]\s+", "", line).strip())
        elif re.match(r"^\s*\d+\.\s+", line):
            flush_list()
            doc.add_paragraph(line.strip(), style="List Number")
        else:
            flush_list()
            cleaned = line.replace("**", "").replace("`", "")
            doc.add_paragraph(cleaned)
    flush_list()


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(subtitle).font.size = Pt(11)


def add_image(doc: Document, image_name: str, caption: str) -> None:
    path = IMAGE_SRC / image_name
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.4))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True


def build_main_docx() -> Path:
    doc = Document()
    set_default_font(doc)
    add_title(doc, "豆包 PPT 生成主资料", "数据库系统课程设计答辩：基于 PostgreSQL 的数字图像集成管理系统")
    doc.add_heading("一、直接粘贴给豆包的总提示词", level=1)
    doc.add_paragraph(DOUBAO_PROMPT)
    doc.add_heading("二、Gemini 生成的页面策划原文", level=1)
    add_markdown_text(doc, GEMINI_OUTPUT.read_text(encoding="utf-8"))
    doc.add_heading("三、推荐优先使用的关键视觉素材", level=1)
    for name in MAIN_IMAGES:
        add_image(doc, name, f"建议 PPT 使用素材：{name}")
    out = OUT_DIR / "上传文件_推荐" / "01_豆包PPT生成主资料.docx"
    doc.save(out)
    return out


def build_report_docx() -> Path:
    doc = Document()
    set_default_font(doc)
    add_title(doc, "课程报告正文参考", "供豆包补全数据库设计和系统实现细节，不要照搬成长段落")
    add_markdown_text(doc, GEMINI_REPORT.read_text(encoding="utf-8"))
    out = OUT_DIR / "上传文件_推荐" / "02_课程报告正文参考.docx"
    doc.save(out)
    return out


def build_atlas_docx() -> Path:
    doc = Document()
    set_default_font(doc)
    add_title(doc, "视觉素材图集", "真实界面截图与数据库图表，请用于 PPT 画面设计")
    doc.add_paragraph("以下图像均来自当前项目交付材料。生成 PPT 时优先使用图表和真实界面，不要用空泛图标替代。")
    for name in ATLAS_IMAGES:
        doc.add_heading(name, level=2)
        add_image(doc, name, name)
    out = OUT_DIR / "上传文件_推荐" / "03_视觉素材图集.docx"
    doc.save(out)
    return out


def copy_optional_materials() -> None:
    optional = OUT_DIR / "可选素材_不用默认上传"
    shutil.copy2(GEMINI_OUTPUT, optional / "Gemini返回的待交付豆包的输出.md")
    shutil.copy2(GEMINI_REPORT, optional / "Gemini返回报告正文.md")
    shutil.copy2(OPTIONAL_PPT, optional / "当前PPT草稿_仅供覆盖检查_不要默认上传.pptx")
    for name in ATLAS_IMAGES:
        shutil.copy2(IMAGE_SRC / name, optional / name)


def write_instructions(outputs: list[Path]) -> None:
    upload_list = "\n".join(f"{i}. `{path.name}`，{path.stat().st_size} 字节" for i, path in enumerate(outputs, start=1))
    (OUT_DIR / "00_使用说明.md").write_text(
        f"""# 豆包 PPT 最终生成交付包

本包按豆包网页端更稳妥的方式准备：不上传 ZIP，不一次性上传大量图片，不上传源码文件夹。默认只上传 `上传文件_推荐` 中 3 个 DOCX 文件，再把 `01_复制到豆包输入框.txt` 全文粘贴到豆包输入框。

## 默认上传文件

{upload_list}

## 操作顺序

1. 打开豆包网页端，选择 `PPT 生成`。
2. 点击添加文件，上传 `上传文件_推荐` 里的 3 个 DOCX。
3. 把 `01_复制到豆包输入框.txt` 全文复制到输入框。
4. 篇幅选择“详细”，要求 12 到 14 页、16:9、可下载 PPTX。
5. 生成后重点检查数据库对象数量、WebDAV/云端边界、AI/NL2SQL 安全边界和截图是否正确。

`可选素材_不用默认上传` 里保留了原始 Gemini 输出、报告 Markdown、当前 PPT 草稿和单张图片。如果豆包生成后缺图，再从里面挑少量图片追加，不要一开始全传。
""",
        encoding="utf-8",
    )
    (OUT_DIR / "01_复制到豆包输入框.txt").write_text(DOUBAO_PROMPT, encoding="utf-8")
    (OUT_DIR / "豆包上传规则核查.md").write_text(
        """# 豆包上传规则核查与本包取舍

核查时间：2026-06-15。

豆包官方公开协议可确认：豆包生成的文档、PPT、表格等文件，以及在豆包对话内上传的文件，会自动保存到 AI 云盘；云盘也提供上传文件或文件夹入口。公开协议没有在可访问文本中给出 PPT 生成模式的单次附件数量、单文件大小、ZIP 支持情况等硬限制。

公开教程和第三方资料对上传上限说法不完全一致，常见保守口径是：优先使用 PDF、DOCX、TXT 等文本型文件；免费用户单文件控制在 50MB 以内；同一会话控制在 3 个文件左右；压缩包、图片型文档和大量直接图片更容易解析失败。

因此本包采用保守方案：默认只上传 3 个 DOCX 文件；每个文件均远低于 50MB；不把 ZIP 作为豆包输入；不默认上传大量 PNG；真实图片全部嵌入 `03_视觉素材图集.docx`，减少附件数量。

参考来源：

1. 豆包用户协议：https://www.doubao.com/legal/terms
2. 豆包 AI 云盘使用须知：https://www.doubao.com/legal/ai_space
3. 公开资料对文档上传常见限制的说明：https://www.php.cn/faq/2484403.html
""",
        encoding="utf-8",
    )


def verify_outputs(outputs: list[Path]) -> None:
    for path in outputs:
        if not path.exists():
            raise FileNotFoundError(path)
        if path.stat().st_size > 50 * 1024 * 1024:
            raise RuntimeError(f"{path.name} 超过 50MB，不适合作为默认上传文件")
    recommended = sorted((OUT_DIR / "上传文件_推荐").glob("*"))
    if len([p for p in recommended if p.is_file()]) != 3:
        raise RuntimeError("默认上传文件数量不是 3")


def main() -> None:
    ensure_inputs()
    reset_output_dir()
    outputs = [build_main_docx(), build_report_docx(), build_atlas_docx()]
    copy_optional_materials()
    write_instructions(outputs)
    verify_outputs(outputs)
    for path in outputs:
        print(f"{path.name}\t{path.stat().st_size}")
    print(f"输出目录：{OUT_DIR}")


if __name__ == "__main__":
    main()
