from __future__ import annotations

import shutil
from pathlib import Path

import win32com.client as win32
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
EXP3 = ROOT / "docs" / "软件工程基础" / "实验3"
DOCX = EXP3 / "数字图像管理系统软件设计规格说明书v1.2_审计修订版.docx"
EXPORT_DIR = ROOT / "C_tmp_visio_exports"


COLORS = {
    "blue": "RGB(221,235,247)",
    "green": "RGB(226,239,218)",
    "orange": "RGB(252,228,214)",
    "purple": "RGB(234,226,247)",
    "yellow": "RGB(255,242,204)",
    "gray": "RGB(242,242,242)",
    "red": "RGB(248,203,173)",
    "white": "RGB(255,255,255)",
}


def set_cell(shape, name: str, value: str) -> None:
    try:
        shape.CellsU(name).FormulaU = value
    except Exception:
        pass


def add_box(page, x: float, y: float, w: float, h: float, text: str, color: str = "blue", font_size: int = 8):
    shape = page.DrawRectangle(x, y, x + w, y + h)
    shape.Text = text
    set_cell(shape, "FillForegnd", COLORS[color])
    set_cell(shape, "LineColor", "RGB(80,80,80)")
    set_cell(shape, "Char.Size", f"{font_size} pt")
    set_cell(shape, "Para.HorzAlign", "1")
    set_cell(shape, "VerticalAlign", "1")
    return shape


def add_title(page, text: str):
    shape = add_box(page, 0.3, 7.72, 11.1, 0.38, text, "gray", 12)
    set_cell(shape, "Char.Style", "17")
    return shape


def add_arrow(page, x1: float, y1: float, x2: float, y2: float, text: str = ""):
    line = page.DrawLine(x1, y1, x2, y2)
    set_cell(line, "EndArrow", "13")
    set_cell(line, "LineWeight", "1.25 pt")
    if text:
        midx = (x1 + x2) / 2
        midy = (y1 + y2) / 2
        add_box(page, midx - 0.42, midy - 0.12, 0.84, 0.24, text, "white", 7)
    return line


def add_note(page, x: float, y: float, w: float, h: float, text: str):
    shape = add_box(page, x, y, w, h, text, "yellow", 7)
    set_cell(shape, "LinePattern", "2")
    return shape


def new_doc(app):
    doc = app.Documents.Add("")
    page = app.ActivePage
    page.PageSheet.CellsU("PageWidth").FormulaU = "11.69 in"
    page.PageSheet.CellsU("PageHeight").FormulaU = "8.27 in"
    return doc, page


def save_doc(doc, page, path: Path, png_path: Path):
    doc.SaveAs(str(path))
    page.Export(str(png_path))
    doc.Close()


def diagram_scan(app):
    doc, page = new_doc(app)
    add_title(page, "图11 目录扫描入库详细活动图")
    lanes = [
        ("用户界面", "blue"),
        ("ScanTask", "green"),
        ("DirectoryScanner", "purple"),
        ("ImageService", "orange"),
        ("DAO / PostgreSQL", "gray"),
    ]
    for i, (name, color) in enumerate(lanes):
        x = 0.3 + i * 2.25
        add_box(page, x, 7.15, 2.0, 0.32, name, color, 9)
        add_box(page, x, 0.45, 2.0, 6.55, "", "white", 7)
    steps = [
        (0, "选择扫描目录\n点击开始扫描", "blue"),
        (1, "创建后台任务\n绑定进度条", "green"),
        (2, "递归遍历目录\n过滤图片扩展名", "purple"),
        (2, "读取尺寸/格式\n跳过损坏文件", "purple"),
        (3, "生成 ImageFile\n生成缩略图", "orange"),
        (4, "findOrCreateDirectory\n批量写入 images", "gray"),
        (4, "更新 thumbnail\n记录扫描结果", "gray"),
        (1, "汇总新增/跳过/失败\n发布进度", "green"),
        (0, "刷新目录树与缩略图\n显示完成状态", "blue"),
    ]
    coords = []
    for idx, (lane, text, color) in enumerate(steps):
        x = 0.45 + lane * 2.25
        y = 6.45 - idx * 0.62
        coords.append((x + 1.0, y + 0.2))
        add_box(page, x, y, 1.7, 0.42, text, color, 7)
    for a, b in zip(coords, coords[1:]):
        add_arrow(page, a[0], a[1] - 0.2, b[0], b[1] + 0.2)
    add_note(page, 8.2, 1.0, 2.5, 0.75, "异常处理：权限不足、坏图、重复路径\n均不应中断整个扫描批次")
    return doc, page


def diagram_components(app):
    doc, page = new_doc(app)
    add_title(page, "图13 子系统与构件职责展开图")
    groups = [
        ("界面交互子系统", 0.4, 5.7, "blue", ["MainController", "ImageViewerController", "ImageEditorController", "SettingsController", "WelcomeDialogController"]),
        ("图片业务子系统", 3.1, 5.7, "green", ["ImageServiceImpl", "EditService", "DirectoryScanner", "ScanTask", "ImageUtil"]),
        ("搜索与 AI 子系统", 5.8, 5.7, "purple", ["AIService", "OpenAICompatibleService", "AiTagStorageService", "SearchService", "AIFallbackManager"]),
        ("数据持久化子系统", 8.5, 5.7, "orange", ["ImageDaoImpl", "DirectoryDaoImpl", "TagDaoImpl", "VersionDaoImpl", "SettingsDaoImpl"]),
        ("配置与启动子系统", 3.1, 2.2, "yellow", ["App / Launcher", "DatabaseSetupDialog", "DatabaseBootstrapService", "AIConfig", "database.properties"]),
        ("PostgreSQL 数据库", 8.5, 2.2, "gray", ["directories", "images", "tags / image_tags", "image_versions", "operation_logs"]),
    ]
    centers = {}
    for title, x, y, color, items in groups:
        add_box(page, x, y + 1.15, 2.25, 0.35, title, color, 9)
        for idx, item in enumerate(items):
            add_box(page, x + 0.18, y + 0.72 - idx * 0.32, 1.9, 0.24, item, color, 7)
        centers[title] = (x + 1.125, y + 0.55)
    links = [
        ("界面交互子系统", "图片业务子系统", "调用服务"),
        ("图片业务子系统", "数据持久化子系统", "写入元数据"),
        ("界面交互子系统", "搜索与 AI 子系统", "检索/识别"),
        ("搜索与 AI 子系统", "数据持久化子系统", "标签结果"),
        ("配置与启动子系统", "界面交互子系统", "启动主界面"),
        ("配置与启动子系统", "PostgreSQL 数据库", "初始化"),
        ("数据持久化子系统", "PostgreSQL 数据库", "JDBC"),
    ]
    for a, b, text in links:
        add_arrow(page, *centers[a], *centers[b], text)
    add_note(page, 0.7, 1.0, 2.6, 0.72, "约束：Controller 不直接拼 SQL；\nDAO 不持有界面状态。")
    return doc, page


def diagram_ai(app):
    doc, page = new_doc(app)
    add_title(page, "图14 AI 标签扫描与智能搜索安全链路图")
    nodes = [
        ("用户选择图片/关键词", 0.5, 6.45, "blue"),
        ("Settings 读取 endpoint\n本机私有配置", 2.6, 6.45, "yellow"),
        ("AIService 构造请求\n限制单批数量", 4.8, 6.45, "green"),
        ("OpenAI-compatible API\n按 fallback 顺序调用", 7.0, 6.45, "purple"),
        ("解析描述/标签/置信度", 9.2, 6.45, "green"),
        ("AiTagStorageService\n标准化标签", 1.5, 4.6, "green"),
        ("TagDaoImpl\n去重写入", 3.8, 4.6, "orange"),
        ("tag_categories / tags\nimage_tags / ai_analysis", 6.1, 4.6, "gray"),
        ("SearchService\n关键词/NL2SQL 分流", 8.4, 4.6, "green"),
        ("SQL 安全校验\n只允许 SELECT", 2.6, 2.75, "red"),
        ("v_image_search\n标签+元数据+描述", 5.0, 2.75, "gray"),
        ("返回图片列表\n按相关度排序", 7.5, 2.75, "blue"),
    ]
    centers = []
    for text, x, y, color in nodes:
        add_box(page, x, y, 1.75, 0.55, text, color, 7)
        centers.append((x + 0.875, y + 0.275))
    for i in range(4):
        add_arrow(page, centers[i][0], centers[i][1], centers[i + 1][0], centers[i + 1][1])
    for a, b in [(4, 5), (5, 6), (6, 7), (7, 8), (8, 9), (9, 10), (10, 11)]:
        add_arrow(page, centers[a][0], centers[a][1], centers[b][0], centers[b][1])
    add_note(page, 8.7, 1.0, 2.2, 0.75, "熔断策略：连续失败达到阈值后\n本会话暂停该 endpoint")
    return doc, page


def diagram_bootstrap(app):
    doc, page = new_doc(app)
    add_title(page, "图15 数据库初始化与离线降级流程图")
    steps = [
        ("启动应用", 0.7, 6.55, "blue"),
        ("读取环境变量\n与本机配置", 2.4, 6.55, "yellow"),
        ("初始化连接池", 4.1, 6.55, "green"),
        ("检测 schema\n是否完整", 5.8, 6.55, "orange"),
        ("加载主界面", 7.5, 6.55, "blue"),
        ("连接失败或缺表", 5.8, 5.05, "red"),
        ("弹出数据库向导", 4.1, 3.75, "yellow"),
        ("用户填写参数\n检测连接", 2.4, 3.75, "yellow"),
        ("创建数据库\n执行 schema/data", 4.1, 2.45, "green"),
        ("保存本机配置\n刷新状态栏", 5.8, 2.45, "green"),
        ("离线浏览降级\n标签/搜索受限", 7.5, 3.75, "gray"),
    ]
    centers = []
    for text, x, y, color in steps:
        add_box(page, x, y, 1.45, 0.55, text, color, 7)
        centers.append((x + 0.725, y + 0.275))
    for a, b in [(0, 1), (1, 2), (2, 3), (3, 4), (3, 5), (5, 6), (6, 7), (7, 8), (8, 9), (9, 4), (5, 10)]:
        add_arrow(page, centers[a][0], centers[a][1], centers[b][0], centers[b][1])
    add_note(page, 0.9, 1.05, 2.8, 0.75, "目标：新机器能启动，有数据库修复入口，\n不能因连接失败直接崩溃。")
    return doc, page


def diagram_versions(app):
    doc, page = new_doc(app)
    add_title(page, "图16 图片编辑与版本历史详细设计图")
    cols = [
        ("ImageEditorController", 0.6, "blue"),
        ("EditService", 3.0, "green"),
        ("ImageUtil", 5.4, "purple"),
        ("VersionDaoImpl", 7.8, "orange"),
        ("images / image_versions", 9.8, "gray"),
    ]
    for name, x, color in cols:
        add_box(page, x, 7.1, 1.7, 0.35, name, color, 8)
        add_box(page, x + 0.8, 1.0, 0.1, 5.9, "", "white", 7)
    actions = [
        (0, "选择工具\n裁剪/绘制/文本", 6.45),
        (1, "生成编辑命令\n校验图片状态", 5.8),
        (2, "合成新图片\n生成缩略图", 5.15),
        (3, "insert image_versions\n标记当前版本", 4.5),
        (4, "更新 images\n当前路径/尺寸/缩略图", 3.85),
        (3, "查询版本时间线", 3.2),
        (0, "刷新画布与版本列表", 2.55),
        (1, "恢复旧版本\n回写当前图片", 1.9),
    ]
    centers = {}
    for idx, (col, text, y) in enumerate(actions):
        x = cols[col][1]
        color = cols[col][2]
        add_box(page, x, y, 1.7, 0.42, text, color, 7)
        centers[idx] = (x + 0.85, y + 0.21)
    for a, b in [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5), (5, 6), (6, 7), (7, 4)]:
        add_arrow(page, centers[a][0], centers[a][1], centers[b][0], centers[b][1])
    add_note(page, 0.9, 0.55, 3.3, 0.35, "设计约束：编辑不直接覆盖原图，必须留下可回退版本。")
    return doc, page


def add_docx_section(exports: list[tuple[str, Path]]):
    document = Document(DOCX)
    marker = "7. 补充详细设计图件"
    if any(marker in p.text for p in document.paragraphs):
        return
    document.add_page_break()
    document.add_heading(marker, level=1)
    p = document.add_paragraph(
        "为提升图件的可实现性，本节补充五张详细设计图，分别覆盖目录扫描入库、子系统构件、AI 标签与搜索安全链路、数据库初始化降级以及图片编辑版本历史。"
        "这些图件均保留为可编辑的 VSDX 文件，并已在提交包中一并提供。"
    )
    p.paragraph_format.first_line_indent = Cm(0.74)
    for idx, (caption, png) in enumerate(exports, start=1):
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(png), width=Cm(15.5))
        cap = document.add_paragraph(f"图7-{idx} {caption}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(10)
    document.save(DOCX)


def main() -> int:
    if EXPORT_DIR.exists():
        shutil.rmtree(EXPORT_DIR)
    EXPORT_DIR.mkdir(parents=True)

    app = win32.DispatchEx("Visio.Application")
    app.Visible = False
    app.AlertResponse = 7
    exports: list[tuple[str, Path]] = []
    diagrams = [
        ("目录扫描入库详细活动图", "图11_目录扫描入库详细活动图.vsdx", diagram_scan),
        ("子系统与构件职责展开图", "图13_子系统与构件职责展开图.vsdx", diagram_components),
        ("AI 标签扫描与智能搜索安全链路图", "图14_AI标签扫描与智能搜索安全链路图.vsdx", diagram_ai),
        ("数据库初始化与离线降级流程图", "图15_数据库初始化与离线降级流程图.vsdx", diagram_bootstrap),
        ("图片编辑与版本历史详细设计图", "图16_图片编辑与版本历史详细设计图.vsdx", diagram_versions),
    ]
    try:
        for caption, filename, builder in diagrams:
            doc, page = builder(app)
            vsdx = EXP3 / filename
            png = EXPORT_DIR / (Path(filename).stem + ".png")
            save_doc(doc, page, vsdx, png)
            exports.append((caption, png))
    finally:
        app.Quit()

    add_docx_section(exports)
    print("generated", len(exports))
    for caption, png in exports:
        print(caption, png)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
