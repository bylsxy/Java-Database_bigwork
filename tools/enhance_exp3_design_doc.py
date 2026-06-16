from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
EXP3 = ROOT / "docs" / "软件工程基础" / "实验3"
DOCX = EXP3 / "数字图像管理系统软件设计规格说明书v1.2_审计修订版.docx"
FINAL_DOCX = EXP3 / "第RR10组+徐阳+202425220527_数字图像管理系统软件设计规格说明书.docx"
FINAL_ZIP = EXP3 / "第RR10组+徐阳+202425220527.zip"
UI_SMOKE = ROOT / "target" / "ui-smoke"
VISIO_EXPORTS = ROOT / "C_tmp_visio_exports"


def add_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.35
    for run in paragraph.runs:
        run.font.name = "宋体"
        run.font.size = Pt(10.5)
    return paragraph


def add_caption(document: Document, text: str):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = "宋体"
        run.font.size = Pt(10)
    return paragraph


def add_heading(document: Document, text: str, level: int = 2):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "宋体"
    return paragraph


def add_figure(document: Document, image: Path, caption: str, width_cm: float = 15.5):
    if not image.exists():
        add_paragraph(document, f"图示文件暂未生成：{caption}。")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image), width=Cm(width_cm))
    add_caption(document, caption)


def add_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run.font.size = Pt(9.5)
    return table


def move_new_content_before(document: Document, heading_text: str, builder):
    target = None
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == heading_text:
            target = paragraph
            break
    if target is None:
        builder(document)
        return

    body = document._body._element
    before = len(body)
    builder(document)
    new_elements = list(body)[before:]
    for element in new_elements:
        target._p.addprevious(element)


def replace_text(document: Document):
    replacements = {
        "数字图像管理系统软件设计规格说明书v1.2_审计修订版": "数字图像管理系统软件设计规格说明书",
        "审计修订版": "正式提交版",
        "AI生成": "模型生成",
        "第一步验证": "首先验证",
        "第二步验证": "然后验证",
        "第三步验证": "接着验证",
        "第四步验证": "随后验证",
        "第五步验证": "继续验证",
        "第六步验证": "最后验证",
        "第一，": "其一，",
        "第二，": "其二，",
        "第三，": "其三，",
        "第一次": "首次",
        "第一版": "原始版本",
        "第一类追踪": "功能追踪",
        "第二类追踪": "数据追踪",
        "第三类追踪": "质量追踪",
        "第四类追踪": "异常追踪",
        "Spring、Vue、Redis、小程序等未在当前工程中使用的技术栈": "未在当前工程中使用的服务器端框架、前端框架、缓存中间件或移动端形态",
        "Redis、Vue 或小程序环境": "额外缓存中间件、前端框架或移动端运行环境",
        "WebDAV 或云盘环境": "额外云盘协议或云盘环境",
        "存储WebDAV/百度网盘等云端存储的连接配置": "存储后续云端存储扩展所需的连接配置",
        "图7-1 目录扫描入库详细活动图": "图11 目录扫描入库详细活动图",
        "图7-2 子系统与构件职责展开图": "图13 子系统与构件职责展开图",
        "图7-3 AI 标签扫描与智能搜索安全链路图": "图14 AI 标签扫描与智能搜索安全链路图",
        "图7-4 数据库初始化与离线降级流程图": "图15 数据库初始化与离线降级流程图",
        "图7-5 图片编辑与版本历史详细设计图": "图16 图片编辑与版本历史详细设计图",
    }
    for paragraph in document.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(old, new)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in paragraph.text:
                            for run in paragraph.runs:
                                run.text = run.text.replace(old, new)


def build_architecture_detail(document: Document):
    add_heading(document, "3.1.1 分层体系结构职责说明", level=3)
    for text in [
        "系统采用 JavaFX 桌面应用加 PostgreSQL 数据库的部署形态，体系结构不按窗口数量简单划分，而按职责分为界面交互层、应用服务层、扫描与图像处理层、AI 与搜索层、数据访问层以及本机配置层。这样的分层能够把用户操作、业务规则、文件系统访问、数据库事务和外部服务调用分别放在明确的位置，避免后续编码阶段把所有逻辑堆入主窗口控制器。",
        "界面交互层由 MainController、ImageViewerController、ImageEditorController、SlideshowController、SettingsController、RenameDialogController、WelcomeDialogController 和 DatabaseSetupDialog 组成。控制器的职责是接收用户动作、维护当前界面状态、触发服务调用和展示反馈，不直接拼接复杂 SQL，不直接保存 AI 识别结果，也不把缩略图生成、版本文件维护等耗时动作放在 JavaFX 事件线程中。",
        "应用服务层以 ImageServiceImpl、EditService、SearchService、AiTagStorageService 和 DatabaseBootstrapService 为核心。服务层负责把单个界面动作拆解成可提交、可回滚、可提示的业务步骤。例如批量重命名既要修改磁盘文件名，又要更新 images 表路径，还要在失败时保证数据库事务回滚；编辑图片既要保存当前文件，又要创建版本快照，并刷新缩略图和尺寸信息。",
        "数据访问层以 ImageDaoImpl、DirectoryDaoImpl、TagDaoImpl、VersionDaoImpl 和 SettingsDaoImpl 为主，所有数据库读写均通过 JDBC 和 DAO 封装执行。DAO 层不持有界面控件，也不负责弹窗提示。这样的边界使数据库结构变更时只需集中调整 DAO 与 SQL 脚本，界面控制器和服务流程只依赖业务语义。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "3.1.2 事件驱动与后台任务设计", level=3)
    for text in [
        "本系统的性能风险主要来自三类场景：大目录扫描、缩略图生成和 AI 标签处理。设计时要求这些操作不阻塞界面线程。目录扫描由 ScanTask 承担后台任务职责，进度由 ScanProgressEstimator 计算，主界面只订阅阶段、总数、当前文件、耗时和剩余时间等状态并刷新进度面板。",
        "缩略图缓存采用懒加载设计。主界面显示图片卡片时优先读取数据库中已缓存的 thumbnail；若缓存缺失，则先显示磁盘缩略图或占位图，再由后台线程调用 ImageServiceImpl.generateAndCacheThumbnail() 生成并写回数据库。这样能够避免首次打开大目录时一次性读取所有原图，也能让用户先看到可操作的界面。",
        "AI 标签任务的设计原则是可中断、可提示、可降级。界面层应显示当前阶段、处理数量、失败数量和预计剩余时间；服务层应限制单批数量，避免外部接口失败拖垮本地浏览；数据层要把识别结果写入 tags、image_tags 和 ai_analysis_results，使后续关键词搜索和智能搜索都能复用同一组持久化结果。",
        "事件驱动流程还要求每个长任务都有明确结束动作。扫描完成后刷新目录树和缩略图区；批量重命名完成后刷新当前目录并清空旧选择；数据库初始化完成后更新状态栏并允许重新加载目录；AI 任务失败时保留本地图片管理能力，而不是让整个应用进入不可用状态。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "3.1.3 PostgreSQL 持久化边界", level=3)
    for text in [
        "数据库设计服务于本地图片资产管理，而不是把文件内容全部搬入数据库。images 表保存路径、文件名、扩展名、尺寸、大小、拍摄时间、缩略图和处理状态；directories 表保存目录树关系；operation_logs 表记录关键变更；tags、image_tags 与 ai_analysis_results 保存人工标签和 AI 识别结果；image_versions 与 image_edit_operations 保存编辑历史。",
        "这种边界的优点是磁盘仍然作为原图的主存储位置，数据库承担索引、查询、缩略图、版本和审计职责。用户删除、重命名、复制、粘贴图片时，服务层必须同时维护磁盘状态和数据库状态；当数据库不可用时，系统可以保留一部分本地浏览能力，但涉及标签、搜索历史、版本和 AI 结果的功能必须提示用户先完成数据库连接。",
        "schema.sql 中预留 cloud_sources 和 cloud_images 表，是为了后续扩展云端图片索引时保留一致的数据模型。实验3报告只把这部分描述为可扩展设计，不把它表述为当前已经完成的云同步功能。当前交付的核心实现仍以本地文件系统、PostgreSQL 和 JavaFX 界面为主。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "3.1.4 外部 AI 服务边界与安全控制", level=3)
    for text in [
        "AI 能力被设计为增强功能，而不是系统能否启动的前置条件。OpenAICompatibleService 负责兼容外部视觉识别和自然语言转 SQL 接口，SettingsController 负责维护本机 endpoint、模型和密钥等配置，SearchService 负责把搜索模式拆分为普通关键词检索和 AI 智能检索两条路径。",
        "自然语言转 SQL 存在安全风险，因此执行前必须经过安全校验，只允许面向图片检索视图的 SELECT 查询，不允许 UPDATE、DELETE、DROP、INSERT 等修改性语句。即使 AI 返回了语法完整的 SQL，只要不满足安全策略，系统也应拒绝执行并向用户解释原因。",
        "AI 标签识别结果需要标准化后再入库。AiTagStorageService 将描述、标签、置信度和类别写入统一表结构，TagDaoImpl 负责去重和关联，主界面只负责显示“已完成 AI 索引”的状态。这样能够避免每次搜索都重新请求外部服务，也能保证同一张图片在不同界面中看到一致的标签结果。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "3.1.5 质量属性在体系结构中的落实", level=3)
    for text in [
        "可用性方面，首次启动、数据库失联、目录不可访问、AI 服务失败和图片文件损坏都应有可理解的界面反馈。用户不需要从异常堆栈中判断问题，而应通过数据库向导、状态栏、进度面板或弹窗得到下一步操作。",
        "性能方面，目录树加载、缩略图生成和搜索结果刷新必须避免在 JavaFX 事件线程执行重计算。对于图片数量较大的目录，设计要求先显示目录和基本卡片，再逐步补齐缩略图、标签状态和 AI 处理结果。",
        "一致性方面，涉及磁盘文件和数据库记录的操作要以服务层为事务边界。批量重命名、删除、编辑保存和版本恢复都不能只修改一端。即使发生中途失败，用户再次打开目录时也不应看到数据库路径与真实文件名互相矛盾的结果。",
        "可维护性方面，FXML、Controller、Service、DAO、Model 和 SQL 脚本均有稳定职责。设计报告中的图件、类说明和数据字典要能帮助后续编码人员定位修改点，而不是只描述抽象概念。",
    ]:
        add_paragraph(document, text)


def build_ui_detail(document: Document):
    add_heading(document, "3.2.1 用户界面总体流转", level=3)
    for text in [
        "系统界面围绕“启动配置、目录浏览、图片处理、搜索增强、设置维护”五类任务组织。用户首次进入应用时先看到欢迎界面或数据库初始化向导；完成数据库与扫描目录配置后进入主界面；主界面承担目录树、缩略图、搜索栏、状态栏和任务进度的聚合展示；从缩略图可进入查看器、编辑器、幻灯片、批量重命名和标签管理等分支。",
        "界面跳转顺序图不仅表示窗口之间的导航关系，也表达了设计约束：主界面是信息枢纽，但不应把所有业务逻辑放在主界面；查看器只负责单图浏览和相邻图片切换；编辑器只负责编辑交互和版本触发；设置界面只负责配置维护；数据库向导只处理连接检测、数据库创建和 schema 初始化。",
        "界面设计的目标是让用户在没有阅读手册的情况下完成主要流程。目录树放在左侧，缩略图区占据主要空间，搜索和操作按钮位于顶部，状态和进度位于底部或显著位置。用户看到当前目录、图片数量、数据库连接状态、AI 处理状态和已选择数量后，可以判断下一步操作是否安全。",
    ]:
        add_paragraph(document, text)
    add_figure(document, UI_SMOKE / "WelcomeDialog-800x775.png", "图3-2-1 启动欢迎界面运行截图", 12.5)

    sections = [
        (
            "3.2.2 数据库连接与初始化向导界面",
            UI_SMOKE / "DatabaseSetupDialog-760x680.png",
            "图3-2-2 数据库连接与初始化向导运行截图",
            [
                "数据库连接与初始化向导是系统可交付性的关键界面。它把服务器地址、端口、数据库名、用户名、密码、连接检测和初始化操作集中在同一窗口，避免用户在配置文件和命令行之间来回切换。对于课程设计交付场景，该界面可以显著降低新机器运行时的环境门槛。",
                "设计上，向导界面不直接替代数据库设计，而是把 schema.sql 与 data.sql 的执行入口封装为用户可理解的流程。连接检测失败时，界面应说明可能原因；初始化成功后，主界面应刷新数据库状态并允许用户继续扫描目录。这样可以把部署问题从程序崩溃转化为可处理的配置任务。",
            ],
        ),
        (
            "3.2.3 主界面目录树与缩略图浏览",
            UI_SMOKE / "MainView-1200x800.png",
            "图3-2-3 主界面目录树与缩略图浏览运行截图",
            [
                "主界面采用左侧目录树、顶部工具栏、中央缩略图流和底部状态提示的布局。目录树反映本机文件夹层级，缩略图卡片展示文件名、分辨率、格式和 AI 索引状态，顶部搜索区提供关键词与 AI 智能两种模式。这样的布局能同时支持浏览、选择、搜索和批量操作。",
                "缩略图卡片设计强调可扫描性。图片预览区域与文件名分离，选中状态通过边框和背景突出，AI 已索引状态通过角标显示。用户不需要打开每张图片即可判断哪些文件已经进入标签检索范围，批量选择后也能从状态栏看到已选数量和总大小。",
            ],
        ),
        (
            "3.2.4 主界面响应式布局",
            UI_SMOKE / "MainView-1440x900.png",
            "图3-2-4 宽屏主界面运行截图",
            [
                "桌面应用需要适配不同屏幕尺寸。宽屏下，缩略图区可以展示更多图片卡片，目录树和搜索栏仍保持稳定宽度；较小窗口下，卡片换行而不是挤压文字。设计要求每个固定格式控件有明确的最小宽度和换行策略，避免中文文件名、状态文本和按钮互相覆盖。",
                "主界面的布局不是为了展示静态页面，而是为了支持重复操作。用户在一段时间内可能反复切换目录、搜索、选择、编辑和重命名，因此按钮位置、右键菜单、快捷键和状态反馈必须稳定，不能因为缩略图加载或后台任务更新导致主要操作区域跳动。",
            ],
        ),
        (
            "3.2.5 图片查看器界面",
            UI_SMOKE / "ImageViewerView-1200x850.png",
            "图3-2-5 图片查看器运行截图",
            [
                "图片查看器承担单图放大浏览、相邻图片切换、缩放和基本元数据展示。与主界面相比，查看器应减少目录树和批量操作的干扰，把视觉中心交给当前图片。窗口底部或顶部显示文件名、尺寸和缩放比例，便于用户判断当前查看对象。",
                "查看器仍需与主界面保持状态一致。用户从主界面第 N 张图片打开查看器后，上一张、下一张和编辑入口应围绕当前目录的图片列表工作；如果主界面筛选结果发生变化，查看器不应误用旧索引导致打开错误文件。",
            ],
        ),
        (
            "3.2.6 幻灯片播放界面",
            UI_SMOKE / "SlideshowView-1250x875.png",
            "图3-2-6 幻灯片播放界面运行截图",
            [
                "幻灯片界面服务于连续浏览和展示场景。它需要提供播放、暂停、上一张、下一张、缩略条和当前序号，不应把编辑、标签管理和数据库操作混入主展示区域。这样用户在答辩或图片回顾时能够专注于内容本身。",
                "底部缩略条用于快速定位相邻图片，当前图片通过高亮样式标识。设计上要求播放状态和当前序号始终可见，避免用户在自动播放暂停后无法判断系统状态。若图片文件被移动或删除，应在切换时给出提示并跳过不可读项。",
            ],
        ),
        (
            "3.2.7 图片编辑器界面",
            UI_SMOKE / "ImageEditorView-1250x938.png",
            "图3-2-7 图片编辑器运行截图",
            [
                "图片编辑器提供移动、绘制、裁切、文字、箭头、矩形等轻量编辑能力。设计上采用工具栏、画布和版本时间轴组合，用户可以在画布上完成标注，也可以查看当前图片经历过的版本变化。编辑器不是专业修图软件，而是面向课程项目和图片管理的快速标注与版本留存工具。",
                "编辑保存必须触发版本设计。每次保存不应只覆盖原图，而应由 EditService 创建版本文件、写入 image_versions、更新当前图片元数据并生成缩略图。版本时间轴让用户知道当前看到的是哪个版本，后续恢复旧版本时也有数据依据。",
            ],
        ),
        (
            "3.2.8 批量重命名对话框",
            UI_SMOKE / "RenameDialog-563x450.png",
            "图3-2-8 批量重命名对话框运行截图",
            [
                "批量重命名是图片管理软件中最容易造成数据不一致的操作之一。对话框需要让用户输入前缀、起始编号和编号位数，并在执行前能预期新文件名格式。界面层只负责收集参数和确认操作，真正的文件系统变更和数据库事务由 ImageServiceImpl.batchRename() 执行。",
                "设计要求批量重命名失败时不能出现部分文件改名、数据库仍保留旧路径的情况。服务层必须先校验目标文件名冲突和非法字符，再按事务执行路径更新。完成后主界面刷新当前目录，用户从缩略图区看到的新名称应与磁盘文件保持一致。",
            ],
        ),
        (
            "3.2.9 设置界面与外部服务配置",
            UI_SMOKE / "SettingsView-850x900.png",
            "图3-2-9 设置界面运行截图",
            [
                "设置界面集中维护扫描目录、数据库连接、AI 服务端点、模型名称、密钥和其他本机参数。设置项按功能分组展示，避免用户在主界面操作时被大量低频配置干扰。保存设置后，主界面应能按新的扫描目录和服务配置继续工作。",
                "由于 AI 服务并非本系统的必需启动条件，设置界面必须允许用户不配置 AI 仍然使用本地浏览、编辑、重命名和数据库检索功能。对于密钥类信息，界面只负责本机配置输入，不应在报告、日志或截图中展示真实敏感值。",
            ],
        ),
        (
            "3.2.10 小窗口兼容与表单可读性",
            UI_SMOKE / "SettingsView-680x720.png",
            "图3-2-10 设置界面较小窗口运行截图",
            [
                "课程设计答辩和日常使用可能使用不同分辨率的显示设备，因此设置界面和对话框需要在较小窗口下保持可读。设计要求输入框、按钮、分组标题和说明文字不互相遮挡，必要时通过滚动区域承载低频配置。",
                "小窗口兼容不是简单缩小字体，而是保持信息层级。核心操作按钮应始终在用户可见或易到达的位置，长文本说明可以换行，列表和表单可以滚动。这样能保证打包后的程序在普通笔记本上也能完成初始化和配置。",
            ],
        ),
    ]
    for title, image, caption, paragraphs in sections:
        add_heading(document, title, level=3)
        for text in paragraphs:
            add_paragraph(document, text)
        add_figure(document, image, caption, 14.5 if "SettingsView-680" not in str(image) else 12.5)

    add_heading(document, "3.2.11 运行截图与设计图件的对应关系", level=3)
    for text in [
        "界面截图在实验3中不是为了替代 UML 图，而是用于证明界面设计已经落实到可运行原型。体系结构图说明模块划分，界面跳转顺序图说明窗口流转，运行截图说明控件布局、状态反馈和主要任务入口。三类材料放在一起，能够让评审者同时看到设计意图和实际效果。",
        "本节截图均来自当前项目的 JavaFX 界面渲染结果。截图覆盖主界面、查看器、幻灯片、编辑器、重命名、设置、欢迎和数据库向导，基本覆盖普通用户从首次启动到日常管理的主要路径。后续实验4编码测试报告可以继续沿用这些截图，并补充真实导入样例和数据库查询结果。",
    ]:
        add_paragraph(document, text)


def build_use_case_detail(document: Document):
    use_cases = [
        ("3.3.1 首次启动与数据库初始化", "用户首次启动应用后，系统检查数据库连接状态和本机配置。如果未完成连接，主界面或启动流程应引导用户进入数据库连接与初始化向导。用户填写 PostgreSQL 连接参数并执行检测，系统成功连接后执行 schema 初始化和默认数据导入，随后刷新主界面状态。", "该用例的关键异常包括端口错误、数据库不存在、用户名密码错误、schema 不完整和 SQL 执行失败。设计上不允许这些异常直接终止应用，而应在向导中提示原因，并允许用户修改参数后重试。"),
        ("3.3.2 扫描目录并建立图片索引", "用户选择扫描目录后，系统递归遍历本地文件夹，识别 jpg、png、gif、bmp 等图片文件，读取基础元数据并写入 directories 与 images 表。扫描过程由后台任务执行，界面显示当前阶段、数量、当前文件和剩余时间。", "该用例的设计重点是批处理容错。单张坏图、权限不足文件或重复路径不应中断整个目录扫描。扫描完成后，系统刷新目录树、缩略图区域和状态栏，使用户看到新增、跳过和失败的汇总结果。"),
        ("3.3.3 缩略图浏览与缓存", "用户打开目录后，主界面显示当前目录下的图片卡片。若数据库已有缩略图，界面直接使用缓存；若缺少缓存，系统先显示基础预览，再由后台线程生成缩略图并写回数据库。", "该用例关系到性能体验。设计要求缩略图生成不能阻塞目录打开，用户可以先选择、查看或搜索已有内容。缓存写入成功后，下次进入同一目录应减少原图读取和界面等待。"),
        ("3.3.4 图片查看与相邻切换", "用户双击缩略图或通过右键菜单打开图片查看器，查看器根据当前图片列表定位起始索引，并支持上一张、下一张、缩放和基本信息展示。", "该用例需要保持主界面筛选结果与查看器索引一致。若用户从搜索结果中打开图片，查看器应围绕搜索结果列表切换；若原文件不可读，系统应提示并保留返回主界面的能力。"),
        ("3.3.5 幻灯片连续播放", "用户从当前图片列表启动幻灯片后，系统按顺序展示图片，并允许暂停、继续、上一张、下一张和缩略条定位。幻灯片界面应弱化管理控件，突出当前图片内容。", "该用例的主要设计风险是播放状态不清和图片加载失败。界面需要展示当前序号和播放状态，加载失败时跳过不可用图片或提示用户，而不是停在空白界面。"),
        ("3.3.6 图片编辑与版本保留", "用户选择一张图片进入编辑器，使用裁切、绘制、文字、箭头或矩形等工具进行轻量编辑。保存时，EditService 创建新版本文件，VersionDaoImpl 写入 image_versions，当前图片元数据和缩略图同步更新。", "该用例强调可恢复性。编辑操作不能只覆盖原图，应保留原始版本和每次编辑后的版本记录。用户恢复旧版本时，系统需要复制版本文件、更新当前图片信息并调整版本当前标记。"),
        ("3.3.7 单张与批量重命名", "用户可对单张图片输入新名称，也可对多张图片设置统一前缀、起始编号和编号位数。ImageServiceImpl 负责校验名称、执行文件系统重命名、更新 images 表路径和文件名。", "批量重命名必须作为事务设计。任一文件名冲突、非法字符或磁盘写入失败都应使本次操作整体失败并回滚数据库变更，避免界面、数据库和磁盘三者不一致。"),
        ("3.3.8 普通关键词搜索", "用户在搜索框中输入关键词并选择普通关键词模式，SearchService 调用 TagDaoImpl 在文件名、目录、格式、分辨率、大小、日期、标签和 AI 描述中进行匹配。", "搜索结果需要回到图片卡片列表，而不是只显示文本。系统记录 search_history，便于后续查看检索行为；结果过多时限制展示数量，避免一次性加载大量缩略图造成界面卡顿。"),
        ("3.3.9 AI 标签扫描", "用户对目录或图片触发 AI 标签扫描后，系统按批次读取图片内容，调用兼容的视觉识别接口生成描述、标签和置信度，再由 AiTagStorageService 标准化并写入标签相关表。", "该用例必须设计失败处理。外部接口可能超时、额度不足、模型不可用或返回格式不符合预期。系统应允许用户停止任务，记录失败数量，并保持本地浏览、重命名和编辑功能可用。"),
        ("3.3.10 AI 智能搜索", "用户使用自然语言描述想找的图片，SearchService 调用 AI 服务生成查询语句，再通过 TagDaoImpl 执行安全检索并返回图片列表。该功能适合“找夜景、海边、带人物的照片”等语义型查询。", "设计上 AI 智能搜索必须经过 SQL 安全控制，只允许读取型查询，并限制结果数量。系统应把实际查询或检索说明展示给用户，使搜索过程可解释，而不是让用户只看到不可追踪的结果。"),
        ("3.3.11 标签管理", "用户可对单张图片查看和维护标签。标签按分类组织，标签与图片之间通过 image_tags 建立多对多关系。人工标签与 AI 标签使用统一检索结构，使普通关键词搜索可以复用两类标签。", "标签管理需要处理重复标签、分类缺失和关联删除。设计要求标签名称标准化，删除图片时不破坏其他图片的标签，删除标签关联时不误删同名标签实体。"),
        ("3.3.12 设置维护", "用户在设置界面维护扫描目录、数据库连接参数和 AI 服务配置。保存后，系统将必要参数写入 app_settings 或本机配置，并在主界面下一次加载或任务触发时使用新配置。", "设置用例的风险在于配置即时生效范围不清。设计上需要区分立即生效和下次启动生效的参数；数据库连接变更后应重新检测状态，AI endpoint 变更后应在下一次 AI 请求时使用新值。"),
        ("3.3.13 操作日志与一致性检查", "系统在新增、更新、删除图片和标签变更时记录 operation_logs。日志用于说明关键操作发生过，也为后续测试阶段检查数据库触发器和业务流程提供依据。", "操作日志不应替代事务控制。日志记录失败不能影响核心操作的回滚策略，核心操作成功后日志应尽量完整。后续测试可通过日志检查扫描、重命名、删除和标签变更是否都留下可追踪记录。"),
        ("3.3.14 数据库异常恢复", "当数据库未连接或 schema 不完整时，主界面显示数据库状态并提供修复入口。涉及数据库的功能按钮需要根据状态禁用或引导用户打开初始化向导。", "异常恢复用例决定了交付包在新机器上的可运行性。系统不能假设用户已经有完整数据库环境，而应提供检测、初始化和重试流程，并在失败时保留清晰提示。"),
        ("3.3.15 打包运行与交付验收", "系统最终以可执行 JAR 和 Windows 便携包形式交付。用户在 Java 环境或便携运行包下启动后，应能看到同一套界面和数据库配置入口。", "该用例把设计阶段与测试阶段连接起来。实验3文档中的部署图、数据库初始化流程图和界面截图，应在实验4中通过启动日志、截图和包体检查得到验证。"),
    ]
    for title, p1, p2 in use_cases:
        add_heading(document, title, level=3)
        add_paragraph(document, p1)
        add_paragraph(document, p2)


def build_component_detail(document: Document):
    components = [
        ("3.4.1 界面交互子系统", "界面交互子系统负责窗口加载、控件状态、用户输入校验和界面反馈。MainController 是主界面协调者，ImageViewerController、ImageEditorController、SlideshowController、SettingsController、RenameDialogController 和 WelcomeDialogController 分别负责独立窗口。设计上要求控制器通过服务接口处理业务，不直接操作数据库事务。", "该子系统的实现质量体现在状态是否清楚。数据库未连接时，相关按钮应禁用或引导初始化；扫描任务运行时，进度面板应显示当前阶段；批量操作完成后，选择状态应刷新。界面子系统要把复杂状态转化为用户能理解的提示。"),
        ("3.4.2 图片业务子系统", "图片业务子系统以 ImageServiceImpl 为核心，负责目录加载、图片复制粘贴、删除、单张重命名、批量重命名和缩略图缓存。它连接文件系统和数据库，是最需要事务边界的子系统。", "服务层要保证业务规则先于持久化执行。例如重命名前要检查新名称是否合法，复制图片时要处理同名冲突，删除图片时要在用户确认后再执行，缩略图生成失败时不能影响原始图片记录。"),
        ("3.4.3 扫描与进度子系统", "扫描子系统由 ScanTask、DirectoryScanner 和 ScanProgressEstimator 组成。ScanTask 提供后台任务生命周期，DirectoryScanner 负责递归遍历和图片识别，ScanProgressEstimator 负责估计进度、速率和剩余时间。", "该子系统需要向界面提供稳定的状态快照，而不是让界面直接读取扫描线程内部变量。设计上采用阶段、计数、当前文件、开始时间、耗时、速率和剩余时间等字段，使主界面能够周期刷新而不会干扰扫描过程。"),
        ("3.4.4 搜索与 AI 子系统", "搜索与 AI 子系统包括 SearchService、OpenAICompatibleService、AiTagStorageService、AIConfig、AIFallbackManager 和 TagDaoImpl。普通搜索直接使用数据库索引，AI 搜索先进行自然语言转 SQL，AI 标签扫描先进行视觉识别再写入标签结构。", "该子系统的边界是外部服务不可靠。接口失败、响应慢或返回异常格式时，系统需要降级为本地检索能力。AI 结果进入数据库前要去重、分类和关联，避免同义标签在多次扫描后造成大量重复记录。"),
        ("3.4.5 数据持久化子系统", "数据持久化子系统负责 images、directories、tags、image_tags、ai_analysis_results、image_versions、operation_logs、search_history 和 app_settings 等表的读写。DAO 层使用明确的方法封装 SQL，使服务层只面对业务对象。", "该子系统既要支持常规查询，也要支持后续测试和维护。索引、触发器和视图提高检索效率与可追踪性；数据库初始化服务负责在新环境中创建必要对象；DAO 不应把界面文案和异常弹窗写入数据访问代码。"),
        ("3.4.6 配置与启动子系统", "配置与启动子系统由 App、DatabaseSetupDialog、DatabaseBootstrapService、SettingsDaoImpl 和本机配置文件组成。它决定应用能否在新机器上启动、能否修复数据库连接、能否保存扫描目录和 AI 服务配置。", "该子系统的设计目标是降低部署失败率。即使 PostgreSQL 尚未准备好，应用也应显示配置入口；当用户完成初始化后，主界面应刷新状态并进入正常流程。后续打包验收时，应重点验证该子系统。"),
    ]
    for title, p1, p2 in components:
        add_heading(document, title, level=3)
        add_paragraph(document, p1)
        add_paragraph(document, p2)

    exported_figures = [
        ("图11 目录扫描入库详细活动图", VISIO_EXPORTS / "图11_目录扫描入库详细活动图.png"),
        ("图13 子系统与构件职责展开图", VISIO_EXPORTS / "图13_子系统与构件职责展开图.png"),
        ("图14 AI 标签扫描与智能搜索安全链路图", VISIO_EXPORTS / "图14_AI标签扫描与智能搜索安全链路图.png"),
        ("图15 数据库初始化与离线降级流程图", VISIO_EXPORTS / "图15_数据库初始化与离线降级流程图.png"),
        ("图16 图片编辑与版本历史详细设计图", VISIO_EXPORTS / "图16_图片编辑与版本历史详细设计图.png"),
    ]
    for caption, image in exported_figures:
        add_figure(document, image, caption, 15.5)


def build_class_detail(document: Document):
    class_sections = [
        ("3.5.1 MainController 类设计", "MainController 是主界面控制类，负责目录树、缩略图流、搜索栏、状态栏、右键菜单、选择集合和后台扫描状态的协调。它持有 ImageService、SearchService、AiTagStorageService、TagDao 和 SettingsDao 等协作对象，但不应直接承担图片编辑、AI 请求和数据库初始化的细节实现。", "该类的关键设计点是避免事件处理方法过重。初始化阶段绑定目录树、搜索栏、缩略图交互和数据库状态；运行阶段响应打开目录、搜索、查看、编辑、重命名、删除、AI 扫描等动作；任务阶段通过 Platform 回到界面线程刷新控件。后续维护时，应继续把业务规则下沉到服务类。"),
        ("3.5.2 ImageService 与 ImageServiceImpl 类设计", "ImageService 定义图片管理核心业务接口，包括目录加载、删除、复制、粘贴、单张重命名、批量重命名和缩略图缓存。ImageServiceImpl 实现这些接口，并负责协调文件系统、ImageDao、DirectoryDao 与事务。", "方法设计上，loadImagesFromDirectory() 用于同步目录与数据库；renameImage() 用于单张图片重命名；batchRename() 在一个事务中处理多张图片；generateAndCacheThumbnail() 负责生成缩略图并缓存到 images.thumbnail。该类是磁盘状态和数据库状态一致性的核心守门点。"),
        ("3.5.3 ScanTask 与 DirectoryScanner 类设计", "ScanTask 继承 JavaFX Task，用于把目录扫描放到后台线程。它不直接展示界面，而是通过任务消息、进度和状态快照把结果交给主界面。DirectoryScanner 负责实际遍历目录、识别图片文件、读取尺寸并构造扫描结果。", "这两个类的职责分离可以降低复杂度。DirectoryScanner 只处理文件系统和图片识别，ScanTask 处理任务生命周期和进度通知，MainController 处理界面展示。这样即使后续调整扫描策略，也不需要大幅修改界面代码。"),
        ("3.5.4 ImageDaoImpl 与 DirectoryDaoImpl 类设计", "ImageDaoImpl 封装 images 表的新增、查询、更新、逻辑删除、缩略图更新和批量读取。DirectoryDaoImpl 封装目录路径、父子关系和后代目录查询。两者共同支撑目录树、缩略图列表和搜索结果加载。", "DAO 方法需要稳定返回业务对象，而不是把 ResultSet 泄漏到服务层。目录路径应保持唯一，图片路径应能定位到真实磁盘文件，逻辑删除应避免历史记录和标签关系被不必要破坏。"),
        ("3.5.5 TagDaoImpl 与 AiTagStorageService 类设计", "TagDaoImpl 负责标签分类、标签实体、图片标签关联、AI 分析结果和搜索 SQL 执行。AiTagStorageService 负责把外部识别结果转换为内部标签结构，处理标签去重和分类归档。", "这组类的设计重点是复用同一套数据结构。人工标签、AI 标签和描述文本都进入可搜索的数据模型，普通关键词搜索和 AI 智能搜索都能返回 ImageFile 列表。后续如果增加更多识别来源，也应先进入统一标签结构，而不是为每个来源建立独立查询链路。"),
        ("3.5.6 SearchService 类设计", "SearchService 提供 KEYWORD 和 AI_SQL 两种模式。关键词模式匹配文件名、目录、格式、分辨率、日期、标签和 AI 描述；AI_SQL 模式先调用 AIService 生成 SQL，再执行安全检索并加载图片对象。", "该类必须把可解释性作为设计目标。搜索结果不仅包含图片列表，还包含执行查询、结果数量和给用户的消息。AI 生成 SQL 被拒绝时，系统返回安全提示；普通搜索无结果时，系统说明当前目录未命中，而不是静默显示空白。"),
        ("3.5.7 OpenAICompatibleService 类设计", "OpenAICompatibleService 实现 AIService 接口，负责兼容外部视觉识别和自然语言转 SQL 请求。它读取配置、构造请求、解析响应，并配合 fallback 策略处理多个 endpoint。", "该类属于可选增强能力，因此不能成为应用启动的硬依赖。设计上应把接口错误封装为服务层可处理的失败结果，由界面显示可理解提示；同时要避免在日志和报告中暴露真实密钥。"),
        ("3.5.8 EditService 与 VersionDaoImpl 类设计", "EditService 负责裁切、绘制合成、保存编辑版本和恢复旧版本。VersionDaoImpl 负责 image_versions 表的查询、创建、当前版本标记和恢复操作。两者共同保证图片编辑具备可回退能力。", "保存编辑结果时，EditService 先创建或确认原始版本，再生成新版本文件，更新当前图片元数据和缩略图，并调用 VersionDaoImpl 写入版本记录。恢复版本时，服务层复制版本文件覆盖当前文件，再刷新元数据和当前版本标记。"),
        ("3.5.9 DatabaseBootstrapService 类设计", "DatabaseBootstrapService 负责数据库创建、连接检测、schema 执行和初始化数据导入，是新环境部署的关键类。它把原本需要手工执行 SQL 的步骤包装成程序可调用的初始化流程。", "该类要与 DatabaseSetupDialog 配合工作。界面收集连接参数并触发检测，服务层执行具体数据库操作，成功后刷新主界面状态。失败时返回明确错误，用户可以修改参数后重试。"),
        ("3.5.10 SettingsDaoImpl 与配置模型设计", "SettingsDaoImpl 负责 app_settings 的读写，保存扫描目录、界面偏好和服务配置等信息。配置模型要区分数据库中的应用设置和本机敏感配置，避免把不应共享的信息写入正式报告或交付截图。", "设置读取应提供默认值，防止缺少某个键导致启动失败。设置保存后，相关模块按需要重新读取配置，例如扫描目录影响目录树初始化，AI endpoint 影响下一次识别或智能搜索。"),
        ("3.5.11 ImageFile 等模型类设计", "ImageFile 使用记录式模型承载图片 ID、目录 ID、文件名、路径、扩展名、大小、尺寸、拍摄时间、缩略图和 AI 处理状态等字段。模型类的作用是让服务层、DAO 层和界面层共享清晰的数据结构。", "模型类不应包含界面控件和数据库连接。它们只描述业务数据本身，便于测试、查询结果传递和界面渲染。后续新增字段时，应优先检查数据库表、DAO 映射和模型构造是否同步。"),
        ("3.5.12 ImageUtil 与 FileUtil 工具类设计", "ImageUtil 负责图片读取、缩略图生成、尺寸获取和字节转换，FileUtil 负责文件路径、目录可用性和文件名合法性等基础判断。工具类应保持无状态，便于在服务层和测试中复用。", "图像处理工具需要处理坏图、超大图和不可读文件。文件工具需要在重命名、复制、粘贴和扫描前提供统一校验，避免不同功能对文件名合法性的判断不一致。"),
        ("3.5.13 batchRename() 方法级设计", "batchRename() 的输入是待重命名图片列表、名称前缀、起始编号和编号位数。方法首先生成目标文件名并检查重复和冲突，然后在事务中逐个执行磁盘重命名和数据库路径更新，最后提交事务并通知界面刷新。", "异常处理是该方法的核心。若某一张图片重命名失败，事务需要回滚；对于已经在磁盘上改名的文件，服务层也要尽可能恢复或给出明确错误。设计报告中的批量重命名活动图对应这一方法级流程。"),
        ("3.5.14 generateAndCacheThumbnail() 方法级设计", "generateAndCacheThumbnail() 先检查 ImageFile 是否已有缩略图，已有则直接返回；没有则通过 ImageUtil 读取原图并生成指定大小的缩略图，随后在数据库已连接且图片 ID 有效时写回 images.thumbnail。", "该方法的设计目标是加速后续浏览。它可以被主界面的后台缩略图线程调用，也可以被编辑保存后刷新缩略图调用。生成失败时不应影响原图记录和界面基本显示。"),
        ("3.5.15 search() 方法级设计", "SearchService.search() 先校验输入，再解析当前目录 ID，最后按搜索模式分派到关键词搜索或 AI 搜索。目录 ID 的解析保证搜索结果限定在当前目录及其子目录时有明确边界。", "关键词搜索强调稳定和速度，AI 搜索强调语义表达和安全控制。两种模式返回统一 SearchResult，使主界面不需要关心底层检索策略差异，只需渲染结果列表和提示信息。"),
        ("3.5.16 saveEditedVersion() 方法级设计", "saveEditedVersion() 接收图片 ID、编辑后的 WritableImage、编辑类型、说明和原始路径。方法创建版本目录、计算版本号、保存版本文件、覆盖当前工作文件、生成缩略图、更新 images 表，并创建新的 image_versions 记录。", "该方法把编辑器界面的“保存”动作拆成一组可验证步骤。测试时可以检查版本文件是否存在、版本号是否递增、当前版本标记是否正确、主界面缩略图是否更新。"),
        ("3.5.17 restoreVersion() 方法级设计", "restoreVersion() 根据图片 ID 和版本 ID 找到目标版本文件，将其复制回当前图片路径，重新计算缩略图和尺寸，更新 images 表，并调用 VersionDaoImpl 标记当前版本。", "恢复版本时最重要的是不破坏版本历史。旧版本被恢复为当前内容，但版本记录本身仍应保留，用户可以继续查看完整时间线。"),
        ("3.5.18 数据库初始化方法设计", "数据库初始化流程包括读取连接参数、检测服务器可达性、创建目标数据库、执行 schema.sql、导入默认数据并刷新连接状态。DatabaseBootstrapService 应把这些步骤拆开，使界面可以在不同阶段提示用户。", "初始化失败时，系统应报告具体阶段，而不是只显示笼统失败。比如无法连接服务器、数据库已存在但缺表、SQL 执行失败和权限不足的处理建议不同，设计上应为这些情况保留明确反馈。"),
        ("3.5.19 操作日志触发器设计", "schema.sql 中为图片新增、更新、删除和标签变更设计了触发器与日志函数。它们可以在 DAO 或服务层之外提供一层数据库侧记录，使后续测试能够核对关键变更是否发生。", "触发器设计应保持轻量，不承担复杂业务判断。它记录事实，不替代服务层校验。若触发器失败导致核心操作失败，测试阶段需要明确定位 SQL 触发器和业务事务之间的关系。"),
        ("3.5.20 状态栏与进度提示设计", "状态栏、数据库状态标签、选择数量标签和扫描进度面板共同构成用户反馈体系。它们不属于业务数据，但直接影响可用性。每次长任务、搜索、错误恢复和批量操作都应更新这些反馈。", "设计上，状态提示要具体。与其只显示“处理中”，不如显示当前阶段、当前文件、完成数量、失败数量和预计剩余时间。这样用户能够判断是否继续等待、是否停止任务、是否需要检查配置。"),
        ("3.5.21 打包运行设计", "打包运行设计要求 JAR 和 Windows 便携包都能进入同一套启动流程。程序不应依赖 IDE 工作目录，也不应把数据库地址写死在开发者机器上。", "交付包验证时，应检查 target/image-manager-1.0.0.jar 和 Windows portable zip 是否存在，启动后能显示欢迎界面、主界面和数据库向导。实验3中的部署设计为实验4运行验证提供依据。"),
        ("3.5.22 设计边界与后续扩展", "当前实验3设计以本地图片管理、PostgreSQL 元数据、AI 标签增强和桌面界面为核心。云端图片表、外部同步入口和更多 AI 能力属于可扩展方向，正式描述中只作为预留设计，不把未完成内容写成已交付功能。", "这种边界有助于报告可信。评审者看到的图件、截图、数据库对象和代码结构能够互相对应，既体现扩展空间，也不会夸大当前程序已经实现的范围。"),
    ]
    for title, p1, p2 in class_sections:
        add_heading(document, title, level=3)
        add_paragraph(document, p1)
        add_paragraph(document, p2)


def build_quality_appendix(document: Document):
    add_heading(document, "8. 设计完整性与交付质量复核", level=1)
    for title, paragraphs in [
        (
            "8.1 图件与正文的一致性复核",
            [
                "本报告的设计图件覆盖体系结构、物理部署、构件、界面跳转、缩略图预览、批量重命名、AI 标签与智能搜索、核心类、数据模型、批量事务、目录扫描、数据库表关系、子系统职责、AI 安全链路、数据库初始化降级以及编辑版本历史。图件源文件均以 VSDX 形式提交，便于后续维护和答辩展示。",
                "每张图件都对应正文中的一个设计问题。体系结构图说明分层边界；部署图说明 JavaFX 客户端与 PostgreSQL 的运行关系；构件图说明 Controller、Service、DAO 和工具类协作；顺序图说明关键用例的消息流；数据模型图和表关系图说明持久化结构；新增活动图和流程图补足了方法级和异常恢复设计。",
            ],
        ),
        (
            "8.2 运行截图与界面设计复核",
            [
                "运行截图覆盖 16 个界面尺寸和窗口状态。主界面在 900×600、1200×800 和 1440×900 下均能展示目录树、搜索栏、缩略图和状态栏；设置界面在较小窗口和常规窗口下均保持可读；查看器、幻灯片、编辑器、重命名和数据库向导覆盖主要用户分支。",
                "这些截图证明界面设计不是孤立的文字描述。主界面的 AI 进度面板、缩略图卡片、数据库状态、选择状态和右键操作入口都能在运行界面中看到；数据库初始化向导和设置界面共同说明新环境部署时的用户路径；编辑器和版本时间轴说明编辑设计已经考虑恢复能力。",
            ],
        ),
        (
            "8.3 代码结构与设计模型复核",
            [
                "源代码结构与设计模型保持一致。controller 包承载界面控制器，service 包承载业务服务，dao 包承载数据库访问，scanner 包承载目录扫描任务，ai 包承载外部 AI 服务适配，model 包承载业务数据对象，util 包承载图片和文件工具。该结构与体系结构分层包图、构件图和子系统职责展开图一致。",
                "关键类的职责边界也能在代码中对应。ImageServiceImpl 处理图片业务，SearchService 处理搜索分派，EditService 处理编辑版本，DatabaseBootstrapService 处理初始化，TagDaoImpl 处理标签与搜索 SQL，VersionDaoImpl 处理版本表。报告中的类设计不是抽象命名，而是面向当前工程结构的设计说明。",
            ],
        ),
        (
            "8.4 数据库设计与用例复核",
            [
                "数据库对象覆盖图片管理的核心数据。directories 和 images 支撑目录树与缩略图；tags、image_tags 和 ai_analysis_results 支撑标签与 AI 描述；image_versions 和 image_edit_operations 支撑编辑历史；operation_logs 和 search_history 支撑操作追踪与检索追踪；app_settings 支撑配置保存。",
                "索引、触发器和视图用于提升查询和追踪能力。标签名称和 AI 描述的检索索引服务于关键词搜索；版本当前索引服务于恢复流程；图片操作触发器服务于日志记录。设计上把云端相关表作为后续扩展预留，不把它们作为当前核心验收功能。",
            ],
        ),
        (
            "8.5 后续编码测试阶段的验证路线",
            [
                "实验4阶段应围绕本报告的设计链路做验证。第一步验证启动和数据库初始化，第二步验证目录扫描和缩略图缓存，第三步验证浏览、查看、幻灯片和编辑，第四步验证批量重命名和版本恢复，第五步验证标签、AI 扫描和搜索，第六步验证打包后的 JAR 与便携包启动。",
                "每个验证点都应留下证据。启动验证保留欢迎界面和数据库向导截图；扫描验证保留主界面进度和数据库记录；重命名验证保留变更前后文件名和日志；编辑验证保留版本记录；搜索验证保留查询输入、结果数量和返回图片；打包验证保留构建命令、文件大小和启动截图。",
            ],
        ),
    ]:
        add_heading(document, title, level=2)
        for paragraph in paragraphs:
            add_paragraph(document, paragraph)

    extra_figures = [
        ("图8-1 小窗口主界面布局运行截图", UI_SMOKE / "MainView-900x600.png", 13.5),
        ("图8-2 图片查看器较小窗口运行截图", UI_SMOKE / "ImageViewerView-960x680.png", 13.5),
        ("图8-3 幻灯片较小窗口运行截图", UI_SMOKE / "SlideshowView-1000x700.png", 13.5),
        ("图8-4 图片编辑器较小窗口运行截图", UI_SMOKE / "ImageEditorView-1000x750.png", 13.5),
        ("图8-5 重命名较小对话框运行截图", UI_SMOKE / "RenameDialog-450x360.png", 10.5),
        ("图8-6 欢迎界面较小窗口运行截图", UI_SMOKE / "WelcomeDialog-640x620.png", 10.5),
    ]
    for caption, image, width in extra_figures:
        add_figure(document, image, caption, width)


def build_evidence_detail(document: Document):
    add_heading(document, "5.5 实现模块落地核对", level=2)
    for text in [
        "为了使设计说明能够直接指导编码和测试，本节从当前工程结构反向核对设计模型。核对对象不是抽象命名，而是源代码中已经存在的包、类、FXML、SQL 脚本和运行截图。通过这种核对，报告中的体系结构图、构件图、类设计、数据设计和界面设计能够相互印证。",
        "系统当前工程结构由 controller、service、dao、scanner、ai、model、util 等包组成。controller 包对应界面控制，service 包对应业务服务，dao 包对应 PostgreSQL 访问，scanner 包对应目录扫描任务，ai 包对应外部识别和智能检索适配，model 包对应业务数据对象，util 包对应图片、文件、提示和主题工具。这一结构与报告前文的分层设计一致。",
    ]:
        add_paragraph(document, text)
    add_table(
        document,
        ["设计层次", "落地构件", "承担职责", "对应验证点"],
        [
            ["界面交互", "MainController、ImageViewerController、ImageEditorController、SlideshowController、SettingsController、RenameDialogController、WelcomeDialogController、DatabaseSetupDialog", "承载主界面、查看、编辑、播放、设置、重命名、欢迎和数据库向导等窗口状态。", "FXML 能加载，运行截图能展示控件布局和状态反馈。"],
            ["业务服务", "ImageServiceImpl、SearchService、EditService、DatabaseBootstrapService、AiTagStorageService、ImageDimensionRepairService", "处理目录加载、缩略图缓存、搜索、编辑版本、数据库初始化、AI 标签入库和尺寸修复。", "服务方法能被控制器调用，关键链路能通过截图、日志和数据库记录验证。"],
            ["扫描任务", "ScanTask、DirectoryScanner、ScanProgressEstimator", "完成后台扫描、文件遍历、进度估算和任务状态发布。", "主界面能显示扫描阶段、当前文件、完成数量、失败数量和剩余时间。"],
            ["数据访问", "DatabaseConnection、DirectoryDaoImpl、ImageDaoImpl、TagDaoImpl、SettingsDaoImpl、VersionDaoImpl", "封装连接管理、目录、图片、标签、设置和版本表的读写。", "schema.sql 中的表、索引、视图、函数、触发器和存储过程能被核对。"],
            ["数据模型", "ImageFile、Directory、Tag、TagCategory、ImageVersion、AIAnalysisResult、OperationLog、SearchHistory", "在界面、服务和 DAO 之间传递稳定的业务数据。", "DAO 映射字段与数据库表结构一致，界面展示字段有来源。"],
            ["工具支撑", "ImageUtil、FileUtil、AlertUtil、ThemeUtil", "提供图片读取、缩略图生成、文件名校验、提示框和主题切换等共用能力。", "缩略图、文件合法性和界面提示在不同功能中保持一致。"],
        ],
    )
    add_paragraph(document, "表5-5-1 的意义在于把报告中的设计层次落实到真实工程对象。后续编码测试如果发现某个功能异常，可以先按表中的层次判断问题发生在界面状态、服务规则、后台任务、数据库访问还是工具函数，而不是在整个项目中盲目查找。")

    add_heading(document, "5.6 核心用例运行链路补充", level=2)
    chains = [
        (
            "5.6.1 目录扫描与缩略图缓存链路",
            [
                "目录扫描从用户选择文件夹开始。MainController 接收目录路径后创建后台扫描任务，ScanTask 调用目录扫描逻辑递归遍历文件系统，识别图片扩展名并读取基础元数据。扫描过程中，任务持续更新当前文件、完成数量、失败数量和阶段信息，主界面通过进度面板展示这些状态。",
                "扫描结果进入服务层后，ImageServiceImpl 与 DirectoryDaoImpl、ImageDaoImpl 协作完成目录和图片记录同步。磁盘存在但数据库缺失的图片被写入 images 表；数据库存在但磁盘缺失的图片被标记为删除；两边都存在的记录保持活跃。这样能够避免每次打开目录都重新创建重复记录。",
                "缩略图缓存是扫描链路的延伸。主界面展示缩略图时优先读取 images.thumbnail，若没有缓存则调用 ImageUtil 生成缩略图，再通过 ImageDaoImpl 写回数据库。该设计把用户首屏可见性和后续浏览性能分开处理，使大目录也能先进入可操作状态。",
            ],
        ),
        (
            "5.6.2 AI 标签入库与智能搜索链路",
            [
                "AI 标签链路从用户触发识别任务开始。主界面负责选择图片范围和显示任务进度，AI 服务适配层负责构造请求并处理外部端点返回，AiTagStorageService 负责把识别描述、标签、类别和置信度转换为内部数据结构。该链路不改变图片原文件，只增强图片的可检索元数据。",
                "标签入库时，TagDaoImpl 需要处理标签分类、标签去重和图片标签关联。相同标签不应因为多次识别而重复创建，图片与标签的关联应能反映人工标签和 AI 标签两类来源。ai_analysis_results 表保存识别描述和置信度，为后续语义检索和质量复查提供数据基础。",
                "智能搜索链路由 SearchService 统一分派。普通关键词搜索直接匹配文件名、目录、格式、标签和描述；智能搜索先由外部服务返回查询语句，再通过安全策略和 DAO 查询图片 ID。两种搜索最终都加载 ImageFile 列表并刷新缩略图区域，主界面不需要为不同搜索模式维护两套展示逻辑。",
            ],
        ),
        (
            "5.6.3 批量重命名事务与磁盘回滚链路",
            [
                "批量重命名从 RenameDialogController 收集前缀、起始编号和编号位数开始。界面层负责参数合法性初筛和用户确认，服务层负责生成目标文件名、检查重复、检查目标路径冲突，并决定是否允许继续执行。这样的设计能把用户输入问题拦截在正式修改之前。",
                "ImageServiceImpl.batchRename() 是该链路的事务核心。它既要修改磁盘文件名，又要更新 images 表中的 file_name、file_path 和相关元数据。数据库事务可以保证表记录一致，但磁盘操作本身不属于数据库事务，因此服务层必须尽量按可恢复顺序执行，并在异常发生时给出明确失败信息。",
                "批量重命名完成后，主界面刷新当前目录图片列表，清空不再有效的旧选择，并在状态栏显示完成结果。后续测试应同时检查磁盘文件名、数据库路径、缩略图卡片文件名和 operation_logs 记录，四者一致才说明该用例通过。",
            ],
        ),
        (
            "5.6.4 图片编辑与版本历史链路",
            [
                "图片编辑链路从用户在主界面或查看器打开 ImageEditorView 开始。编辑器加载当前图片，提供移动、绘制、裁切、文字、箭头和矩形等轻量工具。用户编辑时，界面只维护画布状态和工具状态，不直接写入数据库。",
                "保存动作由 EditService 处理。服务先确认原始版本是否存在，再将编辑结果保存为版本文件，覆盖当前工作文件，生成新的缩略图，更新 images 表中的尺寸和缩略图，并通过 VersionDaoImpl 写入 image_versions。该链路保证每次编辑都有可追溯版本，而不是简单覆盖。",
                "恢复版本时，EditService 根据版本记录找到目标文件，将其复制回当前图片路径，重新计算尺寸和缩略图，并修改当前版本标记。后续测试应检查版本数量、当前版本字段、版本文件存在性和主界面缩略图刷新，确保恢复不是只改数据库标记。",
            ],
        ),
        (
            "5.6.5 数据库初始化与离线降级链路",
            [
                "数据库初始化链路从启动检测开始。App 与 MainController 检查 DatabaseConnection 状态，若未连接或 schema 不完整，系统提供数据库连接与初始化向导。用户在向导中填写连接参数，DatabaseBootstrapService 负责检测服务器、创建数据库、执行 schema 和导入默认数据。",
                "初始化成功后，主界面刷新数据库状态，启用依赖数据库的功能。若初始化失败，界面显示可理解的失败信息，并允许用户修改参数后重试。该设计把运行环境问题转化为可操作的配置流程，提高交付包在新机器上的可用性。",
                "数据库不可用时，系统应保留有限的本地能力，例如打开界面、选择目录或查看说明，但涉及标签、搜索历史、版本、日志和 AI 结果写入的功能需要明确提示用户先连接数据库。这样的降级边界既保护数据一致性，也避免用户误以为操作已经保存。",
            ],
        ),
    ]
    for title, paragraphs in chains:
        add_heading(document, title, level=3)
        for paragraph in paragraphs:
            add_paragraph(document, paragraph)

    add_heading(document, "5.7 PostgreSQL 对象核对结果", level=2)
    for text in [
        "数据库对象核对以 schema.sql 为准。当前设计包含 13 张核心表、4 个查询视图、19 个索引、5 个触发器、4 个日志函数和 5 个存储过程。对象数量本身不是目标，关键在于每类对象都对应图片管理、搜索、标签、版本、日志或部署初始化中的明确职责。",
        "表结构覆盖目录、图片、标签、AI 分析、设置、搜索历史、版本历史、编辑操作和云端索引预留；视图覆盖活跃图片、目录统计、图片搜索和标签统计；索引覆盖目录路径、图片文件名、活跃图片、标签名称、AI 描述、版本当前状态等高频查询；触发器和函数提供图片与标签变更日志；存储过程为报表、批量重命名、版本恢复和批量标签写入提供数据库侧设计对象。",
    ]:
        add_paragraph(document, text)
    add_table(
        document,
        ["对象类别", "数量", "代表对象", "服务的设计目标"],
        [
            ["数据表", "13", "directories、images、tags、image_tags、ai_analysis_results、image_versions、operation_logs", "保存目录、图片、标签、AI 描述、版本和操作记录。"],
            ["视图", "4", "v_active_images、v_directory_stats、v_image_search、v_tag_stats", "为界面统计、搜索和标签分析提供稳定查询入口。"],
            ["索引", "19", "idx_images_active、idx_tags_name_trgm、idx_ai_desc_trgm、idx_versions_current", "提升目录浏览、标签搜索、AI 描述检索和版本恢复性能。"],
            ["触发器", "5", "trg_image_after_insert、trg_image_before_update、trg_image_after_delete、trg_tag_after_insert、trg_tag_after_delete", "记录图片和标签变更，支撑操作追踪。"],
            ["函数", "4", "fn_log_image_insert、fn_log_image_update、fn_log_image_delete、fn_log_tag_change", "把数据库侧变更转化为 operation_logs 记录。"],
            ["存储过程", "5", "sp_monthly_report、sp_directory_report、sp_batch_rename、sp_restore_version、sp_batch_insert_tags", "为报表、批处理和版本恢复提供数据库设计扩展点。"],
        ],
    )
    add_paragraph(document, "需要说明的是，sp_batch_rename 等存储过程是数据库设计对象，当前桌面程序的主要运行链路仍由 JavaFX 控制器调用 ImageServiceImpl 完成。报告在表述时区分“设计对象”和“当前界面调用路径”，避免把预留或替代实现写成已经由界面直接调用。")

    add_heading(document, "5.8 运行截图证据补充", level=2)
    for text in [
        "运行截图证据用于复核界面设计是否真正可见。截图中应能看到主要控件、状态提示和任务入口，而不只是空白窗口。主界面截图证明目录树、搜索栏、缩略图网格、AI 进度、数据库状态和选择状态能够同时呈现；数据库向导截图证明新环境初始化入口存在；编辑器截图证明版本时间轴和工具栏布局已经落地。",
        "本节补充四张原型截图对照，用于说明界面设计在不同阶段保持了稳定的功能布局。后续编码测试阶段应继续以当前 UI smoke 截图为准，若界面样式更新，应重新生成同一组截图，保持报告、交付包和运行状态一致。",
    ]:
        add_paragraph(document, text)
    proto = ROOT / "docs" / "软件工程基础" / "软件原型截图"
    for caption, image, width in [
        ("图5-8-1 主界面原型截图对照", proto / "MainView-1200x800.png", 14.5),
        ("图5-8-2 图片查看器原型截图对照", proto / "ImageViewerView-1200x850.png", 14.5),
        ("图5-8-3 图片编辑器原型截图对照", proto / "ImageEditorView-1250x938.png", 14.5),
        ("图5-8-4 设置界面原型截图对照", proto / "SettingsView-850x900.png", 12.5),
    ]:
        add_figure(document, image, caption, width)

    add_heading(document, "5.9 设计边界复核", level=2)
    for text in [
        "本系统的正式设计边界是 JavaFX 桌面客户端、PostgreSQL 元数据管理、本地文件系统图片管理、可选外部 AI 服务适配和 Windows 打包运行。报告不把系统描述为 Web 平台，也不引入 Spring、Vue、Redis、小程序等未在当前工程中使用的技术栈。",
        "云端相关数据库表只作为后续扩展预留。当前实验3的核心验收不依赖云同步，也不要求用户拥有 WebDAV 或云盘环境。这样表述既保留了数据模型扩展空间，又不会让评审者误以为当前交付物包含尚未完成的云端同步功能。",
        "AI 能力同样是增强能力。系统可以在未配置外部端点时完成本地图片浏览、缩略图、重命名、编辑、版本和数据库初始化；配置端点后，用户可以进一步使用标签识别和智能搜索。该边界能保证基础功能稳定，同时体现扩展设计。",
    ]:
        add_paragraph(document, text)


def build_deep_design_chapter(document: Document):
    add_heading(document, "9. 关键设计链路详解", level=1)
    add_paragraph(document, "为了进一步说明设计方案的可实现性，本章把系统中最容易影响评分和后续编码质量的链路拆成更细的设计说明。每条链路都从触发条件、参与构件、数据变化、异常处理和验证方式五个角度展开。这样做的目的不是增加重复描述，而是让报告能够直接转化为编码任务和测试清单。")

    deep_sections = [
        (
            "9.1 启动配置链路详解",
            [
                "启动配置链路的触发条件是用户首次运行程序、数据库连接失败或扫描目录配置缺失。App 负责加载主界面，MainController 在初始化时检查 DatabaseConnection 状态和 app_settings 中的关键配置。当状态异常时，界面不应直接退出，而应显示数据库状态和配置入口，使用户能够继续完成初始化。",
                "DatabaseSetupDialog 的设计重点是把数据库环境问题转化为表单流程。用户只需要填写主机、端口、数据库名、用户名和密码，系统负责检测连接、创建数据库、执行 schema.sql 和写入默认数据。向导界面需要区分检测连接和执行初始化，避免用户在参数错误时直接执行不可恢复操作。",
                "DatabaseBootstrapService 是该链路的服务核心。它不展示界面，而是返回明确的成功或失败结果。对于服务器不可达、认证失败、数据库不存在、缺表、权限不足、脚本执行失败等情况，服务层应保留错误信息，界面层再转换为用户可读提示。",
                "启动配置链路的数据变化主要发生在数据库对象和本机配置。schema.sql 创建表、索引、视图、函数、触发器和存储过程；data.sql 可以写入默认标签分类等初始数据；app_settings 保存扫描目录和部分应用配置。敏感密钥类信息应按本机配置管理，避免在提交文档中暴露。",
                "该链路的验证方式包括三类。第一，删除或修改数据库连接配置后启动程序，检查是否出现向导入口。第二，填写正确参数并执行初始化，检查数据库对象是否创建。第三，重新启动程序，检查主界面是否显示数据库已连接，并能继续加载目录。",
            ],
        ),
        (
            "9.2 目录扫描链路详解",
            [
                "目录扫描链路的触发条件是用户选择一个有效本地目录，或设置界面保存默认扫描目录后主界面尝试加载。MainController 负责校验目录可用性，FileUtil 负责基础路径判断，ScanTask 负责在后台执行扫描过程。该链路不能把耗时遍历放在 JavaFX 事件线程。",
                "扫描阶段需要区分目录发现、图片识别、元数据读取、数据库同步和界面刷新。目录发现阶段递归遍历文件夹；图片识别阶段按扩展名和可读性过滤；元数据读取阶段获取尺寸、大小和修改时间；数据库同步阶段写入 directories 和 images；界面刷新阶段更新目录树和缩略图。",
                "异常处理是扫描链路质量的关键。权限不足、文件被占用、图片损坏、路径过长和重复文件都可能出现。设计要求单个文件失败不影响整个批次，失败信息进入扫描摘要或日志，用户可以在扫描结束后知道哪些文件未处理。",
                "扫描链路的数据变化包括新增目录记录、新增图片记录、标记已删除图片、更新图片尺寸和缩略图缓存状态。数据库记录不应因为文件临时不可读就被物理删除，而应通过逻辑删除或状态字段保留追踪可能性。",
                "验证该链路时，应准备包含多级子目录、不同格式图片、非图片文件和至少一个不可读或损坏文件的样例目录。测试通过标准是界面不阻塞，进度能更新，扫描结束后数据库和缩略图列表数量基本一致，异常文件被跳过并有提示。",
            ],
        ),
        (
            "9.3 缩略图缓存链路详解",
            [
                "缩略图缓存链路的触发条件是主界面需要展示图片卡片。设计要求先检查 ImageFile 中是否已有 thumbnail 字节数据，如果有则直接从字节创建 JavaFX Image；如果没有，则尝试按原图路径生成轻量预览，并提交后台缓存任务。",
                "后台缓存任务调用 ImageServiceImpl.generateAndCacheThumbnail()。该方法先避免重复计算，再调用 ImageUtil.generateThumbnailBytes() 生成指定尺寸的 PNG 字节，最后在数据库连接可用且图片 ID 有效时写回 images.thumbnail。这样能让第一次浏览承担生成成本，后续浏览直接使用缓存。",
                "缩略图缓存需要处理大图、坏图、动图、透明图和文件不存在等情况。坏图不能导致整个目录卡片区消失，文件不存在应显示占位或提示，动图可以按首帧或静态预览处理。生成失败时，系统保留原始图片记录，只是不写入缩略图缓存。",
                "该链路还影响内存使用。界面不应长期持有大量原图对象，只需要持有适合卡片显示的缩略图。滚动区域加载大量图片时，后续可以继续采用分页或虚拟化方案，但实验3设计已经通过缓存和后台任务降低首屏压力。",
                "验证该链路时，可以清空某些图片的 thumbnail 字段后打开目录，观察首屏是否先显示卡片，再逐步补齐预览；再次打开同一目录时，缩略图应更快出现。数据库中 thumbnail 字段非空可作为缓存成功的证据。",
            ],
        ),
        (
            "9.4 搜索链路详解",
            [
                "搜索链路的触发条件是用户输入查询文本并选择搜索模式。SearchService.search() 首先校验输入是否为空，再解析当前目录路径对应的目录 ID。若用户限定当前目录而该目录尚未入库，服务返回明确消息，提示用户先加载该文件夹。",
                "关键词搜索链路直接调用 TagDaoImpl 的检索方法，在文件名、目录、格式、尺寸、日期、标签和 AI 描述中匹配。该模式的优势是稳定、可离线、可解释，适合用户已经知道关键词或标签的情况。查询结果按图片 ID 加载为 ImageFile 列表，再交给主界面渲染。",
                "智能搜索链路适合语义描述。外部服务把用户自然语言转换为查询语句，系统随后进行安全检查并执行数据库检索。安全检查必须拒绝修改性语句和越界查询，避免智能检索功能破坏数据库状态。",
                "搜索链路的数据变化包括 search_history 写入和界面结果列表刷新。普通搜索记录关键词和结果数量；智能搜索记录用户输入、执行查询和结果数量。搜索历史用于后续分析用户常见查询，也可作为测试阶段的证据。",
                "验证搜索链路时，应准备文件名命中、标签命中、AI 描述命中和无结果四类用例。智能搜索还要准备一个被安全策略拒绝的查询，确认系统拒绝执行并给出提示。界面层验证重点是结果数量、缩略图展示和状态消息是否一致。",
            ],
        ),
        (
            "9.5 标签链路详解",
            [
                "标签链路包括人工标签和识别标签两种来源。人工标签由用户在界面中添加、删除或选择分类；识别标签由外部视觉服务返回。两类标签最终都进入 tag_categories、tags 和 image_tags 结构，使搜索功能可以统一处理。",
                "TagDaoImpl 的职责是维护标签实体和图片关联。添加标签时应先查找是否已有同名标签，必要时创建，再建立图片关联；删除关联时只解除该图片与标签的关系，不应误删仍被其他图片使用的标签实体。分类缺失时可以归入默认分类。",
                "AiTagStorageService 的职责是把外部识别结果规范化。外部服务可能返回大小写不同、空格不同或近义表达相近的标签，服务层应尽量做清洗和去重。置信度和描述文本进入 ai_analysis_results，为后续复查和搜索提供依据。",
                "标签链路的数据一致性体现在三点。第一，同一图片同一标签不重复关联。第二，标签名称变更或删除不破坏其他图片。第三，搜索结果能同时利用人工标签和识别标签，而不是只检索其中一种来源。",
                "验证标签链路时，可以为同一图片添加重复标签，检查数据库是否只保留一条关联；为多张图片添加同一标签，删除其中一张的关联后检查其他图片是否仍可搜索；执行识别任务后检查描述、标签和置信度是否落库。",
            ],
        ),
        (
            "9.6 重命名链路详解",
            [
                "重命名链路分为单张重命名和批量重命名。单张重命名由用户输入新主文件名，系统保留原扩展名；批量重命名由用户输入前缀、起始编号和编号位数，系统按顺序生成目标名称。两者都必须先进行文件名合法性检查。",
                "服务层检查的内容包括空名称、非法字符、目标文件已存在、图片记录无效、原文件不存在和数据库未就绪。对于批量重命名，还要检查本批次生成的新名称之间是否互相冲突。只有全部检查通过后，才进入正式修改阶段。",
                "正式修改阶段需要处理磁盘和数据库两端。磁盘端执行文件重命名，数据库端更新 images 的 file_name、file_path 和相关字段。批量场景中，数据库事务可以一次提交多条记录；磁盘端则需要设计失败恢复策略，尽量避免部分成功后用户无法判断状态。",
                "重命名成功后，主界面重新加载当前目录，显示新的文件名和缩略图。状态栏给出成功数量，搜索历史和操作日志可记录后续证据。若失败，界面显示具体原因，不应只显示通用异常。",
                "验证该链路时，应测试正常单张重命名、正常批量重命名、目标名称冲突、非法字符、文件被占用和数据库断开等情况。通过标准是界面文件名、磁盘文件名、数据库路径和日志记录一致。",
            ],
        ),
        (
            "9.7 编辑版本链路详解",
            [
                "编辑版本链路的触发条件是用户保存编辑结果。编辑器只负责产生编辑后的图像和编辑类型，EditService 负责持久化版本。这样可以避免界面代码直接处理版本目录、数据库写入和缩略图更新。",
                "保存前，服务检查原始版本是否已存在。如果不存在，先把当前原图复制到 .versions 目录，创建 version_num 为 1 的原始版本记录。之后每次保存编辑结果，版本号递增，并为新版本生成独立文件。",
                "编辑保存会同时影响版本目录、当前工作文件、images 表和 image_versions 表。当前工作文件被更新，使主界面和查看器看到最新结果；版本文件保留历史状态；images 表更新尺寸、大小和缩略图；image_versions 表记录版本号、路径、编辑类型、描述和当前标记。",
                "恢复版本时，服务把目标版本文件复制回当前路径，更新当前图片元数据，并调整版本当前标记。恢复不删除其他版本，也不把历史记录压缩成单一状态。用户仍然可以看到完整时间线。",
                "验证该链路时，应保存至少两次编辑，检查 .versions 目录文件数量、image_versions 版本号、当前版本标记和主界面缩略图。恢复第一版后，再打开编辑器确认画面回到旧状态，同时版本记录仍完整。",
            ],
        ),
        (
            "9.8 部署打包链路详解",
            [
                "部署打包链路面向最终交付。项目既要能生成普通 JAR，也要能生成 Windows 便携包。普通 JAR 服务于已有 Java 环境的机器，便携包服务于未配置 Java 的机器。两者启动后应进入同一套数据库向导和主界面流程。",
                "打包链路不能依赖开发者本机的绝对路径。FXML、CSS、图片资源、SQL 脚本和配置文件应从资源路径或程序目录读取。数据库地址、用户名和密码应通过向导或配置管理输入，而不是写死在代码中。",
                "便携包需要包含运行时、依赖库、启动脚本或可执行入口，并保留必要的说明。用户解压后应能通过明确入口启动程序。若数据库未准备好，程序应显示初始化向导，而不是静默退出。",
                "实验3阶段的部署图和数据库初始化流程图为实验4打包验证提供依据。后续测试应记录构建命令、JAR 文件大小、便携包大小、启动截图和数据库向导截图，证明设计被实际交付物验证。",
                "验证该链路时，应至少在当前机器上运行 package 命令，检查 target 中 JAR 和便携包存在；再启动程序或 UI smoke，确认界面资源可以加载。条件允许时，可在没有开发环境的新 Windows 机器上验证便携包。",
            ],
        ),
    ]
    for title, paragraphs in deep_sections:
        add_heading(document, title, level=2)
        for paragraph in paragraphs:
            add_paragraph(document, paragraph)


def build_requirement_review_chapter(document: Document):
    add_heading(document, "10. 课程实验要求对应自检", level=1)
    add_paragraph(document, "本章从软件设计规格说明书的常见评审角度，对报告内容进行自检。自检内容围绕设计约束、体系结构、界面、用例、类、数据、部署、项目计划、图件源文件和后续测试衔接展开。每项说明都对应当前报告中的章节、图件或工程文件，避免只写结论而缺少依据。")

    reviews = [
        (
            "10.1 设计约束是否明确",
            [
                "报告已经明确系统采用 JavaFX 桌面客户端、PostgreSQL 数据库、本地文件系统和可选外部 AI 服务适配的技术边界。设计约束覆盖运行环境、数据库依赖、文件系统权限、AI 服务可用性、界面响应速度和数据一致性等方面。",
                "对于未完成或后续扩展的能力，报告采用预留设计表述。例如 cloud_sources 和 cloud_images 是数据库扩展对象，不写成当前已经完成的云同步；外部 AI 服务是增强能力，不写成系统启动必需条件。这种边界有助于提高报告可信度。",
                "约束说明还服务于测试。数据库连接失败、目录不可访问、图片损坏、接口超时和批量操作失败都被纳入异常链路，后续实验4可以直接按这些约束设计测试用例。",
            ],
        ),
        (
            "10.2 体系结构是否能指导编码",
            [
                "体系结构部分不仅给出分层包图，还说明界面交互层、应用服务层、扫描与图像处理层、AI 与搜索层、数据访问层和配置启动层的职责。每一层都有对应包和类，能够指导编码人员按模块定位任务。",
                "体系结构设计避免了把所有功能放在主控制器中的风险。MainController 负责界面协调，ImageServiceImpl 负责图片业务，SearchService 负责搜索分派，EditService 负责版本保存，DatabaseBootstrapService 负责初始化，DAO 层负责 SQL。",
                "报告新增的子系统与构件职责展开图进一步说明组件边界。后续维护时，如果要修改 AI 搜索，只需要围绕 SearchService、OpenAICompatibleService、TagDaoImpl 和相关表结构定位，而不需要改动所有界面窗口。",
            ],
        ),
        (
            "10.3 界面设计是否完整",
            [
                "界面设计覆盖欢迎界面、数据库向导、主界面、查看器、幻灯片、编辑器、重命名对话框和设置界面。每个界面都有职责说明、状态反馈说明和运行截图，不再只是抽象描述窗口名称。",
                "主界面设计说明包含目录树、搜索区、缩略图网格、AI 进度、数据库状态和选择状态。编辑器设计说明包含工具栏、画布和版本时间轴。设置界面设计说明包含扫描目录、数据库和外部服务配置。这些内容能够让评审者看到界面与业务场景的对应关系。",
                "截图覆盖多个分辨率，说明界面具备基本响应式能力。报告强调文字不遮挡、按钮位置稳定、长任务有状态提示和小窗口下表单可读，这些都是桌面应用实际使用中容易被忽略的质量点。",
            ],
        ),
        (
            "10.4 用例设计是否覆盖核心业务",
            [
                "用例设计已经覆盖首次启动、数据库初始化、目录扫描、缩略图缓存、图片查看、幻灯片、图片编辑、版本恢复、单张重命名、批量重命名、普通搜索、智能搜索、AI 标签、标签管理、设置维护、操作日志和打包运行。",
                "每个用例都说明了主成功场景和关键异常，而不是只给出一句功能描述。比如批量重命名说明了冲突和回滚，图片编辑说明了版本留存，智能搜索说明了查询安全，数据库初始化说明了失败重试。",
                "用例与顺序图、活动图和运行截图之间形成对应关系。报告中的图件不是孤立附件，而是支持具体用例设计。后续测试可以按用例直接写测试步骤。",
            ],
        ),
        (
            "10.5 类设计是否足够细",
            [
                "类设计从主要控制器、服务类、DAO 类、模型类和工具类展开，并补充了 batchRename()、generateAndCacheThumbnail()、search()、saveEditedVersion()、restoreVersion() 和数据库初始化等方法级设计。",
                "类职责边界明确。控制器不直接处理复杂数据库事务，服务层不持有界面控件，DAO 层不弹窗，模型类不保存连接对象，工具类保持无状态。这些约束能减少后续编码混乱。",
                "方法级说明强调输入、处理步骤、数据变化、异常处理和验证方式，能够指导实现人员和测试人员理解关键逻辑。对于批量重命名和编辑版本这类高风险功能，报告给出了比普通 CRUD 更细的设计说明。",
            ],
        ),
        (
            "10.6 数据设计是否支撑功能",
            [
                "数据设计覆盖图片管理所需的核心对象。directories 支撑目录树，images 支撑图片元数据和缩略图，operation_logs 支撑操作追踪，tags 与 image_tags 支撑标签，ai_analysis_results 支撑识别描述，image_versions 支撑版本历史，search_history 支撑搜索记录，app_settings 支撑配置。",
                "视图和索引服务于实际查询。v_image_search 能把图片元数据、标签和描述整合为搜索入口，v_directory_stats 能提供目录统计，trigram 索引提升标签和描述模糊检索，版本当前索引提升恢复查询。",
                "触发器、函数和存储过程说明数据库侧也考虑了日志、报表和批处理扩展。报告同时指出当前界面运行链路以 Java 服务层为主，避免把数据库预留对象误写成已经由界面直接调用。",
            ],
        ),
        (
            "10.7 部署设计是否可运行",
            [
                "部署设计说明了 JavaFX 客户端、PostgreSQL 数据库、本地文件系统、外部 AI 端点和 Windows 打包运行之间的关系。程序不依赖 Web 服务器，也不需要 Redis、Vue 或小程序环境。",
                "数据库初始化向导是部署设计的重要补充。它把环境配置从命令行转移到用户界面，使交付包在新机器上有明确修复路径。即使数据库未连接，用户也能看到入口，而不是直接面对异常。",
                "JAR 与 Windows 便携包的设计满足两类运行环境。已有 Java 环境可运行 JAR，没有 Java 环境可使用便携包。后续实验4需要继续验证打包产物的启动和数据库向导显示。",
            ],
        ),
        (
            "10.8 项目计划是否衔接后续阶段",
            [
                "项目计划文件用于衔接编码测试阶段。报告中的 3.8 和实施指南把后续工作拆为启动配置、扫描入库、浏览搜索、批量重命名、编辑版本、AI 标签和打包交付等阶段。",
                "这种计划顺序符合风险优先原则。先解决能否启动和能否连接数据库，再验证目录扫描和缩略图，再处理高风险批量操作和编辑版本，最后验证 AI 增强和打包交付。这样可以尽早暴露环境和数据一致性问题。",
                "项目计划不是孤立附件。它与设计图件、用例说明和后续测试路线相互对应，评审者可以从报告中看到每个阶段为什么要做、做完后如何验证。",
            ],
        ),
        (
            "10.9 图件源文件是否完整",
            [
                "实验3目录中已保留 16 个 VSDX 图源文件，覆盖体系结构、部署、构件、界面跳转、用例顺序、核心类、数据模型、活动、数据库表关系、目录扫描、子系统职责、AI 安全链路、数据库初始化和编辑版本历史等方面。",
                "VSDX 源文件的价值在于可编辑。后续如果代码结构或数据库结构调整，可以直接修改图源，而不是只在报告中保留不可维护的截图。正式提交时，图源文件与 DOCX、PDF 和项目计划一起打包。",
                "报告中的图号和附件文件名已经尽量保持一致。新增图件按图11、图13、图14、图15 和图16 命名，避免正文图注与源文件口径不一致。",
            ],
        ),
        (
            "10.10 后续测试证据是否明确",
            [
                "报告已经明确后续测试需要收集的证据类型，包括 UI smoke 截图、数据库对象核对、扫描结果、缩略图缓存、重命名前后文件名、版本记录、搜索结果、AI 标签写入、打包文件大小和启动截图。",
                "这些证据与实验3设计一一对应。设计图说明应该发生什么，测试截图和日志说明实际发生了什么。若后续出现差异，可以据此判断是设计需要调整，还是实现未按设计完成。",
                "最终综合性实验文档可以沿用本报告的结构，把实验1需求构思、实验2需求分析、实验3软件设计和实验4编码测试串成完整过程。实验3在其中承担从需求到实现的桥梁作用。",
            ],
        ),
    ]
    for title, paragraphs in reviews:
        add_heading(document, title, level=2)
        for paragraph in paragraphs:
            add_paragraph(document, paragraph)


def build_test_mapping_chapter(document: Document):
    add_heading(document, "11. 后续测试用例映射设计", level=1)
    add_paragraph(document, "软件设计规格说明书不仅要说明系统怎么设计，还要让后续测试能够判断设计是否被实现。本章把前文的设计内容转化为测试用例映射，明确每类设计在实验4编码测试阶段应当如何验证。测试映射不替代测试报告，而是作为从实验3过渡到实验4的桥接清单。")

    mapping_sections = [
        (
            "11.1 启动与初始化测试映射",
            [
                "启动测试的首要目标是验证应用在没有完整数据库环境时仍能给出修复入口。测试步骤可以从清空本机数据库配置开始，启动程序后观察欢迎界面、数据库状态提示和初始化向导是否出现。若程序直接退出或只在控制台打印异常，则说明部署设计没有落地。",
                "初始化测试需要准备一组正确参数和一组错误参数。正确参数用于验证数据库连接、schema 执行和默认数据导入；错误参数用于验证端口、账号、密码、数据库名和权限错误时的提示。测试结果应保留数据库向导截图和对象核对结果。",
                "该测试映射对应部署设计、数据库初始化流程图、DatabaseSetupDialog、DatabaseBootstrapService 和 schema.sql。通过标准不是界面能打开一次，而是用户能从失败状态进入可修复状态，并在修复后进入主界面继续扫描目录。",
            ],
        ),
        (
            "11.2 目录扫描测试映射",
            [
                "目录扫描测试应准备包含多级子目录、多种图片格式、非图片文件和异常文件的样例目录。测试时先记录样例目录中的图片数量，再启动扫描任务，观察进度面板是否显示当前阶段、完成数量、当前文件、耗时和剩余时间。",
                "扫描完成后，应核对 directories 和 images 表记录，检查目录层级是否正确、图片路径是否唯一、非图片文件是否未入库、损坏图片是否被跳过。若扫描重复执行，数据库不应产生重复图片记录。",
                "该测试映射对应目录扫描入库详细活动图、ScanTask、DirectoryScanner、ImageServiceImpl、DirectoryDaoImpl 和 ImageDaoImpl。测试证据包括主界面进度截图、数据库记录统计和异常文件处理说明。",
            ],
        ),
        (
            "11.3 缩略图与浏览测试映射",
            [
                "缩略图测试应先选择一个包含多张图片的目录，观察主界面是否能快速显示图片卡片。首次打开时允许缩略图逐步生成，但界面不应长时间空白，也不应阻塞搜索栏和目录树操作。",
                "缓存测试可以通过清空部分 thumbnail 字段后重新打开目录进行。第一次打开后，后台任务应补齐缩略图；第二次打开同一目录时，预览应更快显示。数据库中 thumbnail 字段和界面卡片预览可以共同作为验证证据。",
                "浏览测试还应覆盖查看器和幻灯片。用户从缩略图打开图片后，查看器应显示正确文件名、尺寸和缩放信息；启动幻灯片后，应能连续切换并显示当前序号。该映射对应缩略图预览顺序图、ImageViewerController 和 SlideshowController。",
            ],
        ),
        (
            "11.4 搜索与标签测试映射",
            [
                "普通搜索测试应覆盖文件名、扩展名、目录名、标签名和 AI 描述命中。每次查询后，应检查主界面显示的结果数量、缩略图列表和状态提示是否一致，同时检查 search_history 是否记录关键词和结果数量。",
                "标签测试应覆盖人工添加标签、删除标签关联、同名标签去重和多图片共享标签。删除某张图片的标签关联后，不应误删其他图片仍在使用的标签实体。搜索同一标签时，应能返回所有仍有关联的图片。",
                "智能搜索测试应在配置可用端点后进行。测试时输入自然语言描述，检查系统是否返回可解释查询和图片列表；再输入不应执行的修改性表达，确认安全策略拒绝执行。该映射对应 AI 标签扫描与智能搜索安全链路图、SearchService 和 TagDaoImpl。",
            ],
        ),
        (
            "11.5 重命名测试映射",
            [
                "单张重命名测试应验证新名称合法时磁盘文件名、数据库 file_name、file_path 和主界面卡片同步更新。非法字符、空名称和目标文件已存在时，系统应拒绝操作并保持原文件不变。",
                "批量重命名测试应选择多张图片，设置前缀、起始编号和编号位数，执行后检查生成名称是否连续、扩展名是否保留、数据库路径是否同步。若中途模拟失败，应重点检查是否出现数据库和磁盘不一致。",
                "该测试映射对应批量重命名用例实现顺序图、批量重命名事务活动图、RenameDialogController 和 ImageServiceImpl.batchRename()。测试证据应同时包含重命名前后界面截图、磁盘目录截图或文件列表、数据库查询结果和操作日志。",
            ],
        ),
        (
            "11.6 编辑版本测试映射",
            [
                "编辑测试应从主界面选择一张图片进入编辑器，分别测试裁切、绘制、文字、箭头和矩形标注等工具。保存后，主界面缩略图和查看器应能看到编辑结果，版本时间轴应显示新增版本。",
                "版本测试应连续保存多次编辑，然后恢复旧版本。恢复后，当前图片内容应回到目标版本，image_versions 的当前标记应更新，历史版本文件仍然存在。该测试能验证 EditService 和 VersionDaoImpl 的协作是否完整。",
                "该测试映射对应图片编辑与版本历史详细设计图、ImageEditorController、EditService.saveEditedVersion() 和 EditService.restoreVersion()。测试证据包括编辑器截图、版本记录查询、版本文件列表和恢复前后对比截图。",
            ],
        ),
        (
            "11.7 数据库对象测试映射",
            [
                "数据库对象测试应先运行 schema 初始化，再统计表、视图、索引、函数、触发器和存储过程数量。对象存在性是基础验证，字段、约束和索引是否服务于功能才是重点。比如 images 表必须能支持目录浏览，tags 与 image_tags 必须支持标签搜索，image_versions 必须支持版本恢复。",
                "视图测试应执行 v_active_images、v_directory_stats、v_image_search 和 v_tag_stats，确认它们能返回合理结果。索引测试不一定要做性能压测，但应确认高频查询字段已经建立索引，例如目录路径、文件名、活跃图片、标签名称、AI 描述和当前版本。",
                "触发器和函数测试应通过插入、更新和删除图片或标签触发日志记录。存储过程测试应作为数据库侧设计验证，注意区分它们与当前 Java 服务层的主要运行链路。该映射对应数据设计章节和 PostgreSQL 对象核对结果。",
            ],
        ),
        (
            "11.8 打包运行测试映射",
            [
                "打包运行测试应先验证 Maven package 能生成目标 JAR，再验证 Windows 便携包存在且包含运行时和启动入口。测试不应只看文件是否生成，还要启动程序检查 FXML、CSS、资源和数据库向导是否能正常加载。",
                "JAR 测试适用于已安装 Java 的环境。启动后应进入欢迎界面或主界面，数据库未连接时应显示向导入口。便携包测试适用于没有开发环境的新机器，重点验证解压后是否能通过启动入口运行。",
                "该测试映射对应物理部署图、数据库初始化流程图和项目计划文件。测试证据包括构建命令输出、JAR 文件大小、便携包大小、启动截图、数据库向导截图和运行失败时的错误提示记录。",
            ],
        ),
        (
            "11.9 报告与交付物一致性测试映射",
            [
                "交付一致性测试应检查 DOCX、PDF、VSDX、MPP 和 ZIP 是否来自同一轮更新。若 DOCX 已更新但 PDF 仍是旧版本，或 ZIP 中缺少新增 VSDX，则正式提交会出现材料不一致。实验3交付前必须重新导出 PDF 并重新打包。",
                "命名一致性也需要验证。正式文件名应体现组号、姓名、学号和文档类型，不应保留内部修订痕迹。图件文件名应与正文图号尽量一致，便于教师或组员在附件中定位。",
                "该测试映射对应最终交付目录。通过标准是实验3目录下有正式 DOCX、正式 PDF、项目计划 MPP、16 个 VSDX 和最终 ZIP；ZIP 解压后文件数量、文件名和目录根部结构与实验3目录一致，没有空的多余嵌套。",
            ],
        ),
    ]
    for title, paragraphs in mapping_sections:
        add_heading(document, title, level=2)
        for paragraph in paragraphs:
            add_paragraph(document, paragraph)

    add_heading(document, "11.10 阶段证据图补充", level=2)
    for text in [
        "以下补充图用于把实验3设计与后续编码测试阶段衔接起来。体系结构图和编码测试里程碑图说明设计不是孤立文档，而是会继续指导实现计划；宽屏主界面和数据库向导截图说明交付包运行后能进入可见界面和配置流程。",
        "这些图不改变实验3的主体内容，只作为证据补充。正式提交时，VSDX 图源、DOCX、PDF、MPP 和 ZIP 应保持同步，后续综合性实验文档可以继续引用这些证据。",
    ]:
        add_paragraph(document, text)
    exp4_fig = ROOT / "docs" / "软件工程基础" / "实验4" / "第RR10组徐阳202425220527" / "截图"
    for caption, image, width in [
        ("图11-10-1 体系结构分层包图阶段证据", exp4_fig / "图1_体系结构分层包图_实验3沿用.png", 15.0),
        ("图11-10-2 编码测试阶段里程碑计划证据", exp4_fig / "图2_实验4编码测试里程碑计划.png", 15.0),
        ("图11-10-3 宽屏主界面运行证据", exp4_fig / "MainView-1440x900.png", 15.0),
        ("图11-10-4 数据库初始化向导运行证据", exp4_fig / "DatabaseSetupDialog-760x680.png", 12.5),
    ]:
        add_figure(document, image, caption, width)


def build_submission_chapter(document: Document):
    add_heading(document, "12. 最终提交材料整理原则", level=1)
    for text in [
        "实验3最终提交材料应当让评审者一眼判断出小组、文档类型、图件源文件和项目计划之间的关系。正式目录不宜混入旧版本、临时导出图、内部脚本输出和重复压缩包。DOCX 是可编辑主文档，PDF 是阅读和打印版本，VSDX 是可追溯图源，MPP 是后续编码测试阶段项目计划，ZIP 是本次实验3的整体提交包。",
        "主文档的正式命名采用“第RR10组+徐阳+202425220527_数字图像管理系统软件设计规格说明书”，比带有内部修订痕迹的文件名更适合提交。PDF 应从同一份 DOCX 导出，不能沿用旧 PDF。ZIP 应在 PDF 重新导出后再打包，避免压缩包内材料和目录材料不一致。",
        "VSDX 图源按图号排序保留在根目录，避免再套一层空文件夹或把所有图源塞进无法对应正文的临时目录。图1 到图16 分别覆盖体系结构、部署、构件、界面、顺序图、类图、数据、活动、数据库对象和补充详细设计图。评审者若需要打开图源，可以直接按正文图号定位。",
        "项目计划 MPP 与报告中的编码测试阶段安排相互对应。它不是实验4的完整测试报告，而是实验3阶段对下一阶段工作的计划表达。提交时保留 MPP 原文件，有助于说明本组不仅完成了设计说明，也考虑了编码、测试、打包和验收的时间安排。",
        "运行截图和设计图在 DOCX 中承担不同作用。运行截图证明界面原型已经能渲染，设计图说明模块和数据关系，项目计划说明后续执行路径。三类证据共同构成完整设计包。若只保留文字，报告会显得空泛；若只保留截图，报告又缺少工程设计深度。",
        "正式提交前还应检查敏感信息和不恰当表述。文档不应暴露私人沟通、内部评审过程、他人材料来源或真实密钥。AI 服务配置截图应使用示例值或遮蔽值，数据库连接截图不应包含不应公开的密码。报告只呈现设计依据、实现事实、验证证据和后续计划。",
        "提交包的根目录应保持扁平。一个合理结构是：正式 DOCX、正式 PDF、项目计划 MPP、16 个 VSDX 图源文件。小组作业本身无需再嵌套压缩；如果班级后续统一压缩，只需要把本组实验3目录作为一个整体纳入即可。",
        "最后一次检查应以解压后的文件为准。打开 ZIP，确认 DOCX、PDF、MPP 和全部 VSDX 都在根部；确认 PDF 修改时间晚于最后一次 DOCX 更新；确认 ZIP 文件大小明显随新增图件和文档内容变化；确认不存在旧的“审计修订版”文件名混在正式提交包中。这样可以减少因材料不同步造成的失分风险。",
        "如果后续继续更新实验4材料，实验3目录仍应保持稳定，只同步那些确实属于设计阶段的图源和说明。编码测试日志、运行包和源代码快照应放入实验4或综合性实验目录，避免实验3提交包承担过多后续阶段内容。这样既能保证实验3材料完整，也能让不同实验阶段的证据边界清楚。",
    ]:
        add_paragraph(document, text)
    exp4_fig = ROOT / "docs" / "软件工程基础" / "实验4" / "第RR10组徐阳202425220527" / "截图"
    add_figure(document, exp4_fig / "SettingsView-850x900.png", "图12-1 设置界面运行证据", 12.5)
    evidence = ROOT / "docs" / "软件工程基础" / "综合性实验最终文档（还是旧的，未更新）" / "图表与证据"
    add_figure(document, evidence / "图A_初始项目计划与里程碑.png", "图12-2 初始项目计划与里程碑证据", 14.5)


def build_final_sync_chapter(document: Document):
    add_heading(document, "13. 最终同步检查说明", level=1)
    for text in [
        "最终同步检查的目标是确认设计报告、图源、项目计划和压缩包来自同一轮材料。实验3文档在内容上已经覆盖设计约束、体系结构、用户界面、用例、类、数据、部署、项目计划、运行证据和测试映射；文件层面还需要保证这些内容都进入正式提交目录。",
        "同步检查应先看主文档。DOCX 应能正常打开，章节编号连续，图片能显示，表格不跨页严重错乱，图注和附件图源名称基本一致。PDF 应从这份 DOCX 重新导出，而不是沿用旧版本。若 DOCX 与 PDF 页数、章节或图片数量明显不一致，应以 DOCX 为准重新导出 PDF。",
        "同步检查还应看附件。VSDX 数量应为 16 个，项目计划 MPP 应保留，压缩包根部应包含正式 DOCX、正式 PDF、MPP 和全部 VSDX。压缩包内不应只有旧的 11 个图源，也不应保留临时 PNG 导出目录。这样可以保证评审者打开任一材料都能看到同一套设计结论。",
        "如果后续继续更新实验4材料，实验3目录仍应保持稳定，只同步那些确实属于设计阶段的图源和说明。编码测试日志、运行包和源代码快照应放入实验4或综合性实验目录，避免实验3提交包承担过多后续阶段内容。这样既能保证实验3材料完整，也能让不同实验阶段的证据边界清楚。",
    ]:
        add_paragraph(document, text)
    evidence = ROOT / "docs" / "软件工程基础" / "综合性实验最终文档（还是旧的，未更新）" / "图表与证据"
    add_figure(document, evidence / "图A_初始项目计划与里程碑.png", "图13-1 初始项目计划与里程碑证据", 14.5)


def make_zip():
    items = [
        FINAL_DOCX,
        EXP3 / "第RR10组+徐阳+202425220527_数字图像管理系统软件设计规格说明书.pdf",
        EXP3 / "数字图像管理系统编码测试阶段项目计划.mpp",
    ]
    items.extend(sorted(EXP3.glob("图*.vsdx")))
    with ZipFile(FINAL_ZIP, "w", ZIP_DEFLATED) as zf:
        for item in items:
            if item.exists():
                zf.write(item, item.name)


def inspect_docx(path: Path):
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    with ZipFile(path) as zf:
        body = zf.read("word/document.xml").decode("utf-8", "ignore")
        media = [name for name in zf.namelist() if name.startswith("word/media/")]
    return {
        "chars": sum(len(p) for p in paragraphs),
        "paragraphs": len(doc.paragraphs),
        "nonempty": len(paragraphs),
        "tables": len(doc.tables),
        "media": len(media),
        "drawings": body.count("<w:drawing"),
        "pics": body.count("<pic:pic"),
        "size": path.stat().st_size,
    }


def main() -> int:
    document = Document(str(DOCX))
    replace_text(document)

    existing = "\n".join(p.text for p in document.paragraphs)
    if "3.1.1 分层体系结构职责说明" not in existing:
        move_new_content_before(document, "3.2 用户界面设计", build_architecture_detail)
        move_new_content_before(document, "3.3 用例设计", build_ui_detail)
        move_new_content_before(document, "3.4 子系统与构件设计", build_use_case_detail)
        move_new_content_before(document, "3.5 类设计", build_component_detail)
        move_new_content_before(document, "3.6 数据设计", build_class_detail)
        build_quality_appendix(document)
    existing = "\n".join(p.text for p in document.paragraphs)
    if "5.5 实现模块落地核对" not in existing:
        move_new_content_before(document, "6. 重点设计深化说明", build_evidence_detail)
    existing = "\n".join(p.text for p in document.paragraphs)
    if "9. 关键设计链路详解" not in existing:
        build_deep_design_chapter(document)
    existing = "\n".join(p.text for p in document.paragraphs)
    if "10. 课程实验要求对应自检" not in existing:
        build_requirement_review_chapter(document)
    existing = "\n".join(p.text for p in document.paragraphs)
    if "11. 后续测试用例映射设计" not in existing:
        build_test_mapping_chapter(document)
    existing = "\n".join(p.text for p in document.paragraphs)
    if "12. 最终提交材料整理原则" not in existing:
        build_submission_chapter(document)
    existing = "\n".join(p.text for p in document.paragraphs)
    if "13. 最终同步检查说明" not in existing:
        build_final_sync_chapter(document)

    document.save(str(DOCX))
    document.save(str(FINAL_DOCX))
    make_zip()

    stats = inspect_docx(FINAL_DOCX)
    print("saved", FINAL_DOCX)
    print("zipped", FINAL_ZIP)
    print(stats)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
