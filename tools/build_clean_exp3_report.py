from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

from enhance_exp3_design_doc import (  # noqa: E402
    DOCX,
    EXP3,
    FINAL_DOCX,
    FINAL_ZIP,
    ROOT,
    UI_SMOKE,
    VISIO_EXPORTS,
    add_caption,
    add_figure,
    add_heading,
    add_paragraph,
    add_table,
    build_architecture_detail,
    build_class_detail,
    build_component_detail,
    build_deep_design_chapter,
    build_evidence_detail,
    build_final_sync_chapter,
    build_quality_appendix,
    build_requirement_review_chapter,
    build_submission_chapter,
    build_test_mapping_chapter,
    build_ui_detail,
    build_use_case_detail,
    inspect_docx,
    make_zip,
    replace_text,
)


def setup_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "宋体"
            style.font.bold = True


def add_title_page(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("数字图像管理系统软件设计规格说明书")
    run.bold = True
    run.font.name = "宋体"
    run.font.size = Pt(22)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("第RR10组  徐阳  202425220527")
    run.font.name = "宋体"
    run.font.size = Pt(14)

    add_table(
        document,
        ["项目", "内容"],
        [
            ["课程", "软件工程基础"],
            ["实验", "实验3 软件设计"],
            ["系统名称", "基于 PostgreSQL 的数字图像集成管理系统"],
            ["提交内容", "软件设计规格说明书、项目计划、VSDX 图源、PDF 阅读版"],
            ["技术边界", "JavaFX 桌面客户端、PostgreSQL、本地文件系统、可选外部 AI 服务适配"],
        ],
    )
    for text in [
        "本文档面向软件设计阶段，说明数字图像管理系统在体系结构、用户界面、用例实现、子系统与构件、类设计、数据设计、部署设计和后续编码测试计划方面的设计方案。",
        "文档强调设计与实现的可追踪性。报告中的图件、运行截图、数据库对象和后续测试映射均围绕当前 JavaFX + PostgreSQL 工程展开，避免把未使用的技术栈或后续扩展能力写成当前交付功能。",
        "提交包同时保留 DOCX、PDF、MPP 和 VSDX 源文件。DOCX 便于继续维护，PDF 便于阅读和打印，VSDX 便于修改设计图，MPP 用于衔接编码测试阶段计划。",
    ]:
        add_paragraph(document, text)
    document.add_page_break()


def add_intro(document: Document) -> None:
    add_heading(document, "1. 引言", level=1)
    sections = [
        ("1.1 编写目的", "本文档用于说明数字图像管理系统在实验3阶段的软件设计结果。文档把实验1的需求构思和实验2的需求分析继续向设计阶段推进，明确系统怎样组织界面、服务、数据访问、数据库对象、外部服务适配和打包运行。后续实验4编码测试阶段可以直接依据本文档拆分任务、编写代码和设计测试用例。"),
        ("1.2 读者对象", "本文档面向课程评审者、小组成员、后续编码测试人员和需要维护系统的开发者。评审者可以通过章节结构和图件判断设计完整性；编码人员可以根据类设计、用例链路和数据设计定位实现边界；测试人员可以根据后续测试映射准备验证证据。"),
        ("1.3 软件系统概述", "数字图像管理系统是一套桌面端图片资产管理工具，主要支持本地目录扫描、缩略图浏览、图片查看、幻灯片播放、批量重命名、轻量编辑、版本历史、标签管理、关键词搜索、可选外部识别标签和 PostgreSQL 元数据管理。系统以本地文件系统保存原始图片，以 PostgreSQL 保存目录、图片、标签、缩略图、版本、搜索和日志等结构化数据。"),
        ("1.4 文档范围", "本文档重点描述软件设计，不展开实验4的完整编码测试报告。文档中的运行截图用于证明界面设计已具备可运行原型，数据库对象用于证明数据设计具有落地依据，VSDX 图源用于证明图件可维护。后续编码日志、运行包和测试记录应在实验4材料中继续补充。"),
        ("1.5 术语说明", "本文档中的 AI 能力指可选的外部兼容端点，用于图片描述、标签识别或自然语言检索辅助；数据库指本机或局域网可访问的 PostgreSQL 实例；缩略图缓存指存入 images.thumbnail 字段的轻量预览数据；版本历史指编辑图片后保存在 image_versions 表和 .versions 目录中的可恢复快照。"),
    ]
    for title, text in sections:
        add_heading(document, title, level=2)
        add_paragraph(document, text)


def add_constraints(document: Document) -> None:
    add_heading(document, "2. 软件设计约束和原则", level=1)
    add_heading(document, "2.1 设计约束", level=2)
    for text in [
        "系统以 JavaFX 桌面应用方式运行，界面资源主要由 FXML 和 CSS 组织。设计必须保证主要界面在常见 Windows 桌面分辨率下可读，按钮、标签、输入框和缩略图卡片不能互相遮挡。",
        "系统以 PostgreSQL 保存结构化数据，但不把原始图片全部写入数据库。数据库负责索引、缩略图、标签、版本、搜索历史和操作日志，磁盘仍然是原图的主要存储位置。",
        "系统要考虑数据库未连接、目录不可访问、图片损坏、外部识别服务不可用和打包后资源路径变化等情况。任何单点故障都不应直接导致用户无法进入修复流程。",
        "外部识别和智能检索属于增强能力。未配置端点时，系统仍应完成本地浏览、缩略图、重命名、编辑、版本和数据库初始化等核心功能。",
        "所有设计表述必须与当前工程事实一致。报告不把系统写成服务器端 Web 平台，不引入未使用的服务器端框架、前端框架、缓存中间件或移动端形态；云端相关表只作为后续扩展预留。",
    ]:
        add_paragraph(document, text)
    add_table(
        document,
        ["约束类别", "约束内容", "设计响应"],
        [
            ["运行环境", "Windows 桌面、JavaFX、PostgreSQL", "提供桌面界面、数据库向导和便携包运行路线。"],
            ["数据一致性", "磁盘文件与数据库路径必须同步", "批量操作由服务层统一处理，必要时回滚或提示。"],
            ["性能体验", "大目录扫描和缩略图生成可能耗时", "采用后台任务、懒加载和缩略图缓存。"],
            ["外部服务", "识别接口可能不可用", "设计为可选能力，失败时保留本地功能。"],
            ["交付维护", "设计图和报告需要可追踪", "保留 VSDX、DOCX、PDF、MPP 和运行截图证据。"],
        ],
    )
    add_heading(document, "2.2 设计原则", level=2)
    for text in [
        "分层清晰。控制器只负责界面状态和事件分派，服务层负责业务规则和事务，DAO 层负责数据库访问，工具类负责无状态辅助能力。任何一层都不应越界承担其他层的主要职责。",
        "先稳定后增强。目录扫描、缩略图浏览、数据库初始化、重命名和编辑版本是基础能力，应优先保证可靠；外部识别和智能检索作为增强能力，在配置可用时提供更好的检索体验。",
        "可恢复优先。批量重命名、编辑保存、版本恢复和数据库初始化都要考虑失败后的状态。设计不追求一次性把所有异常隐藏，而是保证用户知道发生了什么，并尽量保留恢复路径。",
        "证据闭环。每个重要设计点都应能找到对应的图件、类、数据库对象或运行截图。报告不只写“系统支持某功能”，而要说明该功能由哪些构件承担，数据如何变化，后续如何验证。",
    ]:
        add_paragraph(document, text)


def add_design_figures(document: Document) -> None:
    evidence = ROOT / "docs" / "软件工程基础" / "综合性实验最终文档（还是旧的，未更新）" / "图表与证据"
    figures = [
        ("图1 体系结构分层包图", evidence / "图1_体系结构分层包图.png"),
        ("图2 物理部署图", evidence / "图2_物理部署图.png"),
        ("图3 构件图", evidence / "图3_构件图.png"),
        ("图4 用户界面跳转顺序图", evidence / "图4_用户界面跳转顺序图.png"),
        ("图5 缩略图预览用例实现顺序图", evidence / "图5_缩略图预览用例实现顺序图.png"),
        ("图6 批量重命名用例实现顺序图", evidence / "图6_批量重命名用例实现顺序图.png"),
        ("图7 AI 标签扫描与智能搜索用例实现顺序图", evidence / "图7_AI标签扫描与智能搜索用例实现顺序图.png"),
        ("图8 核心设计类图", evidence / "图8_核心设计类图.png"),
        ("图9 数据模型设计图", evidence / "图9_数据模型设计图.png"),
        ("图10 批量重命名事务活动图", evidence / "图10_批量重命名事务活动图.png"),
        ("图11 编码与测试阶段甘特图", evidence / "图11_编码与测试阶段甘特图.png"),
        ("图12 数据库导出表关系图", EXP3 / "图12_数据库导出表关系图.vsdx"),
    ]
    for caption, image in figures:
        if image.suffix.lower() == ".vsdx":
            continue
        add_figure(document, image, caption, 15.5)
    for caption, image in [
        ("图11 目录扫描入库详细活动图", VISIO_EXPORTS / "图11_目录扫描入库详细活动图.png"),
        ("图13 子系统与构件职责展开图", VISIO_EXPORTS / "图13_子系统与构件职责展开图.png"),
        ("图14 AI 标签扫描与智能搜索安全链路图", VISIO_EXPORTS / "图14_AI标签扫描与智能搜索安全链路图.png"),
        ("图15 数据库初始化与离线降级流程图", VISIO_EXPORTS / "图15_数据库初始化与离线降级流程图.png"),
        ("图16 图片编辑与版本历史详细设计图", VISIO_EXPORTS / "图16_图片编辑与版本历史详细设计图.png"),
    ]:
        add_figure(document, image, caption, 15.5)


def add_data_design(document: Document) -> None:
    add_heading(document, "3.6 数据设计", level=2)
    for text in [
        "数据设计围绕图片资产管理的生命周期展开。目录和图片是基础对象，标签和 AI 分析结果是检索对象，版本和编辑操作是恢复对象，搜索历史和操作日志是审计对象，设置表是启动和配置对象。数据库不保存所有原图内容，而保存索引、缩略图、路径、标签和历史记录。",
        "schema.sql 当前包含 13 张基础表、4 个视图、19 个索引、5 个触发器、4 个日志函数和 5 个存储过程。数据库对象数量不是堆砌结果，而是围绕目录浏览、搜索、标签、版本、报表和操作追踪逐步展开。",
        "cloud_sources 和 cloud_images 作为后续扩展预留，当前核心交付仍以本地图片目录、PostgreSQL 元数据和 JavaFX 客户端为主。报告将扩展对象与当前运行链路区分开，避免把预留设计写成已完成能力。",
    ]:
        add_paragraph(document, text)
    add_table(
        document,
        ["数据对象", "主要字段或关系", "支撑功能"],
        [
            ["directories", "dir_path、parent_id、层级关系", "目录树、扫描路径、当前目录范围搜索。"],
            ["images", "file_name、file_path、size、width、height、thumbnail、ai_processed", "缩略图浏览、查看器、重命名、搜索结果展示。"],
            ["tags / image_tags", "标签分类、标签名称、图片关联", "人工标签、识别标签、关键词搜索。"],
            ["ai_analysis_results", "description、confidence、analyzed_at", "图片描述、语义检索、识别结果复查。"],
            ["image_versions", "version_num、file_path、edit_type、is_current", "编辑版本历史、恢复旧版本。"],
            ["operation_logs", "操作类型、对象、时间", "扫描、重命名、删除和标签变更追踪。"],
            ["search_history", "query_text、search_mode、generated_sql、result_count", "搜索行为记录和后续分析。"],
            ["app_settings", "key、value", "扫描目录、界面偏好和运行配置保存。"],
        ],
    )
    add_heading(document, "3.6.1 数据库对象清单", level=3)
    add_table(
        document,
        ["对象类别", "数量", "代表对象", "设计作用"],
        [
            ["基础表", "13", "directories、images、tags、image_tags、ai_analysis_results、image_versions", "保存核心业务数据和扩展元数据。"],
            ["视图", "4", "v_active_images、v_directory_stats、v_image_search、v_tag_stats", "统一查询入口，支撑统计和搜索。"],
            ["索引", "19", "idx_images_active、idx_tags_name_trgm、idx_ai_desc_trgm", "提升目录浏览、标签检索和描述检索速度。"],
            ["触发器", "5", "trg_image_after_insert、trg_tag_after_insert 等", "记录图片和标签变更。"],
            ["函数", "4", "fn_log_image_insert、fn_log_tag_change 等", "为触发器写入操作日志。"],
            ["存储过程", "5", "sp_monthly_report、sp_batch_rename、sp_restore_version 等", "提供报表和批处理设计扩展点。"],
        ],
    )
    add_heading(document, "3.6.2 数据一致性设计", level=3)
    for text in [
        "图片路径、文件名和目录 ID 是连接磁盘与数据库的关键。单张重命名、批量重命名、删除和编辑保存都必须由服务层统一处理，不能让界面直接修改数据库记录，也不能只修改磁盘文件。",
        "标签和 AI 分析结果采用可复用结构。人工标签与识别标签最终都进入 tags 和 image_tags，描述文本进入 ai_analysis_results。这样普通关键词检索和智能检索都可以复用同一组持久化结果。",
        "版本历史采用独立版本文件和数据库记录共同保存。image_versions 记录版本号、路径、尺寸、缩略图、编辑类型和当前标记，.versions 目录保存具体文件。恢复版本时需要同时更新文件、images 表和当前版本标记。",
    ]:
        add_paragraph(document, text)


def add_deployment_plan(document: Document) -> None:
    add_heading(document, "3.7 部署设计", level=2)
    for text in [
        "系统部署为本机桌面应用。JavaFX 客户端运行在用户 Windows 环境中，PostgreSQL 可以运行在本机或可访问的局域网主机上，原始图片保存在本地文件系统，外部识别端点通过 HTTPS 作为可选增强服务接入。",
        "部署设计的核心不是复杂的服务器拓扑，而是保证交付包在普通机器上有明确启动和修复路径。若数据库不可用，用户可以通过数据库初始化向导配置连接并执行 schema；若外部识别端点不可用，系统仍可进行本地图片管理。",
        "打包产物分为普通 JAR 和 Windows 便携包。JAR 面向已有 Java 环境，便携包面向未配置 Java 的机器。两种产物都应进入同一套欢迎界面、主界面和数据库向导流程。",
    ]:
        add_paragraph(document, text)
    add_heading(document, "3.8 编码和测试阶段项目计划", level=2)
    for text in [
        "编码测试阶段计划按风险优先排序：先完成启动配置和数据库初始化，再完成目录扫描和缩略图缓存，随后完成浏览、搜索、批量重命名、编辑版本、标签和 AI 增强，最后完成打包交付与验证。",
        "项目计划文件以 MPP 形式保留，便于继续维护。报告中的后续测试映射与项目计划互相对应，能够让实验3的设计成果自然衔接到实验4的编码测试报告。",
    ]:
        add_paragraph(document, text)


def add_traceability_depth(document: Document) -> None:
    add_heading(document, "14. 需求到设计追踪深化", level=1)
    for text in [
        "软件设计不能脱离需求分析单独存在。本系统从实验1的需求构思、实验2的需求分析，到实验3的软件设计，形成了逐步细化的链路。需求阶段提出的目录管理、图片浏览、批量重命名、标签检索、编辑和数据库管理，在设计阶段分别落到界面、服务、DAO、数据库表和测试映射中。",
        "需求到设计的第一类追踪是功能追踪。目录管理需求对应 MainController 的目录树、ScanTask 的后台扫描、directories 表和目录统计视图；图片浏览需求对应缩略图卡片、ImageViewerController、SlideshowController 和 images 表；重命名需求对应 RenameDialogController、ImageServiceImpl.batchRename() 和操作日志。",
        "需求到设计的第二类追踪是数据追踪。需求中提到的图片名称、路径、格式、大小、尺寸、标签、描述、编辑历史和搜索记录，在数据设计中分别进入 images、tags、image_tags、ai_analysis_results、image_versions、image_edit_operations 和 search_history。这样能够保证每个界面展示字段都有数据库或文件系统来源。",
        "需求到设计的第三类追踪是质量追踪。需求阶段关注易用性、稳定性和可维护性，设计阶段进一步落实为数据库初始化向导、后台扫描任务、缩略图缓存、事务边界、版本历史、错误提示和打包运行路线。质量要求不再停留在口号，而是转化为可检查的构件和测试点。",
        "需求到设计的第四类追踪是异常追踪。需求阶段只会描述用户希望完成什么任务，设计阶段必须补充任务失败时怎么办。例如目录不可访问时提示用户重新选择，数据库不可用时打开初始化向导，外部识别失败时保留本地管理功能，批量重命名失败时避免数据库与磁盘状态分裂。",
    ]:
        add_paragraph(document, text)

    add_table(
        document,
        ["需求来源", "设计落点", "核心类或对象", "后续验证方式"],
        [
            ["目录扫描与管理", "目录树、后台扫描、目录表", "MainController、ScanTask、directories", "扫描样例目录并核对数据库目录层级。"],
            ["图片缩略图浏览", "缩略图卡片、缓存字段", "ImageServiceImpl、ImageUtil、images.thumbnail", "打开大目录并观察首屏响应和缓存写入。"],
            ["批量图片重命名", "重命名对话框、事务活动图", "RenameDialogController、batchRename()", "检查磁盘文件名、数据库路径和日志一致。"],
            ["标签与智能检索", "标签表、搜索服务、安全链路", "SearchService、TagDaoImpl、ai_analysis_results", "测试关键词、标签和自然语言检索。"],
            ["图片编辑与恢复", "编辑器、版本表、版本文件", "ImageEditorController、EditService、image_versions", "保存多版本并恢复旧版本。"],
            ["部署与初始化", "数据库向导、部署图", "DatabaseSetupDialog、DatabaseBootstrapService", "清空配置后重新初始化数据库。"],
        ],
    )

    evidence = ROOT / "docs" / "软件工程基础" / "综合性实验最终文档（还是旧的，未更新）" / "图表与证据"
    for caption, image in [
        ("图14-1 需求构思用例图追踪证据", evidence / "图1_需求构思用例图.png"),
        ("图14-2 需求分析模型追踪证据一", evidence / "图2-1_需求分析模型.png"),
        ("图14-3 需求分析模型追踪证据二", evidence / "图2-2_需求分析模型.png"),
        ("图14-4 需求分析模型追踪证据三", evidence / "图2-3_需求分析模型.png"),
        ("图14-5 需求分析模型追踪证据四", evidence / "图2-4_需求分析模型.png"),
    ]:
        add_figure(document, image, caption, 14.5)

    add_heading(document, "15. 数据字典摘要深化", level=1)
    for text in [
        "数据字典是数据设计可检查性的核心。仅列出表名无法说明系统是否真的考虑了字段、约束和业务语义，因此本节进一步用文字方式概括主要数据表的设计意图。正式提交中的完整字段清单以 schema.sql 和报告数据设计表为准，本节用于帮助阅读者理解字段背后的业务作用。",
        "directories 表的关键在于路径唯一性和父子关系。dir_path 表示真实磁盘路径，parent_id 表示目录层级，目录统计视图可以在此基础上计算每个目录下的图片数量。设计上不把目录树只当作界面控件，而是把它作为数据库中可追踪的数据结构。",
        "images 表是系统最核心的业务表。file_name、file_path 和 extension 支撑文件定位；file_size、width、height、created_at、modified_at 支撑元数据展示；thumbnail 支撑快速预览；ai_processed 支撑识别任务状态；is_deleted 支撑逻辑删除和恢复判断。该表连接了磁盘文件、缩略图界面、搜索结果和日志追踪。",
        "tags、tag_categories 和 image_tags 共同形成标签模型。tag_categories 让标签具有分类，tags 保存标准化标签实体，image_tags 保存图片与标签之间的多对多关系。这样人工标签和识别标签可以进入同一检索结构，避免每类标签单独建一套查询逻辑。",
        "ai_analysis_results 保存外部识别返回的描述、标签摘要、置信度和时间。它不是替代 tags 表，而是补充语义描述和识别结果追踪。普通关键词搜索可以匹配描述文本，智能搜索也可以利用该描述构造更贴近用户表达的查询。",
        "image_versions 和 image_edit_operations 用于保存编辑历史。image_versions 记录每个版本的文件路径、尺寸、缩略图、编辑类型和当前标记；image_edit_operations 可记录具体编辑参数。两张表配合 .versions 目录，使图片编辑具备可恢复性。",
        "operation_logs 和 search_history 用于后续验证和审计。operation_logs 记录图片和标签变更，search_history 记录搜索文本、模式、查询和结果数量。它们不直接决定功能是否可用，但能为测试报告提供证据，也能帮助维护者定位用户操作过程。",
        "app_settings 保存扫描目录、界面偏好和部分运行配置。设置读取应提供默认值，避免缺失配置导致启动失败。涉及敏感信息的配置应避免在正式文档和截图中暴露，报告只说明设计结构，不呈现真实密钥。",
    ]:
        add_paragraph(document, text)
    add_table(
        document,
        ["表名", "关键字段", "设计重点"],
        [
            ["directories", "dir_path、parent_id、created_at", "目录路径唯一，支持层级和统计。"],
            ["images", "file_path、thumbnail、ai_processed、is_deleted", "连接磁盘文件、缩略图、识别状态和逻辑删除。"],
            ["tags", "name、category_id", "标签实体去重，按分类管理。"],
            ["image_tags", "image_id、tag_id、confidence", "图片与标签多对多关联，兼容人工和识别标签。"],
            ["ai_analysis_results", "description、raw_response、confidence", "保存识别描述和结果追踪。"],
            ["image_versions", "version_num、file_path、is_current", "支持编辑版本历史和恢复。"],
            ["operation_logs", "operation_type、target_id、operated_at", "记录图片和标签变更。"],
            ["search_history", "query_text、search_mode、result_count", "记录搜索行为与结果规模。"],
        ],
    )

    add_heading(document, "16. 图件维护说明", level=1)
    for text in [
        "实验3提交包保留 VSDX 图源，是为了让设计具有可维护性。若后续代码结构、数据库表或用例链路发生变化，应优先修改对应 VSDX，再更新 DOCX 中的图像和说明。只在 PDF 中保存静态图片会降低后续维护效率。",
        "图件维护应遵循图号稳定原则。体系结构、部署、构件、界面跳转、顺序图、类图、数据模型、活动图和补充详细设计图已经形成编号。后续新增图件可以追加编号，不应随意重排已有图号，以免正文引用和附件文件名失配。",
        "图件内容应避免过度装饰，重点表达职责、消息、数据和异常路径。对软件工程实验而言，清楚的边界、方向、类名、表名和流程节点比复杂视觉效果更重要。每张图都应能回答一个明确设计问题。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "17. 评审问答准备", level=1)
    for text in [
        "如果评审者询问为什么采用桌面端架构，可以说明本系统的主要操作对象是本地图片目录，用户需要直接扫描磁盘、读取缩略图、重命名文件、编辑图片和维护本机数据库。桌面端能够直接访问文件系统和本机图形界面，部署路径也更符合课程项目的实际交付要求。",
        "如果评审者询问为什么使用 PostgreSQL，可以说明系统不仅保存文件路径，还需要管理标签、描述、搜索历史、版本历史、操作日志、视图、索引和触发器。PostgreSQL 能提供较完整的关系建模能力和查询能力，适合把图片管理从简单文件浏览提升为可检索、可追踪、可恢复的数据管理系统。",
        "如果评审者询问外部识别服务是否影响基础功能，可以说明它被设计为可选增强能力。未配置端点时，系统仍能完成本地目录扫描、缩略图浏览、查看、播放、重命名、编辑、版本恢复和数据库初始化；配置端点后，系统再进一步提供标签识别和智能搜索。",
        "如果评审者询问批量重命名为什么需要重点设计，可以说明该功能同时修改磁盘文件和数据库记录，是最容易产生不一致的操作之一。报告通过顺序图、事务活动图和方法级说明，把参数校验、冲突检测、文件修改、数据库更新、异常处理和界面刷新都列入设计范围。",
        "如果评审者询问编辑版本为什么不直接覆盖原图，可以说明图片编辑具有试错和回退需求。版本历史设计让每次编辑都有文件快照和数据库记录，用户可以恢复旧版本，测试人员也可以通过 image_versions 表和 .versions 目录验证编辑链路是否完整。",
        "如果评审者询问报告中图件与代码是否对应，可以说明图件中的 Controller、Service、DAO、Model、Scanner 和 AI 适配对象都能在当前工程包结构中找到对应类；数据模型图和数据库表关系可以对应 schema.sql；界面设计则通过当前 JavaFX 运行截图进行验证。",
        "如果评审者询问后续如何测试，可以说明报告已经把设计转化为启动初始化、目录扫描、缩略图缓存、搜索标签、批量重命名、编辑版本、数据库对象、打包运行和交付一致性等测试映射。实验4阶段只需要沿这些映射补充真实运行日志、数据库查询和截图证据。",
        "如果评审者询问提交包是否完整，可以说明正式 ZIP 根目录包含正式 DOCX、正式 PDF、项目计划 MPP 和 16 个 VSDX 图源文件，没有空的嵌套目录。DOCX 用于继续编辑，PDF 用于阅读，VSDX 用于图件维护，MPP 用于编码测试阶段计划。",
        "最终回答时应始终围绕系统自身设计、运行证据和交付材料展开，不需要引入与本系统无关的外部比较。",
    ]:
        add_paragraph(document, text)


def build_clean_report() -> None:
    document = Document()
    setup_styles(document)
    add_title_page(document)
    add_intro(document)
    add_constraints(document)

    add_heading(document, "3. 软件设计方案", level=1)
    add_heading(document, "3.1 体系结构设计", level=2)
    build_architecture_detail(document)
    add_design_figures(document)

    add_heading(document, "3.2 用户界面设计", level=2)
    build_ui_detail(document)

    add_heading(document, "3.3 用例设计", level=2)
    build_use_case_detail(document)

    add_heading(document, "3.4 子系统与构件设计", level=2)
    build_component_detail(document)

    add_heading(document, "3.5 类设计", level=2)
    build_class_detail(document)

    add_data_design(document)
    add_deployment_plan(document)

    add_heading(document, "4. 实施指南与评审记录", level=1)
    for text in [
        "编码时应保持分层依赖方向。控制器不直接拼接复杂 SQL，DAO 不处理界面提示，服务层负责业务事务和异常恢复，工具类保持无状态。涉及数据库连接、外部端点和本机路径的配置不写死在源码中。",
        "评审时应重点检查图件、正文、代码结构和数据库对象是否一致。若报告中出现某个设计对象，工程中应能找到对应类、表、截图或计划项；若工程中存在核心功能，报告也应有相应说明。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "5. 设计落地证据与质量复核", level=1)
    for heading, text in [
        ("5.1 设计模型与实现构件的对应关系", "系统按界面层、控制层、服务层、数据访问层和数据库层展开。界面层负责接收操作和展示状态，服务层承担扫描、缩略图、重命名、编辑、搜索、AI 标签和初始化流程，DAO 层封装 PostgreSQL 访问。"),
        ("5.2 界面设计落地证据", "运行截图覆盖欢迎、主界面、查看器、幻灯片、编辑器、设置、重命名和数据库向导，证明界面设计已经落到可运行原型。"),
        ("5.3 核心用例运行链路补充", "目录扫描、AI 标签、批量重命名、编辑版本和数据库初始化是最需要验证的五条链路。报告后续章节分别展开这些链路的触发、构件、数据变化和异常处理。"),
        ("5.4 质量复核与后续编码测试关注点", "质量复核关注界面响应、数据一致性、异常恢复、可维护性和交付同步。后续编码测试应保留截图、日志、数据库查询和打包文件作为证据。"),
    ]:
        add_heading(document, heading, level=2)
        add_paragraph(document, text)
    build_evidence_detail(document)

    add_heading(document, "6. 重点设计深化说明", level=1)
    for text in [
        "本章对前文已经覆盖的界面、用例、构件、类、风险和阶段衔接进行集中说明。它的作用是把设计从“有图有表”推进到“能指导实现和验收”。",
        "用户界面与业务场景要保持对应。主界面服务目录浏览和批量管理，查看器服务单图浏览，编辑器服务版本化编辑，设置界面服务配置维护，数据库向导服务部署修复。",
        "用例设计要覆盖成功路径和异常路径。扫描要处理坏图，重命名要处理冲突，智能搜索要处理安全拒绝，编辑要处理版本恢复，初始化要处理连接失败。",
        "类设计要保持职责清晰。MainController 负责协调，ImageServiceImpl 负责图片业务，SearchService 负责检索，EditService 负责版本，DatabaseBootstrapService 负责初始化，DAO 负责数据库。",
        "风险验收要具体。不能只看界面是否打开，还要检查数据库记录、磁盘文件、缩略图缓存、版本文件、日志记录和打包产物。",
    ]:
        add_paragraph(document, text)

    add_heading(document, "7. 补充详细设计图件", level=1)
    for text in [
        "本章集中说明新增的补充详细设计图件。目录扫描入库详细活动图用于补足大目录扫描、异常文件跳过和数据库同步过程；子系统与构件职责展开图用于补足 Controller、Service、DAO、Scanner 和外部服务适配之间的职责边界；AI 标签扫描与智能搜索安全链路图用于补足外部端点、标签入库、查询安全和结果返回之间的关系。",
        "数据库初始化与离线降级流程图用于说明新环境运行时如何从连接失败进入可修复流程，图片编辑与版本历史详细设计图用于说明编辑保存、缩略图刷新、版本记录和恢复旧版本之间的关系。这五张图均保留为 VSDX 图源，并已在正式提交包中与主文档、PDF 和项目计划一起打包。",
        "补充图件的作用是把关键风险链路画清楚。基础包图、构件图、顺序图和类图说明系统总体结构，补充图件进一步说明后台任务、外部服务、初始化恢复和版本历史这些容易影响实现质量的细节。后续实验4测试时，可以直接按这些图件检查代码和运行证据。",
    ]:
        add_paragraph(document, text)

    build_quality_appendix(document)
    build_deep_design_chapter(document)
    build_requirement_review_chapter(document)
    build_test_mapping_chapter(document)
    build_submission_chapter(document)
    build_final_sync_chapter(document)
    add_traceability_depth(document)
    replace_text(document)

    document.save(str(FINAL_DOCX))
    document.save(str(DOCX))
    stats = inspect_docx(FINAL_DOCX)
    print("saved", FINAL_DOCX)
    print(stats)


if __name__ == "__main__":
    build_clean_report()
