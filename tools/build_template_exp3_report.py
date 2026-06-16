from __future__ import annotations

import shutil
import sys
import re
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import fitz
import win32com.client
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
EXP3 = ROOT / "docs" / "软件工程基础" / "实验3"
MATERIALS = ROOT / "docs" / "软件工程基础" / "2026版《软件工程基础》实验材料"
TEMPLATE_DOC = MATERIALS / "3XX系统软件设计规格说明书.doc"
REVIEW_DIR = ROOT / "target" / "exp3-template-review"
TEMPLATE_DOCX = REVIEW_DIR / "3XX系统软件设计规格说明书.docx"
FINAL_BASENAME = "第RR10组+徐阳+202425220527_数字图像管理系统软件设计规格说明书"
FINAL_DOCX = EXP3 / f"{FINAL_BASENAME}.docx"
FINAL_PDF = EXP3 / f"{FINAL_BASENAME}.pdf"
FINAL_ZIP = EXP3 / "第RR10组+徐阳+202425220527.zip"
LEGACY_DOCX = EXP3 / "数字图像管理系统软件设计规格说明书v1.2_审计修订版.docx"
LEGACY_PDF = EXP3 / "数字图像管理系统软件设计规格说明书v1.2_审计修订版.pdf"
MPP = EXP3 / "数字图像管理系统编码测试阶段项目计划.mpp"
COLLECTED = (
    ROOT
    / "docs"
    / "软件工程基础"
    / "收齐别的同学的作业"
    / "实验3"
    / "第RR10组+徐阳+202425220527"
)
UI_SMOKE = ROOT / "target" / "ui-smoke"
VISIO_EXPORTS = ROOT / "C_tmp_visio_exports"
EVIDENCE = ROOT / "docs" / "软件工程基础" / "综合性实验最终文档（还是旧的，未更新）" / "图表与证据"

TEAM = [
    ("组长", "徐阳", "202425220527", "软工5班", "设计文档统筹、数据库与AI设计、最终整理。"),
    ("成员2", "毕振岚", "202425220501", "软工5班", "主界面、目录树、初始化与部署设计。"),
    ("成员3", "陈厚华", "202425220502", "软工5班", "缩略图、图片服务、DAO与图件复核。"),
    ("成员4", "陈智杰", "202425220506", "软工5班", "批量重命名、右键菜单、事务设计。"),
    ("成员5", "殷浚峰", "202425220528", "软工5班", "查看器、幻灯片、编辑版本与项目计划。"),
]

FIGURE_CREDITS = {
    "图1": ("徐阳", "毕振岚"),
    "图2": ("徐阳", "陈厚华"),
    "图3": ("毕振岚", "徐阳"),
    "图4": ("陈厚华", "陈智杰"),
    "图5": ("陈厚华", "殷浚峰"),
    "图6": ("陈智杰", "徐阳"),
    "图7": ("徐阳", "陈厚华"),
    "图8": ("徐阳", "毕振岚"),
    "图9": ("徐阳", "陈厚华"),
    "图10": ("陈智杰", "徐阳"),
    "图11": ("徐阳", "毕振岚"),
    "图12": ("徐阳", "陈厚华"),
    "图13": ("毕振岚", "徐阳"),
    "图14": ("徐阳", "陈智杰"),
    "图15": ("毕振岚", "徐阳"),
    "图16": ("殷浚峰", "陈厚华"),
    "图17": ("殷浚峰", "徐阳"),
}


def set_east_asia(run, name: str = "宋体") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_paragraph_text(paragraph, text: str, font_size: float | None = None, font_name: str = "宋体") -> None:
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text
    set_east_asia(run, font_name)
    if font_size:
        run.font.size = Pt(font_size)


def add_body(document: Document, text: str):
    paragraph = document.add_paragraph(style="正文2" if "正文2" in [s.name for s in document.styles] else None)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_east_asia(run)
    run.font.size = Pt(10.5)
    return paragraph


def add_heading(document: Document, text: str, level: int = 2):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        set_east_asia(run)
        run.font.bold = True
    return paragraph


def add_caption(document: Document, text: str, bold: bool = False):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.1
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_east_asia(run)
    run.font.size = Pt(9.5)
    run.font.bold = bold
    return paragraph


def credit_for(caption: str) -> tuple[str, str]:
    for prefix, credit in FIGURE_CREDITS.items():
        if caption.startswith(prefix):
            return credit
    return ("徐阳", "毕振岚")


def add_figure(document: Document, image: Path, caption: str, width_cm: float = 15.0):
    if not image.exists():
        add_body(document, f"{caption} 对应图片文件暂未生成，正式提交时以同名 VSDX 图源为准。")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.add_run().add_picture(str(image), width=Cm(width_cm))
    add_caption(document, caption, bold=False)
    drawer, reviewer = credit_for(caption)
    add_caption(document, f"绘制者：{drawer}     审查者：{reviewer}", bold=False)


def add_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_east_asia(run)
                    run.font.size = Pt(9.5)
    return table


def ensure_template_docx() -> None:
    REVIEW_DIR.mkdir(parents=True, exist_ok=True)
    if TEMPLATE_DOCX.exists() and TEMPLATE_DOCX.stat().st_mtime >= TEMPLATE_DOC.stat().st_mtime:
        return
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(TEMPLATE_DOC.resolve()))
        doc.SaveAs(str(TEMPLATE_DOCX.resolve()), FileFormat=16)
        doc.Close(False)
    finally:
        word.Quit()


def ensure_figure12_export() -> None:
    png = VISIO_EXPORTS / "图12_数据库导出表关系图.png"
    vsdx = EXP3 / "图12_数据库导出表关系图.vsdx"
    if png.exists() or not vsdx.exists():
        return
    app = win32com.client.DispatchEx("Visio.Application")
    app.Visible = False
    try:
        doc = app.Documents.Open(str(vsdx.resolve()))
        doc.Pages.Item(1).Export(str(png.resolve()))
        doc.Close()
    finally:
        app.Quit()


def remove_template_body(document: Document) -> None:
    start = None
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "1、引言":
            start = paragraph._p
            break
    if start is None:
        raise RuntimeError("未找到模板正文起始标题：1、引言")
    body = document.element.body
    deleting = False
    for child in list(body):
        if child is start:
            deleting = True
        if deleting and child.tag != qn("w:sectPr"):
            body.remove(child)


def fill_cover_and_history(document: Document) -> None:
    set_paragraph_text(document.paragraphs[0], "文档编号：数字图像管理系统 – SRS – 1.0", 11, "楷体")
    set_paragraph_text(document.paragraphs[11], "数字图像管理系统", 24, "楷体")
    set_paragraph_text(document.paragraphs[12], "软件设计规格说明书", 24, "黑体")
    set_paragraph_text(document.paragraphs[26], "日期：2026年6月15日", 11, "宋体")

    cover = document.tables[0]
    cover.rows[0].cells[1].text = "RR10"
    for idx, member in enumerate(TEAM, start=2):
        cells = cover.rows[idx].cells
        for cell_idx, value in enumerate(member):
            cells[cell_idx].text = value
    for row in cover.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_east_asia(run)
                    run.font.size = Pt(9)

    history = document.tables[1]
    entries = [
        ("1", "2026/6/1", "徐阳", "建立软件设计规格说明书初稿，确定体系结构、部署、数据设计和项目计划框架。", "V1.0"),
        ("2", "2026/6/6", "毕振岚、陈厚华", "补充用户界面设计、目录扫描、缩略图浏览和数据库初始化链路，复核图件表达。", "V1.1"),
        ("3", "2026/6/10", "陈智杰、殷浚峰", "补充批量重命名、图片查看、幻灯片、编辑版本和异常恢复设计。", "V1.2"),
        ("4", "2026/6/15", "徐阳", "基于课程模板完成格式统一、图件绘制者与审查者标注、PDF和提交目录同步。", "V1.3"),
    ]
    for row_idx, entry in enumerate(entries, start=1):
        cells = history.rows[row_idx].cells
        for col_idx, value in enumerate(entry):
            cells[col_idx].text = value


def patch_shared_builders() -> None:
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    import build_clean_exp3_report as clean
    import enhance_exp3_design_doc as enhance

    for module in (clean, enhance):
        module.add_paragraph = add_body
        module.add_heading = add_heading
        module.add_table = add_table
        module.add_figure = add_figure
        module.add_caption = add_caption
    return clean, enhance


def add_intro(document: Document) -> None:
    add_heading(document, "1、引言", 1)
    sections = [
        ("1.1 编写目的", "本文档用于说明数字图像管理系统在软件设计阶段的设计结果。文档承接实验1的需求构思和实验2的软件需求规格说明，将目录扫描、缩略图浏览、图片查看、批量重命名、标签检索、AI辅助识别、图片编辑、版本历史和数据库初始化等需求进一步落实到体系结构、界面、用例、类、数据和部署方案中。后续编码与测试阶段可以据此拆分任务、检查实现边界并准备验证证据。"),
        ("1.2 读者对象", "本文档面向课程评审者、小组成员、后续编码人员、测试人员和维护人员。评审者可以从模板结构、图件、表格和运行截图判断设计是否完整；编码人员可以根据控制器、服务、DAO、模型和数据库对象定位实现职责；测试人员可以根据用例链路、数据设计和项目计划准备实验4的测试用例。"),
        ("1.3 软件项目概述", "数字图像管理系统是一套面向个人用户的桌面端图片资产管理工具。系统采用 JavaFX 构建本地界面，以 PostgreSQL 保存目录、图片、标签、AI分析结果、版本历史、搜索历史和操作日志等元数据，原始图片仍保存在本地文件系统中。系统核心功能包括目录树导航、缩略图预览、图片查看、幻灯片播放、批量重命名、轻量编辑、版本恢复、关键词搜索、标签管理和可选的外部AI识别辅助。"),
        ("1.4 文档概述", "本文档按照老师给出的软件设计规格说明书模板组织。第1章说明编写目的、读者对象、项目概况、术语和参考资料；第2章说明软件设计目标、原则、约束和限制；第3章给出软件体系结构设计、用户界面设计、用例设计、子系统与构件设计、类设计、数据设计、部署设计以及编码测试阶段项目计划。文档中的图件均保留 VSDX 源文件，关键界面使用当前工程运行截图作为设计落地证据。"),
        ("1.5 定义", "本文档中的“AI能力”指可选外部兼容端点提供的图片描述、标签识别和自然语言检索辅助能力；“缩略图缓存”指系统将图片预览数据保存到 images.thumbnail 字段以提升再次浏览速度；“版本历史”指图片编辑后由 image_versions 表和 .versions 目录共同保存的可恢复版本；“数据库初始化”指通过程序界面完成 PostgreSQL 连接检测、schema 建立和默认数据导入的流程。"),
        ("1.6 参考资料", "本文档参考《软件工程基础》实验说明、课程教材中关于软件设计、包图、部署图、顺序图、类设计、数据设计和构件设计的章节，以及本项目实验1、实验2的阶段文档、当前 JavaFX 工程源码、schema.sql、运行截图和 VSDX 设计图源。"),
    ]
    for title, text in sections:
        add_heading(document, title, 2)
        add_body(document, text)


def add_constraints(document: Document) -> None:
    add_heading(document, "2、软件设计约束", 1)
    add_heading(document, "2.1 软件设计目标和原则", 2)
    for text in [
        "系统设计目标是把数字图像管理需求转化为可实现、可测试、可维护的软件结构。设计结果需要说明用户界面如何组织，业务逻辑由哪些服务承担，数据如何进入 PostgreSQL，外部AI服务如何作为增强能力接入，异常情况下用户如何得到明确反馈，后续编码测试阶段如何验证这些设计。",
        "系统坚持分层设计原则。界面控制器负责接收用户动作、维护控件状态和展示反馈；服务层负责业务规则、事务边界和异常恢复；DAO层负责数据库读写；工具类负责无状态的文件和图像处理；模型类负责在各层之间传递清晰的数据结构。",
        "系统坚持可恢复设计原则。目录扫描、批量重命名、图片编辑、版本恢复、数据库初始化和AI任务都可能失败，设计不能只描述成功路径，还要说明冲突、坏图、权限不足、数据库不可用、外部端点失败和打包运行路径变化时的处理方式。",
        "系统坚持证据闭环原则。报告中的主要设计点需要能够对应到 VSDX 图源、当前源码包结构、数据库对象、运行截图或项目计划。文档不把未落地的技术栈写成当前完成内容，而是明确区分当前核心交付与后续可扩展方向。",
    ]:
        add_body(document, text)
    add_table(
        document,
        ["原则", "在本系统中的落实", "对应设计证据"],
        [
            ["模块化", "界面、服务、DAO、模型、工具、AI适配和数据库脚本分层组织。", "体系结构包图、构件图、类设计章节。"],
            ["信息隐藏", "控制器不直接拼复杂SQL，DAO不处理界面提示，AI服务细节封装在适配类中。", "构件职责说明、关键类说明。"],
            ["可追踪", "每个核心需求对应界面、服务、数据对象和后续测试点。", "用例设计、数据设计、测试映射。"],
            ["可恢复", "重命名、编辑、初始化、AI任务均设计异常反馈和恢复路径。", "活动图、版本历史设计、初始化流程图。"],
            ["渐进增强", "外部AI属于增强能力，未配置时不影响本地图片管理。", "AI安全链路、设置界面、部署设计。"],
        ],
    )
    add_heading(document, "2.2 软件设计的约束和限制", 2)
    for text in [
        "运行形态约束。系统当前是 JavaFX 桌面客户端，主要运行于 Windows 桌面环境。界面设计需要适配常见窗口尺寸，目录树、搜索栏、缩略图卡片、状态栏和对话框不能互相遮挡。报告不把系统描述为 Web 前后端分离平台，也不把未使用的移动端形态写入当前交付。",
        "数据存储约束。PostgreSQL 用于保存结构化元数据、缩略图、标签、版本、搜索历史和操作日志，原始图片仍由本地文件系统保存。设计必须处理磁盘文件和数据库记录之间的一致性，尤其是重命名、复制、删除和编辑保存等操作。",
        "外部服务约束。外部AI端点可能未配置、超时、返回异常或额度不足，因此 AI 标签扫描和智能搜索必须设计为可选增强能力。未配置端点时，本地浏览、缩略图、查看、重命名、编辑、版本恢复和数据库初始化仍应可用。",
        "交付约束。实验3提交材料需要包含 DOCX、PDF、VSDX 图源和 MPP 项目计划。小组作业目录应保持扁平清晰，避免空文件夹和不必要的多层嵌套，便于班级最终统一压缩提交。",
    ]:
        add_body(document, text)


def add_design(document: Document) -> None:
    clean, enhance = patch_shared_builders()
    add_heading(document, "3、软件设计", 1)

    add_heading(document, "3.1 软件体系结构设计", 2)
    enhance.build_architecture_detail(document)
    add_design_figures(document)

    add_heading(document, "3.2 用户界面设计", 2)
    enhance.build_ui_detail(document)

    add_heading(document, "3.3 用例设计", 2)
    enhance.build_use_case_detail(document)

    add_heading(document, "3.4 子系统与构件设计", 2)
    original_add_figure = enhance.add_figure
    enhance.add_figure = lambda *args, **kwargs: None
    try:
        enhance.build_component_detail(document)
    finally:
        enhance.add_figure = original_add_figure

    add_heading(document, "3.5 类设计", 2)
    enhance.build_class_detail(document)

    add_data_design(document)
    add_deployment_and_plan(document)
    add_review_mapping(document)


def add_data_design(document: Document) -> None:
    add_heading(document, "3.6 数据设计", 2)
    for text in [
        "数据设计围绕图片资产管理生命周期展开。目录和图片是基础对象，标签和 AI 分析结果是检索对象，版本和编辑操作是恢复对象，搜索历史和操作日志是审计对象，应用设置是启动和配置对象。数据库不保存全部原图内容，而保存索引、缩略图、路径、标签、历史记录和运行状态。",
        "当前 schema.sql 包含目录、图片、标签、标签分类、图片标签关联、AI分析结果、版本历史、编辑操作、搜索历史、操作日志、应用设置和后续扩展预留对象。设计时将 cloud_sources、cloud_images 等对象明确视为后续扩展预留，当前核心交付仍以本地文件系统、PostgreSQL 元数据和 JavaFX 客户端为主。",
        "数据库设计要同时服务功能和测试。目录浏览需要 directories 与 images，关键词搜索需要 tags、image_tags 和 ai_analysis_results，版本恢复需要 image_versions 与版本文件，操作追踪需要 operation_logs，后续实验4可以通过这些表和视图核对运行结果。",
    ]:
        add_body(document, text)
    add_table(
        document,
        ["数据对象", "主要字段或关系", "支撑功能"],
        [
            ["directories", "dir_path、parent_id、层级关系", "目录树、扫描路径、当前目录范围搜索。"],
            ["images", "file_name、file_path、size、width、height、thumbnail、ai_processed", "缩略图浏览、查看器、重命名、搜索结果展示。"],
            ["tags / image_tags", "标签分类、标签名称、图片关联", "人工标签、识别标签、关键词搜索。"],
            ["ai_analysis_results", "description、confidence、analyzed_at", "图片描述、语义检索、识别结果复查。"],
            ["image_versions", "version_num、file_path、edit_type、is_current", "编辑版本历史、恢复旧版本。"],
            ["operation_logs", "操作类型、对象、时间", "扫描、重命名、删除和标签变更追踪。"],
            ["search_history", "query_text、search_mode、generated_sql、result_count", "搜索行为记录和后续分析。"],
            ["app_settings", "key、value", "扫描目录、界面偏好和运行配置保存。"],
        ],
    )
    add_heading(document, "3.6.1 数据库对象清单", 3)
    add_table(
        document,
        ["对象类别", "数量", "代表对象", "设计作用"],
        [
            ["基础表", "13", "directories、images、tags、image_tags、ai_analysis_results、image_versions", "保存核心业务数据和扩展元数据。"],
            ["视图", "4", "v_active_images、v_directory_stats、v_image_search、v_tag_stats", "统一查询入口，支撑统计和搜索。"],
            ["索引", "19", "idx_images_active、idx_tags_name_trgm、idx_ai_desc_trgm", "提升目录浏览、标签检索和描述检索速度。"],
            ["触发器", "5", "trg_image_after_insert、trg_tag_after_insert 等", "记录图片和标签变更。"],
            ["函数", "4", "fn_log_image_insert、fn_log_tag_change 等", "为触发器写入操作日志。"],
            ["存储过程", "5", "sp_monthly_report、sp_batch_rename、sp_restore_version 等", "提供报表和批处理设计扩展点。"],
        ],
    )
    add_heading(document, "3.6.2 数据一致性设计", 3)
    for text in [
        "图片路径、文件名和目录ID是连接磁盘与数据库的关键。单张重命名、批量重命名、删除和编辑保存都必须由服务层统一处理，不能让界面直接修改数据库记录，也不能只修改磁盘文件。",
        "标签和AI分析结果采用可复用结构。人工标签与识别标签最终都进入 tags 和 image_tags，描述文本进入 ai_analysis_results。普通关键词检索和智能检索都可以复用同一组持久化结果。",
        "版本历史采用独立版本文件和数据库记录共同保存。image_versions 记录版本号、路径、尺寸、缩略图、编辑类型和当前标记，.versions 目录保存具体文件。恢复版本时需要同时更新文件、images 表和当前版本标记。",
    ]:
        add_body(document, text)


def add_deployment_and_plan(document: Document) -> None:
    add_heading(document, "3.7 部署设计", 2)
    for text in [
        "系统部署为本机桌面应用。JavaFX 客户端运行在用户 Windows 环境中，PostgreSQL 可以运行在本机或可访问的局域网主机上，原始图片保存在本地文件系统，外部识别端点通过 HTTPS 作为可选增强服务接入。部署设计重点不在复杂服务器拓扑，而在普通机器上能否明确启动、配置、修复和验证。",
        "数据库初始化向导是部署设计的重要组成部分。若数据库不可用，用户可以在程序内填写服务器地址、端口、数据库名、用户名和密码，执行连接检测、数据库创建和 schema 初始化。初始化成功后，主界面刷新连接状态；初始化失败时，界面给出具体阶段和可重试入口。",
        "打包产物分为普通 JAR 和 Windows 便携包。JAR 面向已有 Java 环境，便携包面向未配置 Java 的机器。两种产物都进入同一套欢迎界面、主界面和数据库向导流程，避免开发环境路径被写死到运行逻辑中。",
    ]:
        add_body(document, text)
    add_heading(document, "3.7.1 编码测试阶段项目计划", 3)
    for text in [
        "根据实验说明的加分要求，实验3在设计文档之外保留编码和测试阶段 MPP 项目计划。计划按风险优先排序：先完成启动配置和数据库初始化，再完成目录扫描和缩略图缓存，随后完成浏览、搜索、批量重命名、编辑版本、标签和AI增强，最后完成打包交付与验证。",
        "MPP 文件中的任务标注了负责人、时间安排、产出和前置关系。该计划不是替代实验4测试报告，而是为后续编码和测试提供可执行的阶段安排。实验4只需要沿这些任务补充运行日志、测试结果、数据库查询和截图证据。",
    ]:
        add_body(document, text)
    add_figure(document, EVIDENCE / "图11_编码与测试阶段甘特图.png", "图17 编码与测试阶段甘特图", 15.2)


def add_design_figures(document: Document) -> None:
    ensure_figure12_export()
    figures = [
        ("图1 体系结构分层包图", EVIDENCE / "图1_体系结构分层包图.png"),
        ("图2 物理部署图", EVIDENCE / "图2_物理部署图.png"),
        ("图3 构件图", EVIDENCE / "图3_构件图.png"),
        ("图4 用户界面跳转顺序图", EVIDENCE / "图4_用户界面跳转顺序图.png"),
        ("图5 缩略图预览用例实现顺序图", EVIDENCE / "图5_缩略图预览用例实现顺序图.png"),
        ("图6 批量重命名用例实现顺序图", EVIDENCE / "图6_批量重命名用例实现顺序图.png"),
        ("图7 AI标签扫描与智能搜索用例实现顺序图", EVIDENCE / "图7_AI标签扫描与智能搜索用例实现顺序图.png"),
        ("图8 核心设计类图", EVIDENCE / "图8_核心设计类图.png"),
        ("图9 数据模型设计图", EVIDENCE / "图9_数据模型设计图.png"),
        ("图10 批量重命名事务活动图", EVIDENCE / "图10_批量重命名事务活动图.png"),
        ("图11 目录扫描入库详细活动图", VISIO_EXPORTS / "图11_目录扫描入库详细活动图.png"),
        ("图12 数据库导出表关系图", VISIO_EXPORTS / "图12_数据库导出表关系图.png"),
        ("图13 子系统与构件职责展开图", VISIO_EXPORTS / "图13_子系统与构件职责展开图.png"),
        ("图14 AI标签扫描与智能搜索安全链路图", VISIO_EXPORTS / "图14_AI标签扫描与智能搜索安全链路图.png"),
        ("图15 数据库初始化与离线降级流程图", VISIO_EXPORTS / "图15_数据库初始化与离线降级流程图.png"),
        ("图16 图片编辑与版本历史详细设计图", VISIO_EXPORTS / "图16_图片编辑与版本历史详细设计图.png"),
    ]
    standalone_figures = {
        "图4 用户界面跳转顺序图",
        "图5 缩略图预览用例实现顺序图",
        "图6 批量重命名用例实现顺序图",
        "图7 AI标签扫描与智能搜索用例实现顺序图",
        "图14 AI标签扫描与智能搜索安全链路图",
        "图16 图片编辑与版本历史详细设计图",
    }
    for caption, image in figures:
        if caption in standalone_figures:
            document.add_page_break()
        add_figure(document, image, caption, 15.2)


def add_review_mapping(document: Document) -> None:
    add_heading(document, "3.8 设计评审与课程要求对应", 2)
    add_body(
        document,
        "本节用于把实验说明中的关键要求落实到本报告。体系结构设计已经给出分层包图、部署图和构件图；界面设计给出反映窗口跳转关系的顺序图和当前运行截图；详细设计给出精化后的用例顺序、类设计、方法级活动说明、数据库表和子系统构件设计；文档变更历史记录已在模板第二页填写；每张图均在图题后标注绘制者和审查者。",
    )
    add_table(
        document,
        ["实验3要求", "本文档落实位置", "对应交付文件"],
        [
            ["体系结构设计：包图、部署图、构件图", "3.1 软件体系结构设计", "图1、图2、图3 的 VSDX 与正文。"],
            ["界面设计：截图跳转关系顺序图", "3.2 用户界面设计", "图4、运行截图、界面说明。"],
            ["用例设计：精化后的顺序图", "3.3 用例设计", "图5、图6、图7 及用例说明。"],
            ["类设计：类图、方法和算法说明", "3.5 类设计", "图8、图10、图16 和类方法说明。"],
            ["数据设计：数据库表", "3.6 数据设计", "图9、图12、schema.sql 对象说明。"],
            ["子系统和构件设计", "3.4 子系统与构件设计", "图3、图13 及构件职责说明。"],
            ["评审修订记录", "封面后变更历史页", "DOCX 模板第二页。"],
            ["编码测试阶段计划", "3.7.1 编码测试阶段项目计划", "数字图像管理系统编码测试阶段项目计划.mpp。"],
        ],
    )


def build_docx() -> None:
    ensure_template_docx()
    document = Document(str(TEMPLATE_DOCX))
    fill_cover_and_history(document)
    remove_template_body(document)
    add_intro(document)
    add_constraints(document)
    add_design(document)
    FINAL_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(FINAL_DOCX))
    shutil.copy2(FINAL_DOCX, LEGACY_DOCX)


def update_fields_and_export_pdf() -> None:
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(FINAL_DOCX.resolve()))
        doc.Fields.Update()
        for toc_idx in range(1, doc.TablesOfContents.Count + 1):
            doc.TablesOfContents(toc_idx).Update()
        doc.Save()
        doc.ExportAsFixedFormat(str(FINAL_PDF.resolve()), 17)
        doc.Close(False)
    finally:
        word.Quit()
    shutil.copy2(FINAL_PDF, LEGACY_PDF)


def render_review_pages() -> None:
    pdf = fitz.open(FINAL_PDF)
    for page_idx in range(min(4, pdf.page_count)):
        page = pdf[page_idx]
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
        out = REVIEW_DIR / f"{FINAL_BASENAME}_template_p{page_idx + 1}.png"
        pix.save(out)


def make_zip_and_sync() -> None:
    def figure_key(path: Path):
        match = re.match(r"图(\d+)", path.name)
        return (int(match.group(1)) if match else 999, path.name)

    vsdx_files = sorted(EXP3.glob("图*.vsdx"), key=figure_key)
    required = [FINAL_DOCX, FINAL_PDF, MPP, *vsdx_files]
    missing = [str(p) for p in required if not p.exists()]
    if missing:
        raise FileNotFoundError("缺少提交文件：" + "; ".join(missing))
    with ZipFile(FINAL_ZIP, "w", ZIP_DEFLATED) as zf:
        for path in required:
            zf.write(path, path.name)

    COLLECTED.mkdir(parents=True, exist_ok=True)
    for child in COLLECTED.iterdir():
        if child.is_dir():
            shutil.rmtree(child)
        else:
            child.unlink()
    for path in required:
        shutil.copy2(path, COLLECTED / path.name)


def inspect_outputs() -> dict[str, int | bool]:
    doc = Document(str(FINAL_DOCX))
    chars = sum(len(p.text.strip()) for p in doc.paragraphs)
    with ZipFile(FINAL_ZIP) as zf:
        names = zf.namelist()
    pdf = fitz.open(FINAL_PDF)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "chars": chars,
        "pdf_pages": pdf.page_count,
        "zip_entries": len(names),
        "zip_has_dirs": any(name.endswith("/") for name in names),
        "collected_dirs": sum(1 for p in COLLECTED.iterdir() if p.is_dir()),
        "collected_files": sum(1 for p in COLLECTED.iterdir() if p.is_file()),
        "vsdx_count": len(list(EXP3.glob("图*.vsdx"))),
    }


def main() -> None:
    build_docx()
    update_fields_and_export_pdf()
    render_review_pages()
    make_zip_and_sync()
    print(inspect_outputs())


if __name__ == "__main__":
    main()
